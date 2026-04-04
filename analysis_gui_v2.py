"""
SEISMIC BUILDING SURVEY VIEWER - Modern Edition
Glassmorphic design with embedded maps and building cards
Streamlined for survey data visualization and management
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox, StringVar, scrolledtext
import os
import pandas as pd
import numpy as np
import json
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import logging
import folium
import tempfile
import requests
from PIL import Image, ImageTk
from io import BytesIO
import threading
from docx import Document
from docx.shared import Inches, Pt, RGBColor

# --- Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Color Palette (Glassmorphic) ---
# Dark Blue Theme
DARK_BG = "#0F1419"
DARK_CARD = "#1A1F2E"
DARK_ACCENT = "#2D3748"
DARK_BORDER = "#4A5568"

# Light Blue Theme
LIGHT_BG = "#F0F4F8"
LIGHT_CARD = "#FFFFFF"
LIGHT_ACCENT = "#E2E8F0"
LIGHT_BORDER = "#CBD5E0"

# Accent Colors
PRIMARY = "#00D4FF"      # Cyan
SECONDARY = "#1E90FF"   # Dodger Blue
SUCCESS = "#00D084"     # Green
WARNING = "#FF9800"     # Orange
DANGER = "#FF6B6B"      # Red
TEXT_DARK = "#FFFFFF"
TEXT_LIGHT = "#1A202C"

# --- Data Models ---
@dataclass
class BuildingRecord:
    """Building survey record"""
    building_id: str
    data: Dict[str, Any]
    images: List[Dict[str, str]] = None  # List of {name, url}
    
    def __post_init__(self):
        if self.images is None:
            self.images = []

class ImageDownloader:
    """Handles batch image downloading"""
    
    @staticmethod
    def download_image(url: str, timeout: int = 5) -> Optional[Image.Image]:
        """Download image from URL"""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get(url, timeout=timeout, headers=headers)
            response.raise_for_status()
            img = Image.open(BytesIO(response.content))
            return img
        except Exception as e:
            logger.warning(f"Image download failed {url}: {e}")
            return None
    
    @staticmethod
    def batch_download(records: List[BuildingRecord], output_dir: str, progress_callback=None):
        """Download all images from records"""
        downloaded = 0
        for idx, record in enumerate(records):
            for img_info in record.images:
                url = img_info.get('url')
                name = img_info.get('name', 'image')
                if url:
                    img = ImageDownloader.download_image(url)
                    if img:
                        safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in name)
                        path = os.path.join(output_dir, f"{record.building_id}_{safe_name}.jpg")
                        img.save(path, "JPEG", quality=85)
                        downloaded += 1
            
            if progress_callback:
                progress_callback(idx + 1, len(records))
        
        return downloaded

# --- Main Application ---
class SeismicSurveyViewerApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        self.title("Seismic Building Survey Viewer")
        self.geometry("1700x1000")
        
        # --- State ---
        self.df = None
        self.filtered_df = None
        self.building_records: List[BuildingRecord] = []
        self.columns = []
        self.theme = "dark"  # dark or light
        self.current_building_index = 0
        
        # --- Theme ---
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        # --- Layout ---
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        self._create_sidebar()
        self._create_main_content()
        self._apply_glassmorphic_style()
    
    def _apply_glassmorphic_style(self):
        """Apply glassmorphic styling"""
        # Configure main window
        self.configure(fg_color=DARK_BG)
    
    def _create_sidebar(self):
        """Create left sidebar with navigation"""
        self.sidebar = ctk.CTkFrame(self, width=260, corner_radius=0, fg_color=DARK_CARD)
        self.sidebar.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        self.sidebar.grid_rowconfigure(10, weight=1)
        
        # Logo
        logo_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        logo_frame.pack(fill="x", padx=20, pady=(30, 10))
        
        title = ctk.CTkLabel(logo_frame, text="📊 SURVEY\nVIEWER", 
                            font=("Arial", 22, "bold"), text_color=PRIMARY)
        title.pack()
        
        subtitle = ctk.CTkLabel(logo_frame, text="Building Data Explorer",
                               font=("Arial", 10), text_color=DARK_BORDER)
        subtitle.pack()
        
        # Divider
        divider = ctk.CTkFrame(self.sidebar, height=1, fg_color=DARK_BORDER)
        divider.pack(fill="x", padx=20, pady=20)
        
        # Navigation Section
        nav_label = ctk.CTkLabel(self.sidebar, text="NAVIGATION", 
                                font=("Arial", 10, "bold"), text_color=PRIMARY)
        nav_label.pack(anchor="w", padx=20, pady=(10, 5))
        
        # Nav buttons with hover effect
        nav_buttons = [
            ("📂 Load Data", self._load_data_dialog, PRIMARY),
            ("🏢 Building Cards", self._show_building_cards, SECONDARY),
            ("🗺️ Map View", self._show_map_tab, PRIMARY),
            ("📊 Analytics", self._show_analytics, SUCCESS),
            ("📥 Bulk Download", self._bulk_download_images, WARNING),
            ("📄 Export Passports", self._export_passports, PRIMARY),
        ]
        
        for text, cmd, color in nav_buttons:
            btn = self._create_nav_button(text, cmd, color)
            btn.pack(fill="x", padx=15, pady=5)
        
        # Divider
        divider2 = ctk.CTkFrame(self.sidebar, height=1, fg_color=DARK_BORDER)
        divider2.pack(fill="x", padx=20, pady=20)
        
        # Filter Section - 3-Step Query Builder
        filter_label = ctk.CTkLabel(self.sidebar, text="QUICK FILTER",
                                   font=("Arial", 10, "bold"), text_color=PRIMARY)
        filter_label.pack(anchor="w", padx=20, pady=(10, 5))
        
        # Step 1: Field selection
        ctk.CTkLabel(self.sidebar, text="Field", font=("Arial", 9), text_color=DARK_BORDER).pack(anchor="w", padx=15, pady=(5, 2))
        self.filter_column = ctk.CTkComboBox(self.sidebar, values=[], 
                                            fg_color=DARK_ACCENT, border_color=DARK_BORDER,
                                            button_color=PRIMARY, dropdown_fg_color=DARK_CARD,
                                            command=self._on_filter_column_change)
        self.filter_column.pack(fill="x", padx=15, pady=2)
        self.filter_column.set("Select Field")
        
        # Step 2: Operator selection
        ctk.CTkLabel(self.sidebar, text="Operator", font=("Arial", 9), text_color=DARK_BORDER).pack(anchor="w", padx=15, pady=(8, 2))
        self.filter_operator = ctk.CTkComboBox(self.sidebar, 
                                              values=["=", "!=", ">", "<", ">=", "<=", "contains"],
                                              fg_color=DARK_ACCENT, border_color=DARK_BORDER,
                                              button_color=PRIMARY, dropdown_fg_color=DARK_CARD)
        self.filter_operator.pack(fill="x", padx=15, pady=2)
        self.filter_operator.set("=")
        
        # Step 3: Value selection
        ctk.CTkLabel(self.sidebar, text="Value", font=("Arial", 9), text_color=DARK_BORDER).pack(anchor="w", padx=15, pady=(8, 2))
        self.filter_value = ctk.CTkComboBox(self.sidebar, values=[],
                                           fg_color=DARK_ACCENT, border_color=DARK_BORDER,
                                           button_color=PRIMARY, dropdown_fg_color=DARK_CARD)
        self.filter_value.pack(fill="x", padx=15, pady=2)
        self.filter_value.set("Select Value")
        
        # Filter buttons
        btn_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        btn_frame.pack(fill="x", padx=15, pady=(10, 5))
        
        filter_btn = ctk.CTkButton(btn_frame, text="🔍 Apply", fg_color=SUCCESS,
                                  hover_color="#00B070", command=self._apply_filter, height=32)
        filter_btn.pack(side="left", fill="both", expand=True, padx=(0, 5))
        
        reset_btn = ctk.CTkButton(btn_frame, text="↺ Reset", fg_color=DARK_ACCENT,
                                 hover_color=DARK_BORDER, command=self._reset_filter, height=32)
        reset_btn.pack(side="left", fill="both", expand=True)
        
        # Status
        self.sidebar.status_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        self.sidebar.status_frame.pack(fill="x", padx=20, pady=(40, 20), side="bottom")
        
        self.status_label = ctk.CTkLabel(self.sidebar.status_frame, 
                                        text="No Data Loaded",
                                        font=("Arial", 10), text_color=DARK_BORDER)
        self.status_label.pack(anchor="w", pady=5)
        
        self.record_count = ctk.CTkLabel(self.sidebar.status_frame,
                                        text="Buildings: 0",
                                        font=("Arial", 10, "bold"), text_color=PRIMARY)
        self.record_count.pack(anchor="w", pady=2)
    
    def _create_nav_button(self, text: str, command, color: str) -> ctk.CTkButton:
        """Create navigation button with hover effect"""
        btn = ctk.CTkButton(self.sidebar, text=text, command=command,
                           fg_color=DARK_ACCENT, hover_color=color,
                           text_color=TEXT_DARK, corner_radius=8,
                           font=("Arial", 11), height=40)
        return btn
    
    def _create_main_content(self):
        """Create main content area"""
        self.main_frame = ctk.CTkFrame(self, fg_color=DARK_BG)
        self.main_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.main_frame.grid_columnconfigure(0, weight=1)
        self.main_frame.grid_rowconfigure(0, weight=1)
        
        # Tab view
        self.tabs = ctk.CTkTabview(self.main_frame, fg_color=DARK_CARD, 
                                   text_color=TEXT_DARK, segmented_button_fg_color=DARK_ACCENT,
                                   segmented_button_selected_color=PRIMARY)
        self.tabs.grid(row=0, column=0, sticky="nsew")
        
        # Home tab
        self.tab_home = self.tabs.add("🏠 Home")
        self._setup_home_tab()
        
        # Building Cards tab
        self.tab_cards = self.tabs.add("🏢 Building Cards")
        self._setup_cards_tab()
        
        # Map tab
        self.tab_map = self.tabs.add("🗺️ Map View")
        self._setup_map_tab()
        
        # Analytics tab
        self.tab_analytics = self.tabs.add("📊 Analytics")
        self._setup_analytics_tab()
    
    def _setup_home_tab(self):
        """Setup home/welcome tab"""
        self.tab_home.grid_columnconfigure(0, weight=1)
        self.tab_home.grid_rowconfigure(0, weight=1)
        
        # Welcome card
        welcome_frame = ctk.CTkFrame(self.tab_home, fg_color=DARK_ACCENT, corner_radius=15)
        welcome_frame.place(relx=0.5, rely=0.5, anchor="center", relwidth=0.8, relheight=0.6)
        
        welcome_title = ctk.CTkLabel(welcome_frame, text="Welcome to Survey Viewer",
                                    font=("Arial", 28, "bold"), text_color=PRIMARY)
        welcome_title.pack(pady=(40, 10))
        
        welcome_text = ctk.CTkLabel(welcome_frame,
            text="A modern platform for exploring seismic building survey data.\n\n"
                 "Features:\n"
                 "✓ Load building survey data from Excel files\n"
                 "✓ View building cards with detailed information\n"
                 "✓ Visualize location on interactive maps\n"
                 "✓ Download building images in bulk\n"
                 "✓ Generate building passports\n\n"
                 "Get started by loading your survey data file →",
            font=("Arial", 13), text_color=TEXT_DARK, justify="left")
        welcome_text.pack(pady=20, padx=40)
        
        # Quick start button
        start_btn = ctk.CTkButton(welcome_frame, text="📂 Load Survey File", 
                                 fg_color=PRIMARY, hover_color="#00BBDD",
                                 font=("Arial", 14, "bold"), height=50,
                                 command=self._load_data_dialog)
        start_btn.pack(pady=(20, 40))
    
    def _setup_cards_tab(self):
        """Setup building cards display tab"""
        self.tab_cards.grid_columnconfigure(0, weight=1)
        self.tab_cards.grid_rowconfigure(0, weight=1)
        
        # Create scrollable frame for cards
        self.cards_scroll = ctk.CTkScrollableFrame(self.tab_cards, fg_color="transparent")
        self.cards_scroll.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        empty_label = ctk.CTkLabel(self.cards_scroll, 
                                  text="Load data to view building cards",
                                  font=("Arial", 14), text_color=DARK_BORDER)
        empty_label.pack(pady=40)
        self._cards_empty = True
    
    def _setup_map_tab(self):
        """Setup embedded map tab"""
        self.tab_map.grid_columnconfigure(0, weight=1)
        self.tab_map.grid_rowconfigure(0, weight=1)
        
        map_container = ctk.CTkFrame(self.tab_map, fg_color=DARK_CARD, corner_radius=15)
        map_container.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        map_container.grid_columnconfigure(0, weight=1)
        map_container.grid_rowconfigure(1, weight=1)
        
        # Header
        map_header = ctk.CTkFrame(map_container, fg_color=DARK_ACCENT, corner_radius=15)
        map_header.grid(row=0, column=0, sticky="ew", padx=10, pady=10)
        
        ctk.CTkLabel(map_header, text="🗺️ Interactive Building Map",
                    font=("Arial", 16, "bold"), text_color=PRIMARY).pack(side="left", padx=20, pady=10)
        
        gen_map_btn = ctk.CTkButton(map_header, text="Generate Map", fg_color=PRIMARY,
                                   hover_color="#00BBDD", command=self._generate_map)
        gen_map_btn.pack(side="right", padx=20, pady=10)
        
        # Map display area
        self.map_frame = ctk.CTkFrame(map_container, fg_color=DARK_ACCENT, corner_radius=15)
        self.map_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        
        map_placeholder = ctk.CTkLabel(self.map_frame,
                                      text="📍 Click 'Generate Map' to visualize building locations\n"
                                           "(Requires Latitude/Longitude columns)",
                                      font=("Arial", 13), text_color=DARK_BORDER)
        map_placeholder.place(relx=0.5, rely=0.5, anchor="center")
    
    def _setup_analytics_tab(self):
        """Setup analytics/charts tab"""
        self.tab_analytics.grid_columnconfigure(0, weight=1)
        self.tab_analytics.grid_rowconfigure(0, weight=1)
        
        analytics_frame = ctk.CTkFrame(self.tab_analytics, fg_color=DARK_CARD, corner_radius=15)
        analytics_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        ctk.CTkLabel(analytics_frame, text="📊 Survey Analytics",
                    font=("Arial", 16, "bold"), text_color=PRIMARY).pack(padx=20, pady=(20, 10), anchor="w")
        
        # Stats grid
        stats_frame = ctk.CTkFrame(analytics_frame, fg_color=DARK_ACCENT, corner_radius=10)
        stats_frame.pack(fill="x", padx=20, pady=10)
        
        self.stat_labels = {}
        for key in ['total', 'avg_age', 'materials', 'locations']:
            stat_container = ctk.CTkFrame(stats_frame, fg_color=DARK_BORDER, corner_radius=8)
            stat_container.pack(side="left", fill="both", expand=True, padx=5, pady=10)
            
            title = ctk.CTkLabel(stat_container, text=key.replace('_', ' ').title(),
                               font=("Arial", 10), text_color=DARK_BORDER)
            title.pack(padx=10, pady=(5, 0))
            
            value = ctk.CTkLabel(stat_container, text="--", 
                               font=("Arial", 18, "bold"), text_color=PRIMARY)
            value.pack(padx=10, pady=(0, 5))
            
            self.stat_labels[key] = value
    
    def _load_data_dialog(self):
        """Open file dialog to load data"""
        file_path = filedialog.askopenfilename(
            title="Select Survey Data File",
            filetypes=[("Excel Files", "*.xlsx *.xls"), ("CSV Files", "*.csv"), ("All Files", "*.*")]
        )
        
        if file_path:
            self._load_data(file_path)
    
    def _load_data(self, file_path: str):
        """Load data from Excel/CSV - handles both normal and transposed formats"""
        try:
            # Load file
            if file_path.endswith('.csv'):
                df_raw = pd.read_csv(file_path)
            else:
                df_raw = pd.read_excel(file_path)
            
            # Detect if data is transposed format
            # Transposed: First column has field names, first row has building IDs or empty
            # Normal: First row has field names, first column has building IDs
            
            is_transposed = False
            
            # Check if first column contains common field keywords
            first_col = df_raw.iloc[:, 0].astype(str).str.lower()
            field_keywords = ['building', 'id', 'age', 'material', 'surveyor', 'gps', 'latitude', 'longitude', 'region', 'district']
            
            if any(keyword in first_col.values for keyword in field_keywords):
                is_transposed = True
            
            # Convert if transposed
            if is_transposed:
                # Transpose: rows become columns
                df_raw_t = df_raw.set_index(df_raw.columns[0]).T
                df_raw_t = df_raw_t.reset_index()
                df_raw_t.rename(columns={'index': 'Survey_ID'}, inplace=True)
                self.df = df_raw_t
            else:
                self.df = df_raw
            
            self.filtered_df = self.df.copy()
            self.columns = list(self.df.columns)
            
            # Update filter column options
            self.filter_column.configure(values=self.columns)
            if self.columns:
                self.filter_column.set(self.columns[0])
            
            # Parse records (one per row after transpose)
            self._parse_records()
            
            # Update UI
            self.status_label.configure(text=os.path.basename(file_path))
            self.record_count.configure(text=f"Buildings: {len(self.building_records)}")
            
            # Populate UI
            self._refresh_cards()
            self._update_stats()
            
            # Go to cards tab
            self.tabs.set("🏢 Building Cards")
            
            messagebox.showinfo("Success", f"Loaded {len(self.building_records)} building records!")
            
        except Exception as e:
            logger.error(f"Load error: {e}", exc_info=True)
            messagebox.showerror("Load Error", f"Failed to load file:\n{str(e)}")
    
    def _on_filter_column_change(self, choice: str):
        """Update value dropdown when column selection changes"""
        if choice == "Select Field" or not choice or choice not in self.columns:
            self.filter_value.configure(values=[])
            self.filter_value.set("Select Value")
            return
        
        try:
            # Get unique values from selected column
            unique_values = self.df[choice].dropna().unique()
            # Convert to strings and limit to 100 for performance
            value_list = [str(v) for v in unique_values[:100]]
            self.filter_value.configure(values=sorted(value_list))
            self.filter_value.set("Select Value")
        except Exception as e:
            logger.warning(f"Could not populate filter values: {e}")
    
    def _parse_records(self):
        """Parse DataFrame rows into BuildingRecord objects"""
        self.building_records = []
        
        for idx, row in self.df.iterrows():
            # Find or generate building ID
            id_col = self._find_column(['id', 'building_id', 'building', 'bid', 'building code'])
            building_id = str(row.get(id_col, f"Building_{idx}")) if id_col else f"Building_{idx}"
            
            # Extract images/URLs
            images = []
            for col_name, value in row.items():
                if pd.isna(value):
                    continue
                value_str = str(value)
                if 'http' in value_str.lower() or '.jpg' in value_str.lower() or '.png' in value_str.lower():
                    caption = col_name.replace('_', ' ').strip()
                    images.append({'name': caption or 'Image', 'url': value_str})
            
            record = BuildingRecord(
                building_id=building_id,
                data=row.to_dict(),
                images=images
            )
            
            self.building_records.append(record)
    
    def _find_column(self, keywords: List[str]) -> Optional[str]:
        """Find column by keywords"""
        for col in self.columns:
            if any(kw in col.lower() for kw in keywords):
                return col
        return None
    
    def _refresh_cards(self):
        """Refresh building cards display with collapsible UI"""
        # Clear old cards
        for widget in self.cards_scroll.winfo_children():
            widget.destroy()
        
        if self.filtered_df.empty:
            empty = ctk.CTkLabel(self.cards_scroll, text="No buildings match filter",
                               font=("Arial", 13), text_color=DARK_BORDER)
            empty.pack(pady=40)
            return
        
        # Get important column names for collapsed view
        building_id_col = self._find_column(['id', 'building_id', 'building', 'bid', 'building code'])
        surveyor_col = self._find_column(['surveyor', 'surveyor_name', 'surveyed by'])
        region_col = self._find_column(['region', 'district', 'area', 'location'])
        gps_col = 'GPS COORDINATES'  # Explicit column name
        
        # Create cards for each row
        for idx, row in self.filtered_df.iterrows():
            if idx >= len(self.building_records):
                break
            
            record = self.building_records[idx]
            self._create_collapsible_building_card(record, row, building_id_col, surveyor_col, region_col, gps_col)
    
    def _create_collapsible_building_card(self, record: BuildingRecord, row: pd.Series, 
                                         building_id_col, surveyor_col, region_col, gps_col):
        """Create a single collapsible building card"""
        
        # Outer card container
        card_container = ctk.CTkFrame(self.cards_scroll, fg_color=DARK_CARD, corner_radius=15)
        card_container.pack(fill="x", padx=10, pady=8)
        
        # Collapsed header (always visible)
        header = ctk.CTkFrame(card_container, fg_color=DARK_ACCENT, corner_radius=12)
        header.pack(fill="x", padx=10, pady=10)
        
        # Create state variable for expand/collapse
        expand_state = {"expanded": False}
        details_frame_ref = {"frame": None}
        
        def toggle_expand():
            """Toggle card expansion"""
            expand_state["expanded"] = not expand_state["expanded"]
            
            if expand_state["expanded"]:
                # Expand - show details
                expand_btn.configure(text="▼")
                details_frame_ref["frame"].pack(fill="x", padx=15, pady=10, after=header)
            else:
                # Collapse - hide details
                expand_btn.configure(text="▶")
                details_frame_ref["frame"].pack_forget()
        
        # Expand/Collapse button
        expand_btn = ctk.CTkButton(header, text="▶", font=("Arial", 12), width=30, height=30,
                                  fg_color=PRIMARY, hover_color="#00BBDD",
                                  command=toggle_expand)
        expand_btn.pack(side="left", padx=(10, 15), pady=8)
        
        # Header info (collapsed view)
        info_frame = ctk.CTkFrame(header, fg_color="transparent")
        info_frame.pack(side="left", fill="x", expand=True, padx=5, pady=8)
        
        # Building ID
        building_id = str(row.get(building_id_col, record.building_id)) if building_id_col else record.building_id
        id_label = ctk.CTkLabel(info_frame, text=f"🏢 {building_id}", 
                               font=("Arial", 13, "bold"), text_color=PRIMARY)
        id_label.pack(anchor="w", pady=2)
        
        # Quick info row (Surveyor | Region | GPS)
        quick_info = ctk.CTkFrame(info_frame, fg_color="transparent")
        quick_info.pack(anchor="w", fill="x", pady=(5, 0))
        
        if surveyor_col and surveyor_col in row.index:
            surveyor = str(row.get(surveyor_col, "N/A"))
            ctk.CTkLabel(quick_info, text=f"👤 {surveyor}", font=("Arial", 9), text_color=DARK_BORDER).pack(side="left", padx=(0, 15))
        
        if region_col and region_col in row.index:
            region = str(row.get(region_col, "N/A"))
            ctk.CTkLabel(quick_info, text=f"📍 {region}", font=("Arial", 9), text_color=DARK_BORDER).pack(side="left", padx=(0, 15))
        
        if gps_col in row.index and pd.notna(row.get(gps_col)):
            gps_str = str(row.get(gps_col, "N/A"))
            ctk.CTkLabel(quick_info, text=f"🛰️ {gps_str}", font=("Arial", 9), text_color=DARK_BORDER).pack(side="left")
        
        # Expanded details (hidden by default)
        details_frame = ctk.CTkFrame(card_container, fg_color="transparent")
        details_frame_ref["frame"] = details_frame
        
        # Details content
        details_scroll = ctk.CTkScrollableFrame(details_frame, fg_color="transparent", height=200)
        details_scroll.pack(fill="both", expand=True)
        
        # Display all fields in expanded view
        displayed = 0
        for col_name, value in row.items():
            if pd.isna(value) or displayed >= 30:  # Limit fields shown
                continue
            
            # Skip image URLs in details
            if 'http' in str(value).lower() or '.jpg' in str(value).lower():
                continue
            
            row_frame = ctk.CTkFrame(details_scroll, fg_color="transparent")
            row_frame.pack(fill="x", pady=2, padx=5)
            
            # Field name
            label = ctk.CTkLabel(row_frame, text=f"{col_name}:",
                                font=("Arial", 9, "bold"), text_color=DARK_BORDER, width=150, anchor="w")
            label.pack(side="left", padx=(5, 10))
            
            # Field value
            val_text = str(value)[:60]
            val = ctk.CTkLabel(row_frame, text=val_text,
                              font=("Arial", 9), text_color=TEXT_DARK)
            val.pack(side="left", fill="x", expand=True)
            
            displayed += 1
        
        # Image buttons in expanded view
        if record.images:
            img_divider = ctk.CTkFrame(details_frame, height=1, fg_color=DARK_BORDER)
            img_divider.pack(fill="x", padx=15, pady=10)
            
            img_frame = ctk.CTkFrame(details_frame, fg_color="transparent")
            img_frame.pack(fill="x", padx=15, pady=(0, 10))
            
            ctk.CTkLabel(img_frame, text="Images:", font=("Arial", 10, "bold"), text_color=PRIMARY).pack(anchor="w", pady=(5, 5))
            
            for img_info in record.images[:3]:  # Show max 3
                img_btn = ctk.CTkButton(img_frame, text=f"🖼️ {img_info['name']}",
                                       fg_color=SECONDARY, hover_color="#1870FF", height=32,
                                       command=lambda url=img_info['url']: self._open_image(url))
                img_btn.pack(fill="x", pady=3)
    
    def _open_image(self, url: str):
        """Open image in browser or viewer"""
        try:
            import webbrowser
            webbrowser.open(url)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open image:\n{str(e)}")
    
    def _apply_filter(self):
        """Apply 3-step query builder filter"""
        if self.df is None:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        col = self.filter_column.get()
        operator = self.filter_operator.get()
        val = self.filter_value.get()
        
        if col == "Select Field" or val == "Select Value" or not col or not val or col not in self.columns:
            messagebox.showwarning("Invalid Filter", "Select field, operator, and value")
            return
        
        try:
            temp_df = self.df.copy()
            col_data = temp_df[col]
            
            if operator == "=":
                self.filtered_df = temp_df[col_data.astype(str) == val]
            elif operator == "!=":
                self.filtered_df = temp_df[col_data.astype(str) != val]
            elif operator == "contains":
                self.filtered_df = temp_df[col_data.astype(str).str.contains(val, case=False, na=False)]
            else:
                # For numeric operators, try to convert
                try:
                    col_numeric = pd.to_numeric(col_data, errors='coerce')
                    val_numeric = float(val)
                    
                    if operator == ">":
                        self.filtered_df = temp_df[col_numeric > val_numeric]
                    elif operator == "<":
                        self.filtered_df = temp_df[col_numeric < val_numeric]
                    elif operator == ">=":
                        self.filtered_df = temp_df[col_numeric >= val_numeric]
                    elif operator == "<=":
                        self.filtered_df = temp_df[col_numeric <= val_numeric]
                    else:
                        messagebox.showerror("Error", f"Unknown operator: {operator}")
                        return
                except ValueError:
                    messagebox.showerror("Type Error", f"Cannot use '{operator}' with non-numeric values")
                    return
            
            self._refresh_cards()
            self.record_count.configure(text=f"Buildings: {len(self.filtered_df)}")
            self.tabs.set("🏢 Building Cards")
            
        except Exception as e:
            logger.error(f"Filter error: {e}", exc_info=True)
            messagebox.showerror("Filter Error", str(e))
    
    def _reset_filter(self):
        """Reset filter"""
        if self.df is None:
            return
        self.filtered_df = self.df.copy()
        self.filter_value.delete(0, "end")
        self._refresh_cards()
        self.record_count.configure(text=f"Buildings: {len(self.building_records)}")
    
    def _generate_map(self):
        """Generate and display interactive map with smart coordinate detection - EMBEDDED in window"""
        if not self.building_records:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        try:
            lats = []
            lons = []
            location_data = []  # Store building info with coordinates
            
            # Strategy 1: Look for 'GPS COORDINATES' column/row with comma-separated values
            if 'GPS COORDINATES' in self.columns:
                for idx, record in enumerate(self.building_records):
                    try:
                        gps_str = record.data.get('GPS COORDINATES', "")
                        if pd.isna(gps_str) or not gps_str:
                            continue
                        
                        # Parse comma-separated lat,lon
                        gps_str = str(gps_str).strip()
                        coords = gps_str.split(',')
                        if len(coords) >= 2:
                            lat = float(coords[0].strip())
                            lon = float(coords[1].strip())
                            
                            if -90 <= lat <= 90 and -180 <= lon <= 180:
                                lats.append(lat)
                                lons.append(lon)
                                location_data.append({'building_id': record.building_id, 'lat': lat, 'lon': lon})
                    except (ValueError, IndexError, TypeError):
                        continue
            
            # Strategy 2: Look for separate Latitude/Longitude columns
            if not lats or not lons:
                lat_col = self._find_column(['latitude', 'lat', 'y'])
                lon_col = self._find_column(['longitude', 'lon', 'x'])
                
                if lat_col and lon_col:
                    for record in self.building_records:
                        try:
                            lat = float(record.data.get(lat_col, 0)) if lat_col in record.data else 0
                            lon = float(record.data.get(lon_col, 0)) if lon_col in record.data else 0
                            
                            if -90 <= lat <= 90 and -180 <= lon <= 180:
                                lats.append(lat)
                                lons.append(lon)
                                location_data.append({'building_id': record.building_id, 'lat': lat, 'lon': lon})
                        except (ValueError, TypeError):
                            continue
            
            # Strategy 3: Ask user to select columns
            if not lats or not lons:
                # Show dialog for manual selection
                select_window = ctk.CTkToplevel(self)
                select_window.title("Select Coordinate Columns")
                select_window.geometry("400x200")
                select_window.resizable(False, False)
                
                ctk.CTkLabel(select_window, text="Coordinate columns not detected.\nPlease select manually:",
                            font=("Arial", 12)).pack(pady=20)
                
                lat_frame = ctk.CTkFrame(select_window, fg_color="transparent")
                lat_frame.pack(fill="x", padx=20, pady=5)
                
                ctk.CTkLabel(lat_frame, text="Latitude:", width=100).pack(side="left", padx=5)
                lat_combo = ctk.CTkComboBox(lat_frame, values=self.columns, width=200)
                lat_combo.pack(side="left", padx=5)
                
                lon_frame = ctk.CTkFrame(select_window, fg_color="transparent")
                lon_frame.pack(fill="x", padx=20, pady=5)
                
                ctk.CTkLabel(lon_frame, text="Longitude:", width=100).pack(side="left", padx=5)
                lon_combo = ctk.CTkComboBox(lon_frame, values=self.columns, width=200)
                lon_combo.pack(side="left", padx=5)
                
                def proceed_map():
                    lat_col = lat_combo.get()
                    lon_col = lon_combo.get()
                    if not lat_col or not lon_col:
                        messagebox.showwarning("Selection Required", "Select both columns")
                        return
                    select_window.destroy()
                    self._create_map_with_columns(lat_col, lon_col)
                
                btn = ctk.CTkButton(select_window, text="Generate Map", command=proceed_map)
                btn.pack(pady=20)
                
                return
            
            # Create map with detected coordinates - EMBEDDED display
            self._display_map_embedded(lats, lons, location_data)
            
        except Exception as e:
            logger.error(f"Map error: {e}", exc_info=True)
            messagebox.showerror("Map Error", str(e))
    
    def _display_map_embedded(self, lats: list, lons: list, location_data: list):
        """Display map embedded in app window using canvas visualization"""
        try:
            # Clear map frame
            for widget in self.map_frame.winfo_children():
                widget.destroy()
            
            if not lats or not lons:
                messagebox.showerror("No Coordinates", "No valid coordinates found")
                return
            
            # Create scrollable frame for map info and controls
            map_container = ctk.CTkScrollableFrame(self.map_frame, fg_color="transparent")
            map_container.pack(fill="both", expand=True, padx=10, pady=10)
            
            # Map title and stats
            title_frame = ctk.CTkFrame(map_container, fg_color=DARK_ACCENT, corner_radius=10)
            title_frame.pack(fill="x", pady=(0, 15))
            
            ctk.CTkLabel(title_frame, text=f"🗺️ Building Locations ({len(location_data)} buildings)",
                        font=("Arial", 14, "bold"), text_color=PRIMARY).pack(anchor="w", padx=15, pady=10)
            
            min_lat, max_lat = min(lats), max(lats)
            min_lon, max_lon = min(lons), max(lons)
            center_lat = (min_lat + max_lat) / 2
            center_lon = (min_lon + max_lon) / 2
            
            stats = f"Center: ({center_lat:.4f}, {center_lon:.4f}) | Range: Lat [{min_lat:.2f}, {max_lat:.2f}] Lon [{min_lon:.2f}, {max_lon:.2f}]"
            ctk.CTkLabel(title_frame, text=stats, font=("Arial", 9), text_color=DARK_BORDER).pack(anchor="w", padx=15, pady=(0, 10))
            
            # Canvas for visual map representation
            canvas_frame = ctk.CTkFrame(map_container, fg_color=DARK_ACCENT, corner_radius=10, height=300)
            canvas_frame.pack(fill="x", pady=(0, 15))
            
            try:
                import tkinter as tk
                canvas = tk.Canvas(canvas_frame, width=600, height=300, bg='#2D3748', highlightthickness=0)
                canvas.pack(padx=10, pady=10)
                
                # Draw map background
                canvas.create_rectangle(0, 0, 600, 300, fill='#1A1F2E', outline=PRIMARY, width=2)
                
                # Normalize coordinates to canvas size
                if max_lon != min_lon and max_lat != min_lat:
                    lon_range = max_lon - min_lon
                    lat_range = max_lat - min_lat
                    
                    for loc in location_data:
                        # Map coordinates to canvas
                        x = 20 + ((loc['lon'] - min_lon) / lon_range) * 560
                        y = 280 - ((loc['lat'] - min_lat) / lat_range) * 260
                        
                        # Draw point
                        canvas.create_oval(x-4, y-4, x+4, y+4, fill=PRIMARY, outline=TEXT_DARK, width=2)
                        
                        # Tooltip on hover would require more complex code
            except Exception as e:
                logger.warning(f"Canvas rendering failed: {e}")
            
            # Location list with links
            locations_label = ctk.CTkLabel(map_container, text="Buildings and Coordinates:",
                                          font=("Arial", 12, "bold"), text_color=PRIMARY)
            locations_label.pack(anchor="w", pady=(15, 5))
            
            for loc in location_data[:50]:  # Show first 50
                loc_frame = ctk.CTkFrame(map_container, fg_color=DARK_ACCENT, corner_radius=8)
                loc_frame.pack(fill="x", pady=3)
                
                loc_text = f"🏢 {loc['building_id']} → ({loc['lat']:.6f}, {loc['lon']:.6f})"
                ctk.CTkLabel(loc_frame, text=loc_text, font=("Arial", 10), text_color=TEXT_DARK).pack(anchor="w", padx=10, pady=8)
            
            # Action buttons
            btn_frame = ctk.CTkFrame(map_container, fg_color="transparent")
            btn_frame.pack(fill="x", pady=15)
            
            # Generate full map file
            def generate_full_map():
                self._generate_full_map_file(lats, lons, location_data)
            
            gen_btn = ctk.CTkButton(btn_frame, text="🌐 Generate Full Interactive Map File",
                                   fg_color=PRIMARY, hover_color="#00BBDD",
                                   command=generate_full_map, height=40)
            gen_btn.pack(fill="x", pady=5)
            
            messagebox.showinfo("Map Embedded", f"Map displayed with {len(location_data)} buildings in the window!")
            
        except Exception as e:
            logger.error(f"Embedded map error: {e}", exc_info=True)
            messagebox.showerror("Map Error", str(e))
    
    def _generate_full_map_file(self, lats: list, lons: list, location_data: list):
        """Generate full Folium map file and open in browser"""
        try:
            center_lat = np.mean(lats)
            center_lon = np.mean(lons)
            
            m = folium.Map(location=[center_lat, center_lon], zoom_start=11)
            
            # Add markers
            for loc in location_data:
                popup_text = f"<b>{loc['building_id']}</b><br>Lat: {loc['lat']:.6f}<br>Lon: {loc['lon']:.6f}"
                
                folium.CircleMarker(
                    location=[loc['lat'], loc['lon']],
                    radius=6,
                    popup=folium.Popup(popup_text, max_width=250),
                    color=PRIMARY,
                    fill=True,
                    fillOpacity=0.7
                ).add_to(m)
            
            # Save
            html_file = tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False)
            m.save(html_file.name)
            
            # Open
            import webbrowser
            webbrowser.open('file://' + html_file.name)
            messagebox.showinfo("Success", f"Full map opened in browser!")
            
        except Exception as e:
            logger.error(f"Full map generation error: {e}", exc_info=True)
            messagebox.showerror("Error", str(e))
    
    def _show_building_cards(self):
        """Switch to building cards tab"""
        self.tabs.set("🏢 Building Cards")
    
    def _show_map_tab(self):
        """Switch to map tab"""
        self.tabs.set("🗺️ Map View")
    
    def _show_analytics(self):
        """Switch to analytics tab"""
        self.tabs.set("📊 Analytics")
        self._update_stats()
    
    def _update_stats(self):
        """Update analytics statistics"""
        if not self.building_records:
            return
        
        try:
            # Total buildings
            self.stat_labels['total'].configure(text=str(len(self.building_records)))
            
            # Average age
            age_col = self._find_column(['age', 'year', 'built'])
            if age_col:
                ages = pd.to_numeric(self.df[age_col], errors='coerce').dropna()
                if len(ages) > 0:
                    avg_age = ages.mean()
                    self.stat_labels['avg_age'].configure(text=f"{avg_age:.0f} yrs")
            
            # Material count
            mat_col = self._find_column(['material', 'type', 'construction'])
            if mat_col:
                unique_materials = self.df[mat_col].nunique()
                self.stat_labels['materials'].configure(text=str(unique_materials))
            
            # Locations
            lat_col = self._find_column(['latitude', 'lat'])
            if lat_col:
                valid_coords = self.df[lat_col].notna().sum()
                self.stat_labels['locations'].configure(text=str(valid_coords))
        
        except Exception as e:
            logger.warning(f"Stats update error: {e}")
    
    def _bulk_download_images(self):
        """Bulk download all images from records"""
        if not self.building_records:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        # Count total images
        total_images = sum(len(r.images) for r in self.building_records)
        if total_images == 0:
            messagebox.showinfo("No Images", "No image URLs found in data")
            return
        
        # Ask for output directory
        output_dir = filedialog.askdirectory(title="Select Output Directory for Images")
        if not output_dir:
            return
        
        # Show progress window
        progress_window = ctk.CTkToplevel(self)
        progress_window.title("Downloading Images")
        progress_window.geometry("400x150")
        progress_window.resizable(False, False)
        
        status_label = ctk.CTkLabel(progress_window, text="Downloading...",
                                   font=("Arial", 12))
        status_label.pack(pady=(20, 10))
        
        progress_bar = ctk.CTkProgressBar(progress_window, width=300)
        progress_bar.set(0)
        progress_bar.pack(pady=10)
        
        count_label = ctk.CTkLabel(progress_window, text="0 / 0",
                                  font=("Arial", 11))
        count_label.pack(pady=10)
        
        def update_progress(current, total):
            progress_bar.set(current / max(total, 1))
            count_label.configure(text=f"{current} / {total}")
            progress_window.update()
        
        # Download in thread
        def download_thread():
            downloaded = ImageDownloader.batch_download(
                self.building_records, output_dir, update_progress
            )
            progress_window.destroy()
            messagebox.showinfo("Complete", f"Downloaded {downloaded} images!")
        
        threading.Thread(target=download_thread, daemon=True).start()
    
    def _export_passports(self):
        """Export building passports to Word document"""
        if not self.building_records:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        
        if not save_path:
            return
        
        try:
            doc = Document()
            doc.add_heading('Building Survey Passports', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Total Buildings: {len(self.filtered_df)}')
            doc.add_page_break()
            
            id_col = self._find_column(['id', 'building_id']) or self.columns[0]
            
            for idx, row in self.filtered_df.iterrows():
                record = self.building_records[idx] if idx < len(self.building_records) else None
                
                # Building heading
                building_id = row.get(id_col, f"Building_{idx}")
                doc.add_heading(f'Building: {building_id}', level=1)
                
                # Data table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Field'
                hdr_cells[1].text = 'Value'
                
                for col, value in row.items():
                    if pd.notna(value) and 'http' not in str(value).lower():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(col)
                        row_cells[1].text = str(value)[:100]
                
                # Images
                if record and record.images:
                    doc.add_heading('Images', level=2)
                    for img_info in record.images:
                        p = doc.add_paragraph()
                        p.add_run(f"Link: ").bold = True
                        p.add_run(img_info['url']).italic = True
                
                doc.add_page_break()
            
            doc.save(save_path)
            messagebox.showinfo("Success", f"Exported {len(self.filtered_df)} passports!")
        
        except Exception as e:
            logger.error(f"Export error: {e}", exc_info=True)
            messagebox.showerror("Export Error", str(e))

# --- Main ---
if __name__ == "__main__":
    app = SeismicSurveyViewerApp()
    app.mainloop()
