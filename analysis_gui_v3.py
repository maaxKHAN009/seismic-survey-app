"""
SEISMIC BUILDING SURVEY VIEWER - Modern Edition
Glassmorphic design with embedded maps and building cards
Streamlined for survey data visualization and management
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox, StringVar, scrolledtext
import os
import pandas as pd
import numpy as np
import json
from typing import Dict, List, Optional, Any
from dataclasses import dataclass
from datetime import datetime
import logging
import folium
import tempfile
import requests
from PIL import Image, ImageTk
from io import BytesIO
import threading
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
import matplotlib
matplotlib.use('TkAgg')
import matplotlib.pyplot as plt
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
from matplotlib.figure import Figure
from matplotlib.widgets import RectangleSelector
import contextily as ctx

# --- Logging ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Color Palette (Glassmorphic) ---
# Dark Blue Theme
DARK_BG = "#0F1419"
DARK_CARD = "#1A1F2E"
DARK_ACCENT = "#2D3748"
DARK_BORDER = "#4A5568"

# Light Blue Theme
LIGHT_BG = "#F0F4F8"
LIGHT_CARD = "#FFFFFF"
LIGHT_ACCENT = "#E2E8F0"
LIGHT_BORDER = "#CBD5E0"

# Accent Colors
PRIMARY = "#00D4FF"      # Cyan
SECONDARY = "#1E90FF"   # Dodger Blue
SUCCESS = "#00D084"     # Green
WARNING = "#FF9800"     # Orange
DANGER = "#FF6B6B"      # Red
TEXT_DARK = "#FFFFFF"
TEXT_LIGHT = "#1A202C"

# --- Data Models ---
@dataclass
class BuildingRecord:
    """Building survey record"""
    building_id: str
    data: Dict[str, Any]
    images: List[Dict[str, str]] = None  # List of {name, url}
    
    def __post_init__(self):
        if self.images is None:
            self.images = []

class ImageDownloader:
    """Handles batch image downloading"""
    
    @staticmethod
    def download_image(url: str, timeout: int = 5) -> Optional[Image.Image]:
        """Download image from URL"""
        try:
            headers = {'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'}
            response = requests.get(url, timeout=timeout, headers=headers)
            response.raise_for_status()
            img = Image.open(BytesIO(response.content))
            return img
        except Exception as e:
            logger.warning(f"Image download failed {url}: {e}")
            return None
    
    @staticmethod
    def batch_download(records: List[BuildingRecord], output_dir: str, progress_callback=None):
        """Download all images from records"""
        downloaded = 0
        for idx, record in enumerate(records):
            for img_info in record.images:
                url = img_info.get('url')
                name = img_info.get('name', 'image')
                if url:
                    img = ImageDownloader.download_image(url)
                    if img:
                        safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in name)
                        path = os.path.join(output_dir, f"{record.building_id}_{safe_name}.jpg")
                        img.save(path, "JPEG", quality=85)
                        downloaded += 1
            
            if progress_callback:
                progress_callback(idx + 1, len(records))
        
        return downloaded

# --- Main Application ---
class SeismicSurveyViewerApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        self.title("Seismic Building Survey Viewer")
        self.geometry("1700x1000")
        
        # --- State ---
        self.df = None
        self.filtered_df = None
        self.building_records: List[BuildingRecord] = []
        self.columns = []
        self.theme = "dark"  # dark or light
        self.current_building_index = 0
        self.map_html_file = None  # Store map file path
        
        # --- Theme ---
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("blue")
        
        # --- Layout ---
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        self._create_sidebar()
        self._create_main_content()
        self._apply_glassmorphic_style()
    
    def _apply_glassmorphic_style(self):
        """Apply glassmorphic styling"""
        # Configure main window
        self.configure(fg_color=DARK_BG)
    
    def _create_sidebar(self):
        """Create left sidebar with navigation"""
        self.sidebar = ctk.CTkFrame(self, width=260, corner_radius=0, fg_color=DARK_CARD)
        self.sidebar.grid(row=0, column=0, sticky="nsew", padx=0, pady=0)
        self.sidebar.grid_rowconfigure(10, weight=1)
        
        # Logo
        logo_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        logo_frame.pack(fill="x", padx=20, pady=(30, 10))
        
        title = ctk.CTkLabel(logo_frame, text="📊 SURVEY\nVIEWER", 
                            font=("Arial", 22, "bold"), text_color=PRIMARY)
        title.pack()
        
        subtitle = ctk.CTkLabel(logo_frame, text="Building Data Explorer",
                               font=("Arial", 10), text_color=DARK_BORDER)
        subtitle.pack()
        
        # Divider
        divider = ctk.CTkFrame(self.sidebar, height=1, fg_color=DARK_BORDER)
        divider.pack(fill="x", padx=20, pady=20)
        
        # Navigation Section
        nav_label = ctk.CTkLabel(self.sidebar, text="NAVIGATION", 
                                font=("Arial", 10, "bold"), text_color=PRIMARY)
        nav_label.pack(anchor="w", padx=20, pady=(10, 5))
        
        # Nav buttons with hover effect
        nav_buttons = [
            ("📂 Load Data", self._load_data_dialog, PRIMARY),
            ("🏢 Building Cards", self._show_building_cards, SECONDARY),
            ("🗺️ Map View", self._show_map_tab, PRIMARY),
            ("📊 Analytics", self._show_analytics, SUCCESS),
            ("📥 Bulk Download", self._bulk_download_images, WARNING),
            ("📄 Export Passports", self._export_passports, PRIMARY),
        ]
        
        for text, cmd, color in nav_buttons:
            btn = self._create_nav_button(text, cmd, color)
            btn.pack(fill="x", padx=15, pady=5)
        
        # Divider
        divider2 = ctk.CTkFrame(self.sidebar, height=1, fg_color=DARK_BORDER)
        divider2.pack(fill="x", padx=20, pady=20)
        
        # Filter Section - 3-Step Query Builder
        filter_label = ctk.CTkLabel(self.sidebar, text="QUERY BUILDER",
                                   font=("Arial", 10, "bold"), text_color=PRIMARY)
        filter_label.pack(anchor="w", padx=20, pady=(10, 5))
        
        # Step 1: Select Field
        field_label = ctk.CTkLabel(self.sidebar, text="1️⃣ Select Field:",
                                  font=("Arial", 9, "bold"), text_color=SECONDARY)
        field_label.pack(anchor="w", padx=15, pady=(5, 2))
        
        self.filter_column = ctk.CTkComboBox(self.sidebar, values=[], 
                                            fg_color=DARK_ACCENT, border_color=DARK_BORDER,
                                            button_color=PRIMARY, dropdown_fg_color=DARK_CARD)
        self.filter_column.pack(fill="x", padx=15, pady=2)
        self.filter_column.set("Field")
        
        # Step 2: Select Operator
        operator_label = ctk.CTkLabel(self.sidebar, text="2️⃣ Select Operator:",
                                     font=("Arial", 9, "bold"), text_color=SECONDARY)
        operator_label.pack(anchor="w", padx=15, pady=(8, 2))
        
        self.filter_operator = ctk.CTkComboBox(self.sidebar, 
                                              values=["Contains", "Equals", "Starts With", "Ends With", "Not Contains"],
                                              fg_color=DARK_ACCENT, border_color=DARK_BORDER,
                                              button_color=PRIMARY, dropdown_fg_color=DARK_CARD)
        self.filter_operator.pack(fill="x", padx=15, pady=2)
        self.filter_operator.set("Contains")
        
        # Step 3: Enter Value
        value_label = ctk.CTkLabel(self.sidebar, text="3️⃣ Enter Value:",
                                  font=("Arial", 9, "bold"), text_color=SECONDARY)
        value_label.pack(anchor="w", padx=15, pady=(8, 2))
        
        self.filter_value = ctk.CTkEntry(self.sidebar, placeholder_text="Value to search...",
                                        fg_color=DARK_ACCENT, border_color=DARK_BORDER,
                                        text_color=TEXT_DARK)
        self.filter_value.pack(fill="x", padx=15, pady=2)
        
        filter_btn = ctk.CTkButton(self.sidebar, text="🔍 Apply Filter", fg_color=SUCCESS,
                                  hover_color="#00B070", command=self._apply_filter)
        filter_btn.pack(fill="x", padx=15, pady=(5, 3))
        
        reset_btn = ctk.CTkButton(self.sidebar, text="↺ Reset", fg_color=DARK_ACCENT,
                                 hover_color=DARK_BORDER, command=self._reset_filter)
        reset_btn.pack(fill="x", padx=15, pady=2)
        
        # Status
        self.sidebar.status_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        self.sidebar.status_frame.pack(fill="x", padx=20, pady=(40, 20), side="bottom")
        
        self.status_label = ctk.CTkLabel(self.sidebar.status_frame, 
                                        text="No Data Loaded",
                                        font=("Arial", 10), text_color=DARK_BORDER)
        self.status_label.pack(anchor="w", pady=5)
        
        self.record_count = ctk.CTkLabel(self.sidebar.status_frame,
                                        text="Buildings: 0",
                                        font=("Arial", 10, "bold"), text_color=PRIMARY)
        self.record_count.pack(anchor="w", pady=2)
    
    def _create_nav_button(self, text: str, command, color: str) -> ctk.CTkButton:
        """Create navigation button with hover effect"""
        btn = ctk.CTkButton(self.sidebar, text=text, command=command,
                           fg_color=DARK_ACCENT, hover_color=color,
                           text_color=TEXT_DARK, corner_radius=8,
                           font=("Arial", 11), height=40)
        return btn
    
    def _create_main_content(self):
        """Create main content area"""
        self.main_frame = ctk.CTkFrame(self, fg_color=DARK_BG)
        self.main_frame.grid(row=0, column=1, sticky="nsew", padx=20, pady=20)
        self.main_frame.grid_columnconfigure(0, weight=1)
        self.main_frame.grid_rowconfigure(0, weight=1)
        
        # Tab view
        self.tabs = ctk.CTkTabview(self.main_frame, fg_color=DARK_CARD, 
                                   text_color=TEXT_DARK, segmented_button_fg_color=DARK_ACCENT,
                                   segmented_button_selected_color=PRIMARY)
        self.tabs.grid(row=0, column=0, sticky="nsew")
        
        # Home tab
        self.tab_home = self.tabs.add("🏠 Home")
        self._setup_home_tab()
        
        # Building Cards tab
        self.tab_cards = self.tabs.add("🏢 Building Cards")
        self._setup_cards_tab()
        
        # Map tab
        self.tab_map = self.tabs.add("🗺️ Map View")
        self._setup_map_tab()
        
        # Analytics tab
        self.tab_analytics = self.tabs.add("📊 Analytics")
        self._setup_analytics_tab()
    
    def _setup_home_tab(self):
        """Setup home/welcome tab"""
        self.tab_home.grid_columnconfigure(0, weight=1)
        self.tab_home.grid_rowconfigure(0, weight=1)
        
        # Welcome card
        welcome_frame = ctk.CTkFrame(self.tab_home, fg_color=DARK_ACCENT, corner_radius=15)
        welcome_frame.place(relx=0.5, rely=0.5, anchor="center", relwidth=0.8, relheight=0.6)
        
        welcome_title = ctk.CTkLabel(welcome_frame, text="Welcome to Survey Viewer",
                                    font=("Arial", 28, "bold"), text_color=PRIMARY)
        welcome_title.pack(pady=(40, 10))
        
        welcome_text = ctk.CTkLabel(welcome_frame,
            text="A modern platform for exploring seismic building survey data.\n\n"
                 "Features:\n"
                 "✓ Load building survey data from Excel files\n"
                 "✓ View building cards with detailed information\n"
                 "✓ Visualize location on interactive maps\n"
                 "✓ Download building images in bulk\n"
                 "✓ Generate building passports\n\n"
                 "Get started by loading your survey data file →",
            font=("Arial", 13), text_color=TEXT_DARK, justify="left")
        welcome_text.pack(pady=20, padx=40)
        
        # Quick start button
        start_btn = ctk.CTkButton(welcome_frame, text="📂 Load Survey File", 
                                 fg_color=PRIMARY, hover_color="#00BBDD",
                                 font=("Arial", 14, "bold"), height=50,
                                 command=self._load_data_dialog)
        start_btn.pack(pady=(20, 40))
    
    def _setup_cards_tab(self):
        """Setup building cards display tab"""
        self.tab_cards.grid_columnconfigure(0, weight=1)
        self.tab_cards.grid_rowconfigure(0, weight=1)
        
        # Create scrollable frame for cards
        self.cards_scroll = ctk.CTkScrollableFrame(self.tab_cards, fg_color="transparent")
        self.cards_scroll.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        empty_label = ctk.CTkLabel(self.cards_scroll, 
                                  text="Load data to view building cards",
                                  font=("Arial", 14), text_color=DARK_BORDER)
        empty_label.pack(pady=40)
        self._cards_empty = True
    
    def _setup_map_tab(self):
        """Setup embedded map tab"""
        self.tab_map.grid_columnconfigure(0, weight=1)
        self.tab_map.grid_rowconfigure(0, weight=1)
        
        map_container = ctk.CTkFrame(self.tab_map, fg_color=DARK_CARD, corner_radius=15)
        map_container.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        map_container.grid_columnconfigure(0, weight=1)
        map_container.grid_rowconfigure(1, weight=1)
        
        # Header
        map_header = ctk.CTkFrame(map_container, fg_color=DARK_ACCENT, corner_radius=15)
        map_header.grid(row=0, column=0, sticky="ew", padx=10, pady=10)
        
        ctk.CTkLabel(map_header, text="🗺️ Interactive Building Map",
                    font=("Arial", 16, "bold"), text_color=PRIMARY).pack(side="left", padx=20, pady=10)
        
        gen_map_btn = ctk.CTkButton(map_header, text="Generate Map", fg_color=PRIMARY,
                                   hover_color="#00BBDD", command=self._generate_map)
        gen_map_btn.pack(side="right", padx=20, pady=10)
        
        # Map display area
        self.map_frame = ctk.CTkFrame(map_container, fg_color=DARK_ACCENT, corner_radius=15)
        self.map_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        
        map_placeholder = ctk.CTkLabel(self.map_frame,
                                      text="📍 Click 'Generate Map' to visualize building locations\n"
                                           "(Requires Latitude/Longitude columns)",
                                      font=("Arial", 13), text_color=DARK_BORDER)
        map_placeholder.place(relx=0.5, rely=0.5, anchor="center")
    
    def _setup_analytics_tab(self):
        """Setup analytics/charts tab with chart type and axis selectors"""
        self.tab_analytics.grid_columnconfigure(0, weight=1)
        self.tab_analytics.grid_rowconfigure(1, weight=1)
        
        analytics_frame = ctk.CTkFrame(self.tab_analytics, fg_color=DARK_CARD, corner_radius=15)
        analytics_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=(10, 10))
        
        ctk.CTkLabel(analytics_frame, text="📊 Survey Analytics & Charts",
                    font=("Arial", 16, "bold"), text_color=PRIMARY).pack(padx=20, pady=(15, 10), anchor="w")
        
        # Chart Configuration Frame
        config_frame = ctk.CTkFrame(analytics_frame, fg_color=DARK_ACCENT, corner_radius=10)
        config_frame.pack(fill="x", padx=20, pady=10)
        
        # Chart Type
        type_frame = ctk.CTkFrame(config_frame, fg_color="transparent")
        type_frame.pack(fill="x", padx=15, pady=10)
        
        ctk.CTkLabel(type_frame, text="Chart Type:", font=("Arial", 11, "bold"), text_color=SECONDARY).pack(side="left", padx=(0, 10))
        self.chart_type = ctk.CTkComboBox(type_frame, values=["Bar Chart", "Pie Chart", "Line Chart", "Scatter Plot", "Histogram"],
                                         fg_color=DARK_CARD, border_color=DARK_BORDER, button_color=PRIMARY, dropdown_fg_color=DARK_CARD)
        self.chart_type.pack(side="left", fill="x", expand=True, padx=(0, 15))
        self.chart_type.set("Bar Chart")
        
        # X-Axis Field
        x_frame = ctk.CTkFrame(config_frame, fg_color="transparent")
        x_frame.pack(fill="x", padx=15, pady=5)
        
        ctk.CTkLabel(x_frame, text="X-Axis Field:", font=("Arial", 11, "bold"), text_color=SECONDARY).pack(side="left", padx=(0, 10))
        self.chart_x_field = ctk.CTkComboBox(x_frame, values=[],
                                            fg_color=DARK_CARD, border_color=DARK_BORDER, button_color=PRIMARY, dropdown_fg_color=DARK_CARD)
        self.chart_x_field.pack(side="left", fill="x", expand=True, padx=(0, 15))
        self.chart_x_field.set("Field")
        
        # Y-Axis Field
        y_frame = ctk.CTkFrame(config_frame, fg_color="transparent")
        y_frame.pack(fill="x", padx=15, pady=5)
        
        ctk.CTkLabel(y_frame, text="Y-Axis Field:", font=("Arial", 11, "bold"), text_color=SECONDARY).pack(side="left", padx=(0, 10))
        self.chart_y_field = ctk.CTkComboBox(y_frame, values=[],
                                            fg_color=DARK_CARD, border_color=DARK_BORDER, button_color=PRIMARY, dropdown_fg_color=DARK_CARD)
        self.chart_y_field.pack(side="left", fill="x", expand=True, padx=(0, 15))
        self.chart_y_field.set("Field")
        
        # Generate Button
        btn_frame = ctk.CTkFrame(config_frame, fg_color="transparent")
        btn_frame.pack(fill="x", padx=15, pady=(10, 15))
        
        generate_btn = ctk.CTkButton(btn_frame, text="📈 Generate Chart", fg_color=PRIMARY,
                                    hover_color="#00BBDD", command=self._generate_chart)
        generate_btn.pack(side="left", padx=(0, 10))
        
        reset_chart_btn = ctk.CTkButton(btn_frame, text="↺ Clear", fg_color=DARK_ACCENT,
                                       hover_color=DARK_BORDER, command=self._clear_chart)
        reset_chart_btn.pack(side="left")
        
        # Chart Display Area
        self.chart_frame = ctk.CTkFrame(self.tab_analytics, fg_color=DARK_CARD, corner_radius=15)
        self.chart_frame.grid(row=1, column=0, sticky="nsew", padx=10, pady=(0, 10))
        
        chart_placeholder = ctk.CTkLabel(self.chart_frame,
                                        text="📉 Select fields and click 'Generate Chart' to display visualization",
                                        font=("Arial", 13), text_color=DARK_BORDER)
        chart_placeholder.place(relx=0.5, rely=0.5, anchor="center")
        
        self.stat_labels = {}
    
    def _load_data_dialog(self):
        """Open file dialog to load data"""
        file_path = filedialog.askopenfilename(
            title="Select Survey Data File",
            filetypes=[("Excel Files", "*.xlsx *.xls"), ("CSV Files", "*.csv"), ("All Files", "*.*")]
        )
        
        if file_path:
            self._load_data(file_path)
    
    def _load_data(self, file_path: str):
        """Load data from Excel/CSV with column-wise survey format"""
        try:
            # Load file
            if file_path.endswith('.csv'):
                df_raw = pd.read_csv(file_path)
                self.hyperlinks_dict = {}  # No hyperlinks in CSV
            else:
                df_raw = pd.read_excel(file_path)
                # Extract hyperlinks from Excel file using openpyxl
                self.hyperlinks_dict = self._extract_excel_hyperlinks(file_path)
                logger.info(f"Extracted {len(self.hyperlinks_dict)} hyperlinks from Excel file")
            
            # Column A contains field names, Columns B+ contain survey data
            # Set first column as index (field names)
            if df_raw.shape[1] > 1:
                field_col_name = df_raw.columns[0]
                df_raw = df_raw.set_index(field_col_name)
            
            self.df = df_raw.copy()
            self.filtered_df = self.df.copy()
            # Survey names are the column headers (B, C, D... columns)
            self.columns = list(self.df.columns)
            
            # Update filter column options (field names)
            field_names = list(self.df.index)
            self.filter_column.configure(values=field_names)
            if field_names:
                self.filter_column.set(field_names[0])
            
            # Update chart field options (field names)
            self.chart_x_field.configure(values=field_names)
            self.chart_y_field.configure(values=field_names)
            if len(field_names) > 0:
                self.chart_x_field.set(field_names[0])
            if len(field_names) > 1:
                self.chart_y_field.set(field_names[1])
            
            # Parse records
            self._parse_records()
            
            # Update UI
            self.status_label.configure(text=os.path.basename(file_path))
            self.record_count.configure(text=f"Buildings: {len(self.building_records)}")
            
            # Populate UI
            self._refresh_cards()
            self._update_stats()
            
            # Debug: Count total images found
            total_images = sum(len(r.images) for r in self.building_records)
            logger.info(f"Loaded {len(self.building_records)} buildings with {total_images} total images")
            
            # Go to cards tab
            self.tabs.set("🏢 Building Cards")
            
            messagebox.showinfo("Success", f"Loaded {len(self.building_records)} building records!\n({total_images} images found)")
            
        except Exception as e:
            logger.error(f"Load error: {e}", exc_info=True)
            messagebox.showerror("Load Error", f"Failed to load file:\n{str(e)}")
    
    def _extract_excel_hyperlinks(self, file_path: str) -> Dict[tuple, Dict[str, str]]:
        """Extract hyperlinks from Excel file using openpyxl, including display text"""
        hyperlinks = {}
        
        try:
            wb = load_workbook(file_path)
            ws = wb.active
            
            # Iterate through all cells to find hyperlinks
            for row_idx, row in enumerate(ws.iter_rows(), start=1):
                for col_idx, cell in enumerate(row, start=1):
                    if cell.hyperlink:
                        # Get both the URL and the display text
                        hyperlink_url = cell.hyperlink.target
                        cell_value = cell.value or ""  # Display text
                        key = (row_idx, col_idx)
                        hyperlinks[key] = {
                            'url': hyperlink_url,
                            'text': str(cell_value).strip()
                        }
                        logger.debug(f"Found hyperlink at R{row_idx}C{col_idx}: {cell_value} -> {hyperlink_url}")
            
            wb.close()
            logger.info(f"Total hyperlinks found: {len(hyperlinks)}")
            return hyperlinks
        
        except Exception as e:
            logger.error(f"Error extracting hyperlinks: {e}")
            return {}
    
    def _parse_records(self):
        """Parse DataFrame columns into BuildingRecord objects (each column = one survey)"""
        self.building_records = []
        
        # Get the field names (index) - these correspond to Excel rows
        field_names = list(self.df.index)
        
        # Each column in the DataFrame is one survey
        for col_idx, col_name in enumerate(self.df.columns):
            col_data = self.df[col_name]
            
            # Build the data dictionary for this survey
            data_dict = {}
            hyperlinks_for_column = {}  # Store hyperlinks for this column
            
            for field_idx, (field_name, value) in enumerate(col_data.items()):
                data_dict[field_name] = value
                
                # Map to Excel coordinates to find hyperlinks
                # Excel row = field_idx + 2 (row 1 is header, Excel is 1-indexed)
                # Excel col = col_idx + 2 (col 1 is field names)
                excel_row = field_idx + 2
                excel_col = col_idx + 2
                hyperlink_key = (excel_row, excel_col)
                
                if hyperlink_key in self.hyperlinks_dict:
                    hyperlinks_for_column[field_name] = self.hyperlinks_dict[hyperlink_key]
                    logger.debug(f"Found hyperlink for {field_name}: {self.hyperlinks_dict[hyperlink_key]}")
            
            # Store hyperlinks with the record
            data_dict['__hyperlinks__'] = hyperlinks_for_column
            
            # Get building ID - look for BUILDING ID field
            building_id = "Unknown"
            for key in data_dict:
                if 'building' in str(key).lower() and 'id' in str(key).lower():
                    building_id = str(data_dict[key])
                    break
            
            if not building_id or building_id == "Unknown":
                building_id = f"Survey_{col_idx + 1}"
            
            # Extract images from the data
            images = self._extract_image_urls_from_dict(data_dict)
            
            record = BuildingRecord(
                building_id=building_id,
                data=data_dict,
                images=images
            )
            
            self.building_records.append(record)
        
        # Store original records for filter reset
        self.original_records = self.building_records.copy()
    
    def _extract_image_urls(self, row: pd.Series) -> List[Dict[str, str]]:
        """Extract image URLs from row hyperlinks"""
        images = []
        
        # Look for columns with URLs or image-related data
        for col_name, value in row.items():
            if pd.isna(value):
                continue
            
            value_str = str(value)
            
            # Check if looks like URL
            if 'http' in value_str.lower():
                # Extract image caption from column name
                caption = col_name.replace('_', ' ').replace('http', '').strip()
                images.append({
                    'name': caption or 'Image',
                    'url': value_str
                })
            elif '.jpg' in value_str.lower() or '.png' in value_str.lower():
                caption = col_name.replace('_', ' ').strip()
                images.append({
                    'name': caption or 'Image',
                    'url': value_str
                })
        
        return images
    
    def _extract_image_urls_from_dict(self, data_dict: Dict[str, Any]) -> List[Dict[str, str]]:
        """Extract image URLs from data dictionary, prioritizing hyperlinks and their display text"""
        images = []
        
        # Extract hyperlinks dictionary (if present)
        hyperlinks = data_dict.get('__hyperlinks__', {})
        logger.debug(f"Extracting images with {len(hyperlinks)} hyperlinks available")
        
        for field_name, value in data_dict.items():
            # Skip the special hyperlinks dict entry
            if field_name == '__hyperlinks__':
                continue
            
            if pd.isna(value) or value is None or value == '':
                continue
            
            value_str = str(value).strip()
            field_name_lower = str(field_name).lower()
            
            # Check if field name suggests it's an image field
            is_image_field = any(keyword in field_name_lower for keyword in 
                               ['image', 'photo', 'picture', 'sketch', 'elevation', 'img', 
                                'url', 'link', 'hyperlink', 'jpeg', 'jpg', 'png', 'related'])
            
            # Check if hyperlink exists for this field
            has_hyperlink = field_name in hyperlinks
            
            # If image field or has hyperlink, extract URL
            if is_image_field or has_hyperlink:
                # Prefer hyperlink if available
                url = None
                caption = None
                
                if has_hyperlink:
                    hyperlink_data = hyperlinks[field_name]
                    url = hyperlink_data.get('url')
                    caption = hyperlink_data.get('text')  # Use display text as caption
                    logger.debug(f"Using hyperlink for {field_name}: {url}, caption: {caption}")
                elif value_str.startswith('http://') or value_str.startswith('https://'):
                    url = value_str
                    caption = field_name
                    logger.debug(f"Using value URL for {field_name}: {url}")
                
                if url and caption:
                    images.append({
                        'name': caption,
                        'url': url
                    })
                    logger.info(f"Added image: {caption} -> {url[:50]}...")
        
        logger.info(f"Total images extracted: {len(images)}")
        return images
    
    def _find_column(self, keywords: List[str]) -> Optional[str]:
        """Find column by keywords"""
        for col in self.columns:
            if any(kw in col.lower() for kw in keywords):
                return col
        return None
    
    def _refresh_cards(self):
        """Refresh building cards display"""
        # Clear old cards
        for widget in self.cards_scroll.winfo_children():
            widget.destroy()
        
        if self.building_records:
            for idx, record in enumerate(self.building_records):
                self._create_building_card(record, idx)
        else:
            empty = ctk.CTkLabel(self.cards_scroll, text="No buildings match filter",
                               font=("Arial", 13), text_color=DARK_BORDER)
            empty.pack(pady=40)
    
    def _create_building_card(self, record: BuildingRecord, idx: int):
        """Create a single building card"""
        if not record:
            return
        
        # Card frame
        card = ctk.CTkFrame(self.cards_scroll, fg_color=DARK_CARD, corner_radius=15)
        card.pack(fill="x", padx=10, pady=10)
        
        # Header
        header = ctk.CTkFrame(card, fg_color=DARK_ACCENT, corner_radius=12)
        header.pack(fill="x", padx=15, pady=(15, 10))
        
        ctk.CTkLabel(header, text=f"🏢 {record.building_id}",
                    font=("Arial", 14, "bold"), text_color=PRIMARY).pack(anchor="w", padx=10, pady=8)
        
        # Data grid
        data_frame = ctk.CTkFrame(card, fg_color="transparent")
        data_frame.pack(fill="x", padx=15, pady=10)
        
        # Display key data (limit to 10 fields)
        displayed_fields = 0
        for field_name, value in record.data.items():
            if displayed_fields >= 10 or pd.isna(value):
                continue
            
            value_str = str(value)
            # Skip URLs and image files in main display
            if 'http' not in value_str.lower() and '.jpg' not in value_str.lower() and '.png' not in value_str.lower():
                row_frame = ctk.CTkFrame(data_frame, fg_color="transparent")
                row_frame.pack(fill="x", pady=2)
                
                label = ctk.CTkLabel(row_frame, text=f"{field_name}:",
                                    font=("Arial", 10, "bold"), text_color=DARK_BORDER, width=150)
                label.pack(side="left", padx=(0, 10))
                
                val = ctk.CTkLabel(row_frame, text=value_str[:70],
                                  font=("Arial", 10), text_color=TEXT_DARK)
                val.pack(side="left", fill="x", expand=True)
                
                displayed_fields += 1
        
        # Image buttons
        if record.images:
            image_divider = ctk.CTkFrame(card, height=1, fg_color=DARK_BORDER)
            image_divider.pack(fill="x", padx=15, pady=10)
            
            ctk.CTkLabel(card, text="📸 Images/Links:",
                        font=("Arial", 10, "bold"), text_color=SECONDARY).pack(anchor="w", padx=15, pady=(0, 5))
            
            image_frame = ctk.CTkFrame(card, fg_color="transparent")
            image_frame.pack(fill="x", padx=15, pady=(0, 15))
            
            for img_info in record.images[:5]:  # Show max 5 image buttons
                # Truncate button name if too long
                btn_name = img_info['name'][:35]  # Limit to 35 chars
                if len(img_info['name']) > 35:
                    btn_name += "..."
                
                img_btn = ctk.CTkButton(image_frame, text=f"🖼️ {btn_name}",
                                       fg_color=SECONDARY, hover_color="#1870FF",
                                       font=("Arial", 9),
                                       command=lambda url=img_info['url']: self._open_image(url))
                img_btn.pack(side="left", padx=3, pady=2, fill="x", expand=True)
    
    def _open_image(self, url: str):
        """Open image in browser or viewer"""
        try:
            import webbrowser
            webbrowser.open(url)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to open image:\n{str(e)}")
    
    def _apply_filter(self):
        """Apply 3-step filter using field, operator, and value"""
        if not self.building_records:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        field_name = self.filter_column.get()
        operator = self.filter_operator.get()
        filter_val = self.filter_value.get()
        
        if field_name == "Field" or not filter_val or not operator:
            messagebox.showwarning("Invalid Filter", "Select field, operator, and enter value")
            return
        
        try:
            # Store original records
            if not hasattr(self, 'original_records'):
                self.original_records = self.building_records.copy()
            
            # Filter building records based on field, operator, and value
            filtered_records = []
            for record in self.original_records:
                if field_name in record.data:
                    field_value = str(record.data[field_name]).lower()
                    search_val = filter_val.lower()
                    
                    match = False
                    if operator == "Contains":
                        match = search_val in field_value
                    elif operator == "Equals":
                        match = field_value == search_val
                    elif operator == "Starts With":
                        match = field_value.startswith(search_val)
                    elif operator == "Ends With":
                        match = field_value.endswith(search_val)
                    elif operator == "Not Contains":
                        match = search_val not in field_value
                    
                    if match:
                        filtered_records.append(record)
            
            self.building_records = filtered_records
            self._refresh_cards()
            self.record_count.configure(text=f"Buildings: {len(filtered_records)}")
            messagebox.showinfo("Filter Applied", f"Found {len(filtered_records)} matching buildings")
        except Exception as e:
            messagebox.showerror("Filter Error", str(e))
    
    def _reset_filter(self):
        """Reset filter"""
        if self.df is None:
            return
        # Reload all building records
        self._parse_records()
        self.filter_value.delete(0, "end")
        self.filter_column.set("Field")
        self.filter_operator.set("Contains")
        self._refresh_cards()
        self.record_count.configure(text=f"Buildings: {len(self.building_records)}")
        messagebox.showinfo("Filter Reset", f"Showing all {len(self.building_records)} buildings")
    
    def _generate_map(self):
        """Generate and display embedded OSM with interactive pan/zoom and region selection"""
        if not self.building_records:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        try:
            # Extract coordinates from GPS COORDINATES field
            lats = []
            lons = []
            building_ids = []
            valid_records = []
            
            for record in self.building_records:
                # Look for GPS COORDINATES field
                gps_coords = None
                for field_name, value in record.data.items():
                    if 'gps' in str(field_name).lower() and 'coordinates' in str(field_name).lower():
                        gps_coords = value
                        break
                
                if gps_coords:
                    # Parse lat,lon from GPS coordinates string
                    try:
                        gps_str = str(gps_coords).strip()
                        parts = gps_str.split(',')
                        if len(parts) >= 2:
                            lat = float(parts[0].strip())
                            lon = float(parts[1].strip())
                            
                            if -90 <= lat <= 90 and -180 <= lon <= 180:
                                lats.append(lat)
                                lons.append(lon)
                                building_ids.append(record.building_id)
                                valid_records.append(record)
                    except:
                        pass
            
            if not lats or not lons:
                messagebox.showerror("No Valid Coordinates", "Could not find valid GPS coordinates\nExpected format: 'latitude, longitude'")
                return
            
            center_lat = np.mean(lats)
            center_lon = np.mean(lons)
            
            # Clear map frame
            for widget in self.map_frame.winfo_children():
                widget.destroy()
            
            # Show loading message
            loading_label = ctk.CTkLabel(self.map_frame, 
                                        text="📡 Loading OpenStreetMap tiles...",
                                        font=("Arial", 12), text_color=PRIMARY)
            loading_label.pack(pady=20)
            self.map_frame.update()
            
            # Create matplotlib figure for map with proper geographic aspect ratio
            fig, ax = plt.subplots(figsize=(14, 8), dpi=100)
            fig.patch.set_facecolor(DARK_CARD)
            
            # Set the map extent
            margin = 0.02
            min_lon, max_lon = min(lons) - margin, max(lons) + margin
            min_lat, max_lat = min(lats) - margin, max(lats) + margin
            
            ax.set_xlim(min_lon, max_lon)
            ax.set_ylim(min_lat, max_lat)
            
            # Fix aspect ratio to show map correctly (equal scaling for lat/lon)
            aspect = np.cos(np.radians(center_lat))
            ax.set_aspect(1/aspect, adjustable='box')
            
            # Create scatter plot with coordinates (in lat/lon space)
            scatter = ax.scatter(lons, lats, c=PRIMARY, s=300, marker='o', 
                               edgecolors='white', linewidth=2.5, alpha=0.8, zorder=5, label='Buildings')
            
            # Add building labels with background
            for lon, lat, bid in zip(lons, lats, building_ids):
                ax.annotate(bid, (lon, lat), xytext=(8, 8), textcoords='offset points',
                           fontsize=9, color='white', fontweight='bold', 
                           bbox=dict(boxstyle='round,pad=0.3', facecolor='#00D4FF', alpha=0.7, edgecolor='white', linewidth=1),
                           zorder=6)
            
            # Add center marker
            ax.plot(center_lon, center_lat, marker='*', markersize=25, color='red', 
                   markeredgecolor='white', markeredgewidth=2, zorder=4, label='Center')
            
            # Add OSM basemap tiles using contextily
            try:
                ctx.add_basemap(ax, crs='EPSG:4326', source=ctx.providers.OpenStreetMap.Mapnik, 
                               zoom=13, alpha=0.9)
            except Exception as e:
                logger.warning(f"Could not load OSM tiles: {e}. Using plain background.")
                ax.set_facecolor('#D0E8F2')
            
            # Set labels and title
            ax.set_xlabel('Longitude', fontsize=12, color=TEXT_DARK, fontweight='bold')
            ax.set_ylabel('Latitude', fontsize=12, color=TEXT_DARK, fontweight='bold')
            ax.set_title(f'Seismic Survey Locations - OpenStreetMap ({len(valid_records)} buildings)\n[Drag rectangle to select region, Use toolbar to pan/zoom]', 
                        fontsize=13, color=PRIMARY, fontweight='bold', pad=15)
            
            # Style axes
            ax.tick_params(colors=TEXT_DARK, labelsize=10)
            for spine in ax.spines.values():
                spine.set_edgecolor(DARK_BORDER)
                spine.set_linewidth(1.5)
            
            # Add grid
            ax.grid(True, alpha=0.3, color='gray', linestyle='--', linewidth=0.7, zorder=0)
            
            # Add legend
            ax.legend(loc='upper right', fontsize=11, framealpha=0.95, edgecolor=DARK_BORDER)
            
            # Tight layout
            fig.tight_layout()
            
            # Remove loading label
            loading_label.destroy()
            
            # Embed in tkinter with navigation toolbar
            canvas = FigureCanvasTkAgg(fig, master=self.map_frame)
            canvas.draw()
            
            # Add matplotlib toolbar for pan/zoom
            toolbar_frame = ctk.CTkFrame(self.map_frame, fg_color=DARK_ACCENT, corner_radius=0)
            toolbar_frame.pack(fill="x", padx=0, pady=0)
            
            toolbar = NavigationToolbar2Tk(canvas, toolbar_frame)
            toolbar.update()
            toolbar_frame.pack(fill="x", side="top")
            
            canvas.get_tk_widget().pack(fill="both", expand=True, padx=15, pady=(10, 5))
            
            # Store map data for region selection
            self.map_lons = lons
            self.map_lats = lats
            self.map_building_ids = building_ids
            self.map_valid_records = valid_records
            self.map_scatter = scatter
            self.map_ax = ax
            self.map_fig = fig
            self.map_canvas = canvas
            self.map_center_lat = center_lat
            self.map_center_lon = center_lon
            
            # Add rectangle selector for region filtering
            def on_select(eclick, erelease):
                x1, y1 = eclick.xdata, eclick.ydata
                x2, y2 = erelease.xdata, erelease.ydata
                
                min_x, max_x = min(x1, x2), max(x1, x2)
                min_y, max_y = min(y1, y2), max(y1, y2)
                
                # Filter buildings within selected region
                filtered_lon = []
                filtered_lat = []
                filtered_bid = []
                filtered_records = []
                
                for lon, lat, bid, record in zip(lons, lats, building_ids, valid_records):
                    if min_x <= lon <= max_x and min_y <= lat <= max_y:
                        filtered_lon.append(lon)
                        filtered_lat.append(lat)
                        filtered_bid.append(bid)
                        filtered_records.append(record)
                
                if filtered_lon:
                    self._redraw_map_with_selection(filtered_lon, filtered_lat, filtered_bid, 
                                                    len(filtered_records), f"({len(filtered_records)}/{len(valid_records)} buildings in region)")
                    # Update building records to show only selected
                    original_records = self.building_records.copy()
                    self.building_records = filtered_records
                    self._refresh_cards()
                    # Restore for next map generation
                    self.building_records = original_records
            
            # Create rectangle selector
            rect_selector = RectangleSelector(
                ax, on_select,
                useblit=True,
                button=[1],  # Left mouse button
                minspanx=5, minspany=5,
                spancoords='pixels',
                interactive=True
            )
            rect_selector.set_active(True)
            
            # Add info frame below map
            info_frame = ctk.CTkFrame(self.map_frame, fg_color=DARK_ACCENT, corner_radius=10)
            info_frame.pack(fill="x", padx=15, pady=(10, 15))
            
            info_text = f"✓ OpenStreetMap Embedded | Center: {center_lat:.4f}, {center_lon:.4f} | Buildings: {len(valid_records)}\n💡 Tip: Draw rectangle on map to filter buildings by region"
            info_label = ctk.CTkLabel(info_frame, text=info_text,
                                     font=("Arial", 10), text_color=SUCCESS, justify="left")
            info_label.pack(anchor="w", padx=15, pady=10)
            
        except Exception as e:
            logger.error(f"Map error: {e}", exc_info=True)
            messagebox.showerror("Map Error", str(e))
    
    def _redraw_map_with_selection(self, lons, lats, building_ids, count, status):
        """Redraw map highlighting selected region buildings"""
        try:
            # Clear scatter and annotations
            for artist in self.map_ax.collections:
                artist.remove()
            for txt in self.map_ax.texts:
                txt.remove()
            
            # Plot selected buildings (green)
            scatter_selected = self.map_ax.scatter(lons, lats, c='#00FF00', s=350, marker='o', 
                                                  edgecolors='white', linewidth=2.5, alpha=0.9, 
                                                  zorder=5, label='Selected Buildings')
            
            # Plot unselected buildings (dimmed)
            for i, (lon, lat) in enumerate(zip(self.map_lons, self.map_lats)):
                if lon not in lons:
                    self.map_ax.scatter(lon, lat, c=PRIMARY, s=200, marker='o', 
                                      edgecolors='white', linewidth=1.5, alpha=0.3, zorder=4)
            
            # Add labels for selected buildings
            for lon, lat, bid in zip(lons, lats, building_ids):
                self.map_ax.annotate(bid, (lon, lat), xytext=(8, 8), textcoords='offset points',
                                    fontsize=9, color='white', fontweight='bold', 
                                    bbox=dict(boxstyle='round,pad=0.3', facecolor='#00FF00', 
                                            alpha=0.8, edgecolor='white', linewidth=1),
                                    zorder=6)
            
            # Update title
            self.map_ax.set_title(f'Region Selection - {status}', 
                                 fontsize=13, color='#00FF00', fontweight='bold')
            
            # Update legend
            self.map_ax.legend(loc='upper right', fontsize=11, framealpha=0.95, edgecolor=DARK_BORDER)
            
            self.map_fig.tight_layout()
            self.map_canvas.draw()
            
        except Exception as e:
            logger.error(f"Redraw error: {e}")
    
    def _open_map_browser(self):
        """Open the interactive map in browser"""
        if hasattr(self, 'map_html_file') and self.map_html_file:
            try:
                os.startfile(self.map_html_file)
            except Exception as e:
                messagebox.showerror("Error", f"Failed to open map:\n{str(e)}")
    
    def _show_building_cards(self):
        """Switch to building cards tab"""
        self.tabs.set("🏢 Building Cards")
    
    def _show_map_tab(self):
        """Switch to map tab"""
        self.tabs.set("🗺️ Map View")
    
    def _show_analytics(self):
        """Switch to analytics tab"""
        self.tabs.set("📊 Analytics")
        self._update_stats()
    
    def _update_stats(self):
        """Update analytics statistics"""
        if not self.building_records:
            return
        
        try:
            # Total buildings
            self.stat_labels['total'].configure(text=str(len(self.building_records)))
            
            # Average age
            age_col = self._find_column(['age', 'year', 'built'])
            if age_col:
                ages = pd.to_numeric(self.df[age_col], errors='coerce').dropna()
                if len(ages) > 0:
                    avg_age = ages.mean()
                    self.stat_labels['avg_age'].configure(text=f"{avg_age:.0f} yrs")
            
            # Material count
            mat_col = self._find_column(['material', 'type', 'construction'])
            if mat_col:
                unique_materials = self.df[mat_col].nunique()
                self.stat_labels['materials'].configure(text=str(unique_materials))
            
            # Locations
            lat_col = self._find_column(['latitude', 'lat'])
            if lat_col:
                valid_coords = self.df[lat_col].notna().sum()
                self.stat_labels['locations'].configure(text=str(valid_coords))
        
        except Exception as e:
            logger.warning(f"Stats update error: {e}")
    
    def _bulk_download_images(self):
        """Bulk download all images from records"""
        if not self.building_records:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        # Count total images
        total_images = sum(len(r.images) for r in self.building_records)
        if total_images == 0:
            messagebox.showinfo("No Images", "No image URLs found in data")
            return
        
        # Ask for output directory
        output_dir = filedialog.askdirectory(title="Select Output Directory for Images")
        if not output_dir:
            return
        
        # Show progress window
        progress_window = ctk.CTkToplevel(self)
        progress_window.title("Downloading Images")
        progress_window.geometry("400x150")
        progress_window.resizable(False, False)
        
        status_label = ctk.CTkLabel(progress_window, text="Downloading...",
                                   font=("Arial", 12))
        status_label.pack(pady=(20, 10))
        
        progress_bar = ctk.CTkProgressBar(progress_window, width=300)
        progress_bar.set(0)
        progress_bar.pack(pady=10)
        
        count_label = ctk.CTkLabel(progress_window, text="0 / 0",
                                  font=("Arial", 11))
        count_label.pack(pady=10)
        
        def update_progress(current, total):
            progress_bar.set(current / max(total, 1))
            count_label.configure(text=f"{current} / {total}")
            progress_window.update()
        
        # Download in thread
        def download_thread():
            downloaded = ImageDownloader.batch_download(
                self.building_records, output_dir, update_progress
            )
            progress_window.destroy()
            messagebox.showinfo("Complete", f"Downloaded {downloaded} images!")
        
        threading.Thread(target=download_thread, daemon=True).start()
    
    def _export_passports(self):
        """Export building passports to Word document"""
        if not self.building_records:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")]
        )
        
        if not save_path:
            return
        
        try:
            doc = Document()
            doc.add_heading('Building Survey Passports', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Total Buildings: {len(self.filtered_df)}')
            doc.add_page_break()
            
            id_col = self._find_column(['id', 'building_id']) or self.columns[0]
            
            for idx, row in self.filtered_df.iterrows():
                record = self.building_records[idx] if idx < len(self.building_records) else None
                
                # Building heading
                building_id = row.get(id_col, f"Building_{idx}")
                doc.add_heading(f'Building: {building_id}', level=1)
                
                # Data table
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Light Grid Accent 1'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Field'
                hdr_cells[1].text = 'Value'
                
                for col, value in row.items():
                    if pd.notna(value) and 'http' not in str(value).lower():
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(col)
                        row_cells[1].text = str(value)[:100]
                
                # Images
                if record and record.images:
                    doc.add_heading('Images', level=2)
                    for img_info in record.images:
                        p = doc.add_paragraph()
                        p.add_run(f"Link: ").bold = True
                        p.add_run(img_info['url']).italic = True
                
                doc.add_page_break()
            
            doc.save(save_path)
            messagebox.showinfo("Success", f"Exported {len(self.filtered_df)} passports!")
        
        except Exception as e:
            logger.error(f"Export error: {e}", exc_info=True)
            messagebox.showerror("Export Error", str(e))
    
    def _generate_chart(self):
        """Generate chart based on selected type and fields"""
        if not self.building_records:
            messagebox.showwarning("No Data", "Load data first")
            return
        
        chart_type = self.chart_type.get()
        x_field = self.chart_x_field.get()
        y_field = self.chart_y_field.get()
        
        if x_field == "Field" or y_field == "Field" or not chart_type:
            messagebox.showwarning("Invalid Selection", "Select chart type and both X and Y axis fields")
            return
        
        try:
            # Prepare data
            x_data = []
            y_data = []
            
            for record in self.building_records:
                if x_field in record.data and y_field in record.data:
                    x_val = record.data[x_field]
                    y_val = record.data[y_field]
                    
                    try:
                        # Try to convert to numeric
                        x_num = float(x_val) if str(x_val).replace('.', '', 1).replace('-', '', 1).isdigit() else x_val
                        y_num = float(y_val) if str(y_val).replace('.', '', 1).replace('-', '', 1).isdigit() else y_val
                        
                        x_data.append(x_num)
                        y_data.append(y_num)
                    except:
                        continue
            
            if not x_data or not y_data:
                messagebox.showerror("No Data", "No valid data found for selected fields")
                return
            
            # Clear previous chart
            for widget in self.chart_frame.winfo_children():
                widget.destroy()
            
            # Create figure
            fig, ax = plt.subplots(figsize=(12, 6), dpi=100)
            fig.patch.set_facecolor(DARK_CARD)
            
            # Plot based on chart type
            try:
                if chart_type == "Bar Chart":
                    ax.bar(range(len(x_data)), y_data, color=PRIMARY, edgecolor='white', linewidth=1.5)
                    ax.set_xlabel(x_field, fontsize=11, fontweight='bold', color=TEXT_DARK)
                    ax.set_ylabel(y_field, fontsize=11, fontweight='bold', color=TEXT_DARK)
                    ax.set_title(f"{y_field} by {x_field} (Bar Chart)", fontsize=13, fontweight='bold', color=PRIMARY)
                    
                elif chart_type == "Pie Chart":
                    # For pie chart, use counts
                    from collections import Counter
                    if isinstance(x_data[0], str):
                        counts = Counter(x_data)
                        ax.pie(counts.values(), labels=counts.keys(), autopct='%1.1f%%', 
                              colors=[PRIMARY, SECONDARY, SUCCESS, WARNING] * 10, startangle=90)
                    else:
                        ax.pie(y_data[:10], labels=[f"Item {i}" for i in range(len(y_data[:10]))], 
                              autopct='%1.1f%%', colors=[PRIMARY, SECONDARY, SUCCESS, WARNING] * 10)
                    ax.set_title(f"{x_field} Distribution", fontsize=13, fontweight='bold', color=PRIMARY)
                    
                elif chart_type == "Line Chart":
                    ax.plot(range(len(x_data)), y_data, marker='o', color=PRIMARY, linewidth=2.5, 
                           markersize=6, markeredgecolor='white', markeredgewidth=1.5)
                    ax.set_xlabel(x_field, fontsize=11, fontweight='bold', color=TEXT_DARK)
                    ax.set_ylabel(y_field, fontsize=11, fontweight='bold', color=TEXT_DARK)
                    ax.set_title(f"{y_field} by {x_field} (Line Chart)", fontsize=13, fontweight='bold', color=PRIMARY)
                    ax.grid(True, alpha=0.3, color='gray', linestyle='--')
                    
                elif chart_type == "Scatter Plot":
                    ax.scatter(x_data, y_data, s=100, c=PRIMARY, alpha=0.7, edgecolors='white', linewidth=1.5)
                    ax.set_xlabel(x_field, fontsize=11, fontweight='bold', color=TEXT_DARK)
                    ax.set_ylabel(y_field, fontsize=11, fontweight='bold', color=TEXT_DARK)
                    ax.set_title(f"{y_field} vs {x_field} (Scatter Plot)", fontsize=13, fontweight='bold', color=PRIMARY)
                    ax.grid(True, alpha=0.3, color='gray', linestyle='--')
                    
                elif chart_type == "Histogram":
                    ax.hist(y_data, bins=15, color=PRIMARY, edgecolor='white', linewidth=1.5, alpha=0.8)
                    ax.set_xlabel(y_field, fontsize=11, fontweight='bold', color=TEXT_DARK)
                    ax.set_ylabel("Frequency", fontsize=11, fontweight='bold', color=TEXT_DARK)
                    ax.set_title(f"Distribution of {y_field} (Histogram)", fontsize=13, fontweight='bold', color=PRIMARY)
                    ax.grid(True, alpha=0.3, axis='y', color='gray', linestyle='--')
                
                # Style axes
                ax.tick_params(colors=TEXT_DARK, labelsize=10)
                for spine in ax.spines.values():
                    spine.set_edgecolor(DARK_BORDER)
                    spine.set_linewidth(1.5)
                ax.set_facecolor(DARK_ACCENT)
                
                plt.tight_layout()
                
                # Embed in tkinter
                canvas = FigureCanvasTkAgg(fig, master=self.chart_frame)
                canvas.draw()
                canvas.get_tk_widget().pack(fill="both", expand=True, padx=10, pady=10)
                
                # Store reference
                self.chart_canvas = canvas
                self.chart_figure = fig
                
            except Exception as e:
                logger.error(f"Charting error: {e}")
                messagebox.showerror("Chart Error", f"Failed to create chart:\n{str(e)}")
        
        except Exception as e:
            logger.error(f"Chart generation error: {e}", exc_info=True)
            messagebox.showerror("Error", str(e))
    
    def _clear_chart(self):
        """Clear the chart display"""
        for widget in self.chart_frame.winfo_children():
            widget.destroy()
        
        placeholder = ctk.CTkLabel(self.chart_frame,
                                  text="📊 Select fields and click 'Generate Chart' to display visualization",
                                  font=("Arial", 13), text_color=DARK_BORDER)
        placeholder.place(relx=0.5, rely=0.5, anchor="center")
        
        # Close matplotlib figure if exists
        if hasattr(self, 'chart_figure'):
            plt.close(self.chart_figure)

# --- Main ---
if __name__ == "__main__":
    app = SeismicSurveyViewerApp()
    app.mainloop()
