"""
SEISMIC BUILDING VULNERABILITY ANALYZER - Advanced Edition
Enhanced with AI-driven risk assessment, clustering, and intelligent recommendations
Aligned with BuildingForm.tsx dynamic schema structure
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox, scrolledtext
import os
import sys
import traceback
import ctypes
import threading
import requests
import shutil
import pandas as pd
import numpy as np
import json
from typing import Dict, List, Tuple, Optional, Any
from dataclasses import dataclass
from abc import ABC, abstractmethod
from datetime import datetime
import logging
import folium
import webbrowser
import tempfile
import matplotlib.pyplot as plt
import matplotlib.patches as mpatches
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
from matplotlib.figure import Figure

# --- Logging Setup ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# --- Configuration ---
ctk.set_appearance_mode("Light")
ctk.set_default_color_theme("dark-blue")

# --- Color Palette ---
NAVY = "#001F3F"
AQUA = "#39CCCC"
MAROON = "#85144B"
GREEN = "#2ECC40"
YELLOW = "#FFDC00"
WHITE = "#FFFFFF"
GREY = "#F5F5F5"
LIGHT_RED = "#FF4136"
ORANGE = "#FF851B"

# Default fields that seismic analysis should look for
SEISMIC_FIELD_KEYWORDS = {
    'building_age': ['year', 'age', 'construction', 'built'],
    'material_type': ['material', 'concrete', 'stone', 'brick', 'wood', 'steel'],
    'structure_type': ['structure', 'type', 'frame', 'bearing wall'],
    'stories': ['stories', 'floors', 'levels', 'height'],
    'foundation_type': ['foundation', 'footing', 'base'],
    'damage_level': ['damage', 'crack', 'deterioration', 'condition'],
    'modifications': ['modification', 'alteration', 'repair', 'upgrade'],
    'location': ['latitude', 'longitude', 'lat', 'lon', 'location'],
    'building_id': ['id', 'building_id', 'bid', 'code'],
}

# =============================================
# 1. DATA MODELS & VALIDATORS
# =============================================

@dataclass
class BuildingRecord:
    """Typed building record with validation"""
    building_id: str
    survey_date: Optional[str] = None
    surveyor_name: Optional[str] = None
    district: Optional[str] = None
    material_type: Optional[str] = None
    building_age: Optional[float] = None
    stories: Optional[int] = None
    foundation_type: Optional[str] = None
    damage_level: Optional[str] = None
    latitude: Optional[float] = None
    longitude: Optional[float] = None
    structural_system: Optional[str] = None
    aspect_ratio: Optional[float] = None
    raw_data: Dict[str, Any] = None
    comments: str = ""  # Combined from all comment fields
    
    def __post_init__(self):
        if self.raw_data is None:
            self.raw_data = {}

class DataValidator:
    """Validates and cleans building data"""
    
    @staticmethod
    def validate_location(lat: float, lon: float) -> bool:
        """Validate GPS coordinates"""
        return -90 <= lat <= 90 and -180 <= lon <= 180
    
    @staticmethod
    def validate_building_age(age: float) -> bool:
        """Age should be 0-500 years"""
        return 0 <= age <= 500
    
    @staticmethod
    def validate_stories(stories: int) -> bool:
        """Stories should be 1-100"""
        return 1 <= stories <= 100
    
    @staticmethod
    def safe_float(value: Any) -> Optional[float]:
        """Safely convert to float"""
        try:
            if pd.isna(value) or value == '': return None
            return float(value)
        except:
            return None
    
    @staticmethod
    def safe_int(value: Any) -> Optional[int]:
        """Safely convert to int"""
        try:
            if pd.isna(value) or value == '': return None
            return int(float(value))
        except:
            return None

# =============================================
# 2. RISK CALCULATION ENGINE
# =============================================

class VulnerabilityCalculator(ABC):
    """Abstract base for all vulnerability calculators"""
    
    @abstractmethod
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        """
        Returns: (score 0-100, details dict)
        """
        pass

class AgeRiskCalculator(VulnerabilityCalculator):
    """Building age is critical seismic indicator"""
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        if not record.building_age:
            return 0, {"status": "missing"}
        
        age = record.building_age
        # Pre-1980 unengineered, 1980-2000 partial, 2000+ modern codes
        if age > 40:
            score = 35
            interpretation = "High risk - pre-earthquake code era"
        elif age > 25:
            score = 20
            interpretation = "Moderate risk - transition period"
        else:
            score = 10
            interpretation = "Lower risk - modern construction codes"
        
        return score, {
            "score": score,
            "age_years": age,
            "interpretation": interpretation,
            "weight": 0.35  # 35% of total vulnerability
        }

class MaterialRiskCalculator(VulnerabilityCalculator):
    """Different materials have different seismic behavior"""
    
    MATERIAL_RISKS = {
        'unreinforced': 40, 'stone': 35, 'brick': 30,
        'adobe': 40, 'mud': 40,
        'wood': 15, 'rc': 10, 'concrete': 10, 'steel': 5,
    }
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        if not record.material_type:
            return 0, {"status": "missing"}
        
        material = record.material_type.lower()
        
        # Fuzzy matching
        score = 0
        for key, val in self.MATERIAL_RISKS.items():
            if key in material:
                score = val
                break
        
        if score == 0:
            score = 20  # Unknown material = moderate risk
        
        return score, {
            "score": score,
            "material": record.material_type,
            "interpretation": f"{record.material_type} - inherent vulnerability",
            "weight": 0.25
        }

class FoundationRiskCalculator(VulnerabilityCalculator):
    """Foundation quality critical in earthquakes"""
    
    FOUNDATION_RISKS = {
        'unknown': 30, 'surface': 35, 'poor': 40,
        'adequate': 15, 'deep': 10, 'piled': 5,
    }
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        if not record.foundation_type:
            return 0, {"status": "missing"}
        
        foundation = record.foundation_type.lower()
        
        score = 0
        for key, val in self.FOUNDATION_RISKS.items():
            if key in foundation:
                score = val
                break
        
        if score == 0:
            score = 20  # Unknown = moderate
        
        return score, {
            "score": score,
            "foundation": record.foundation_type,
            "interpretation": f"Foundation type: {record.foundation_type}",
            "weight": 0.20
        }

class HeightRiskCalculator(VulnerabilityCalculator):
    """Tall buildings amplify seismic waves"""
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        if not record.stories:
            return 0, {"status": "missing"}
        
        stories = record.stories
        if stories > 7:
            score = 25
            interp = "High-rise amplifies ground motion"
        elif stories > 4:
            score = 15
            interp = "Mid-rise moderate amplification"
        else:
            score = 5
            interp = "Low-rise naturally stable"
        
        return score, {
            "score": score,
            "stories": stories,
            "interpretation": interp,
            "weight": 0.10
        }

class CommentAnalyzer:
    """Simple lightweight comment analyzer to extract intensity from free text"""
    KEYWORD_SCORES = {
        'collapsed': 30,
        'collapse': 30,
        'severe': 25,
        'critical': 25,
        'major': 20,
        'heavy': 20,
        'severe cracking': 25,
        'crack': 10,
        'damage': 12,
        'deterioration': 8,
        'tilt': 20,
        'leaning': 20,
    }

    def analyze_comments(self, text: str) -> Tuple[int, Dict[str, int]]:
        """Return a coarse score (0-30) and matched keywords counts"""
        if not text:
            return 0, {}
        t = text.lower()
        score = 0
        matches = {}
        for kw, val in self.KEYWORD_SCORES.items():
            if kw in t:
                score = max(score, val)
                matches[kw] = matches.get(kw, 0) + 1
        return score, matches

class ConditionRiskCalculator(VulnerabilityCalculator):
    """Current damage and deterioration"""
    
    CONDITION_KEYWORDS = {
        'severe': 25, 'heavy': 25, 'critical': 30,
        'moderate': 15, 'fair': 10,
        'good': 5, 'excellent': 0,
        'crack': 10, 'damage': 12, 'deterioration': 8,
    }
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        # If no explicit damage text, still attempt to infer from comments
        score = 0
        if record.damage_level and pd.notna(record.damage_level):
            damage = str(record.damage_level).lower()
            for keyword, val in self.CONDITION_KEYWORDS.items():
                if keyword in damage:
                    score = max(score, val)

        # Use comment analyzer for qualitative clues
        try:
            analyzer = CommentAnalyzer()
            comment_score, comment_details = analyzer.analyze_comments(record.comments or "")
        except Exception:
            comment_score, comment_details = 0, {}

        # Combine damage-derived score and comment-derived score (take max)
        score = max(score, comment_score)

        return score, {
            "score": score,
            "damage_level": record.damage_level,
            "comment_analysis": bool(record.comments),
            "comment_details": comment_details,
            "interpretation": f"Current condition: {record.damage_level}",
            "weight": 0.10
        }

class StructuralSystemCalculator(VulnerabilityCalculator):
    """Analyze primary lateral resisting system"""
    
    SYSTEM_RISKS = {
        'frame': 10, 'rc frame': 10, 'steel frame': 5,
        'wall': 30, 'bearing wall': 30, 'shear wall': 15,
        'mixed': 20, 'combined': 20,
        'unknown': 25,
    }
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        system_text = record.raw_data.get('PRIMARY LATERAL RESISTING SYSTEM', '')
        if not system_text or pd.isna(system_text):
            return 0, {"status": "missing"}
        
        system = str(system_text).lower()
        score = 0
        for key, val in self.SYSTEM_RISKS.items():
            if key in system:
                score = val
                break
        if score == 0:
            score = 20
        
        return score, {
            "score": score,
            "system": system_text,
            "interpretation": f"Structural system: {system_text}",
            "weight": 0.15
        }

class AspectRatioCalculator(VulnerabilityCalculator):
    """Calculate building aspect ratio vulnerability"""
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        length = record.raw_data.get('length') or record.raw_data.get('BUILDING DIMENSION [LONGITUDINAL DIMENSION (FT)]')
        width = record.raw_data.get('width') or record.raw_data.get('BUILDING DIMENSION [TRANSVERSE DIMENSION (FT)]')
        
        try:
            length = float(length) if length else None
            width = float(width) if width else None
        except:
            return 0, {"status": "missing"}
        
        if not length or not width or width == 0:
            return 0, {"status": "missing"}
        
        aspect_ratio = length / width
        
        # High aspect ratio = weak in transverse direction
        if aspect_ratio > 3:
            score = 25
            interp = "High aspect ratio - weak transverse direction"
        elif aspect_ratio > 2:
            score = 15
            interp = "Moderate aspect ratio"
        else:
            score = 5
            interp = "Favorable aspect ratio"
        
        return score, {
            "score": score,
            "aspect_ratio": round(aspect_ratio, 2),
            "length": length,
            "width": width,
            "interpretation": interp,
            "weight": 0.08
        }

class ConnectionsQualityCalculator(VulnerabilityCalculator):
    """Assess critical connections (roof, floor, walls)"""
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        roof_conn = str(record.raw_data.get('ROOF-TO-WALL CONNECTION', '')).lower()
        floor_conn = str(record.raw_data.get('FLOOR-TO-WALL', '')).lower()
        wall_conn = str(record.raw_data.get('WALL-TO-WALL CONNECTION', '')).lower()
        
        conn_text = f"{roof_conn} {floor_conn} {wall_conn}"
        
        score = 15  # Default moderate risk
        if any(w in conn_text for w in ['good', 'strong', 'tied', 'connected']):
            score = 5
        elif any(w in conn_text for w in ['poor', 'weak', 'loose', 'missing']):
            score = 30
        
        return score, {
            "score": score,
            "roof_connection": roof_conn,
            "floor_connection": floor_conn,
            "wall_connection": wall_conn,
            "interpretation": "Connection quality assessment",
            "weight": 0.12
        }

class DamageDetailCalculator(VulnerabilityCalculator):
    """Parse specific damage types from structured data"""
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        damage_types = record.raw_data.get('OBSERVED DAMAGE TYPE(S)', '')
        
        if not damage_types or pd.isna(damage_types):
            return 0, {"status": "missing"}
        
        damage_str = str(damage_types).lower()
        
        # Weighted damage scoring
        score = 0
        damages_found = []
        
        damage_weights = {
            'collapse': 40, 'severe cracking': 30, 'wall failure': 30,
            'crack': 15, 'separation': 20, 'tilt': 25,
            'settlement': 18, 'none': 0
        }
        
        for damage, weight in damage_weights.items():
            if damage in damage_str:
                score = max(score, weight)
                damages_found.append(damage)
        
        return score, {
            "score": score,
            "damages": damages_found,
            "interpretation": f"Damage types: {', '.join(damages_found)}",
            "weight": 0.10
        }

class SeismicVulnerabilityIndex:
    """Composite index combining multiple risk factors"""
    
    def __init__(self):
        self.calculators = [
            AgeRiskCalculator(),
            MaterialRiskCalculator(),
            StructuralSystemCalculator(),
            AspectRatioCalculator(),
            ConnectionsQualityCalculator(),
            HeightRiskCalculator(),
            ConditionRiskCalculator(),
            DamageDetailCalculator(),
        ]
    
    def calculate(self, record: BuildingRecord) -> Tuple[float, Dict[str, Any]]:
        """
        Returns: (vulnerability_score 0-100, detailed_breakdown)
        """
        scores = {}
        details = {}
        total_weight = 0
        weighted_score = 0
        
        for calc in self.calculators:
            score, detail = calc.calculate(record)
            calc_name = calc.__class__.__name__
            scores[calc_name] = score
            details[calc_name] = detail
            
            if 'weight' in detail:
                weight = detail['weight']
                weighted_score += score * weight
                total_weight += weight
        
        # Normalize to 0-100
        final_score = weighted_score / total_weight if total_weight > 0 else 0
        
        # Normalize to 0-100 scale
        final_score = min(100, final_score * 1.5)  # Scale up since each component maxes at 40
        
        return round(final_score, 1), {
            "components": details,
            "recommendations": self._generate_recommendations(final_score, record),
            "risk_level": self._interpret_score(final_score),
            "priority": self._calculate_priority(final_score, record)
        }
    
    def _interpret_score(self, score: float) -> str:
        if score >= 75: return "CRITICAL"
        if score >= 55: return "HIGH"
        if score >= 35: return "MODERATE"
        if score >= 15: return "LOW"
        return "MINIMAL"
    
    def _generate_recommendations(self, score: float, record: BuildingRecord) -> List[str]:
        """Generate actionable recommendations"""
        recs = []
        
        if record.building_age and record.building_age > 40:
            recs.append("🔴 URGENT: Pre-code building - Seismic assessment recommended")
        
        if record.material_type and any(m in record.material_type.lower() for m in ['stone', 'brick', 'unreinforced']):
            recs.append("⚠️ Unreinforced masonry - High priority for retrofit")
        
        if record.damage_level and any(d in record.damage_level.lower() for d in ['crack', 'damage']):
            recs.append("🔧 Address existing damage before retrofitting")
        
        if score >= 75:
            recs.append("🚨 IMMEDIATE: Priority for seismic retrofit programme")
        elif score >= 55:
            recs.append("📋 Schedule detailed engineering evaluation")
        else:
            recs.append("✓ Monitor condition, conduct 5-year review")
        
        return recs
    
    def _calculate_priority(self, score: float, record: BuildingRecord) -> int:
        """Priority ranking 1-10 (10 = highest)"""
        priority = int(score / 10)
        
        # Boost if seems to be residential/institutional
        if record.raw_data.get('building_type'):
            bt = str(record.raw_data.get('building_type')).lower()
            if any(t in bt for t in ['school', 'hospital', 'government', 'residential']):
                priority += 1
        
        return min(10, priority)

# =============================================
# 3. CLUSTERING & PATTERN ANALYSIS
# =============================================

class BuildingTypologyAnalyzer:
    """Identify building types and clusters"""
    
    TYPOLOGIES = {
        'Early Unreinforced': {
            'age_min': 40,
            'material': ['stone', 'brick', 'adobe', 'unreinforced'],
            'risk_baseline': 70
        },
        'Transition Era': {
            'age_min': 25, 'age_max': 40,
            'material': ['rc', 'concrete'],
            'risk_baseline': 45
        },
        'Modern Code': {
            'age_max': 25,
            'material': ['rc', 'steel'],
            'risk_baseline': 25
        },
        'High-Rise': {
            'stories_min': 8,
            'risk_baseline': 40
        },
    }
    
    @classmethod
    def classify(cls, record: BuildingRecord) -> str:
        """Classify building into typology"""
        
        if record.building_age and record.building_age > 40:
            if record.material_type and any(m in record.material_type.lower() for m in ['stone', 'brick']):
                return "Early Unreinforced"
        
        if record.building_age and 25 <= record.building_age <= 40:
            return "Transition Era"
        
        if record.stories and record.stories > 7:
            return "High-Rise"
        
        if record.building_age and record.building_age < 25:
            return "Modern Code"
        
        return "Unclassified"
    
    @classmethod
    def analyze_cluster(cls, records: List[BuildingRecord], typology: str) -> Dict[str, Any]:
        """Analyze cluster statistics"""
        if not records:
            return {}
        
        df = pd.DataFrame([{
            'building_id': r.building_id,
            'age': r.building_age or 0,
            'stories': r.stories or 0,
        } for r in records])
        
        return {
            'count': len(records),
            'avg_age': df['age'].mean(),
            'avg_stories': df['stories'].mean(),
            'age_range': (df['age'].min(), df['age'].max()),
        }

# =============================================
# 4. MAIN APPLICATION
# =============================================

class SeismicAnalyzerApp(ctk.CTk):
    def __init__(self):
        super().__init__()
        
        self.title("Seismic Building Vulnerability Analyzer - Advanced Edition")
        self.geometry("1600x950")
        
        # State
        self.df = None
        self.filtered_df = None
        self.building_records = []  # Typed records
        self.vulnerability_cache = {}
        self.display_limit = 50
        self.cols = []
        self.link_map = {}
        self.risk_engine = SeismicVulnerabilityIndex()
        
        # Field mapping
        self.field_mapping = self._detect_field_mapping()
        
        # Grid
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)
        
        self.create_sidebar()
        
        self.tab_view = ctk.CTkTabview(self, fg_color="transparent")
        self.tab_view.grid(row=0, column=1, sticky="nsew", padx=20, pady=10)
        
        self.tab_dashboard = self.tab_view.add("📊 Vulnerability Dashboard")
        self.tab_inspector = self.tab_view.add("🗂️ All Buildings")
        self.tab_comparison = self.tab_view.add("⚖️ Compare")
        self.tab_clustering = self.tab_view.add("🎯 Typology Analysis")
        self.tab_map = self.tab_view.add("🗺️ Map View")
        self.tab_help = self.tab_view.add("❓ Help & Guide")
        
        self.setup_dashboard()
        self.setup_inspector()
        self.setup_comparison()
        self.setup_clustering()
        self.setup_map()
        self.setup_help()
    
    def _detect_field_mapping(self) -> Dict[str, Optional[str]]:
        """Detect which columns map to seismic fields"""
        return {key: None for key in SEISMIC_FIELD_KEYWORDS}
    
    def create_sidebar(self):
        self.sidebar = ctk.CTkFrame(self, width=280, corner_radius=0, fg_color=NAVY)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_rowconfigure(20, weight=1)
        
        # Logo
        self.logo = ctk.CTkLabel(self.sidebar, text="SEISMIC\nANALYZER",
                                font=("Arial", 20, "bold"), text_color=WHITE)
        self.logo.grid(row=0, column=0, padx=20, pady=(30, 20))
        
        # File Section
        ctk.CTkLabel(self.sidebar, text="DATA SOURCE", font=("Arial", 11, "bold"),
                    text_color=AQUA, anchor="w").grid(row=1, column=0, padx=20, sticky="w")
        
        self.load_btn = ctk.CTkButton(self.sidebar, text="📂 Load Excel File",
                                     command=self.load_file, fg_color=AQUA, text_color=NAVY)
        self.load_btn.grid(row=2, column=0, padx=20, pady=5, sticky="ew")
        
        self.merge_btn = ctk.CTkButton(self.sidebar, text="🔗 Merge Datasets",
                                      command=self.merge_datasets, fg_color="#334E68")
        self.merge_btn.grid(row=3, column=0, padx=20, pady=5, sticky="ew")
        
        self.status_label = ctk.CTkLabel(self.sidebar, text="No File Loaded",
                                        text_color="gray", font=("Arial", 10))
        self.status_label.grid(row=4, column=0, padx=20, pady=(0, 20), sticky="w")
        
        # Analysis Tools
        ctk.CTkLabel(self.sidebar, text="ANALYSIS TOOLS", font=("Arial", 11, "bold"),
                    text_color=AQUA, anchor="w").grid(row=5, column=0, padx=20, sticky="w", pady=(10, 5))
        
        self.calc_vuln_btn = ctk.CTkButton(self.sidebar, text="🔬 Calculate Vulnerability",
                                          command=self.calculate_vulnerability,
                                          fg_color=MAROON, text_color=WHITE)
        self.calc_vuln_btn.grid(row=6, column=0, padx=20, pady=5, sticky="ew")
        
        self.cluster_btn = ctk.CTkButton(self.sidebar, text="📍 Analyze Typologies",
                                        command=self.analyze_clusters,
                                        fg_color=ORANGE, text_color=WHITE)
        self.cluster_btn.grid(row=7, column=0, padx=20, pady=5, sticky="ew")
        
        # Export Tools
        ctk.CTkLabel(self.sidebar, text="EXPORT & REPORTS", font=("Arial", 11, "bold"),
                    text_color=AQUA, anchor="w").grid(row=8, column=0, padx=20, sticky="w", pady=(10, 5))
        
        self.passport_btn = ctk.CTkButton(self.sidebar, text="📘 Passports + Risk",
                                         command=self.generate_passport_thread,
                                         fg_color=ORANGE)
        self.passport_btn.grid(row=9, column=0, padx=20, pady=5, sticky="ew")
        
        self.pdf_btn = ctk.CTkButton(self.sidebar, text="📄 Risk Report PDF",
                                    command=self.export_pdf, fg_color=WHITE, text_color=NAVY)
        self.pdf_btn.grid(row=10, column=0, padx=20, pady=5, sticky="ew")
        
        self.export_csv_btn = ctk.CTkButton(self.sidebar, text="📊 Export to CSV",
                                           command=self.export_csv, fg_color="#334E68")
        self.export_csv_btn.grid(row=11, column=0, padx=20, pady=5, sticky="ew")
        
        self.geojson_btn = ctk.CTkButton(self.sidebar, text="🌍 To QGIS (GeoJSON)",
                                        command=self.export_geojson, fg_color="#334E68")
        self.geojson_btn.grid(row=12, column=0, padx=20, pady=5, sticky="ew")
        
        # Visualization Tools
        ctk.CTkLabel(self.sidebar, text="VISUALIZATION", font=("Arial", 11, "bold"),
                    text_color=AQUA, anchor="w").grid(row=13, column=0, padx=20, sticky="w", pady=(10, 5))
        
        self.map_btn = ctk.CTkButton(self.sidebar, text="🗺️ Interactive Map",
                                    command=self.generate_interactive_map, fg_color=MAROON)
        self.map_btn.grid(row=14, column=0, padx=20, pady=5, sticky="ew")
        
        self.charts_btn = ctk.CTkButton(self.sidebar, text="📊 Analysis Charts",
                                       command=self.generate_risk_charts, fg_color=ORANGE)
        self.charts_btn.grid(row=15, column=0, padx=20, pady=5, sticky="ew")
        
        # Query Builder
        ctk.CTkLabel(self.sidebar, text="QUICK FILTER", font=("Arial", 11, "bold"),
                    text_color=AQUA, anchor="w").grid(row=16, column=0, padx=20, sticky="w", pady=(10, 5))
        
        self.query_col = ctk.CTkOptionMenu(self.sidebar, values=["Column"],
                                          fg_color="#334E68", button_color="#243B53")
        self.query_col.grid(row=17, column=0, padx=20, pady=5, sticky="ew")
        
        self.query_op = ctk.CTkOptionMenu(self.sidebar,
                                         values=["=", "!=", ">", "<", ">=", "<=", "contains"],
                                         fg_color="#334E68", button_color="#243B53")
        self.query_op.set("=")
        self.query_op.grid(row=18, column=0, padx=20, pady=5, sticky="ew")
        
        self.query_val = ctk.CTkEntry(self.sidebar, placeholder_text="Value...",
                                     fg_color="#102A43", border_color="#334E68", text_color=WHITE)
        self.query_val.grid(row=19, column=0, padx=20, pady=5, sticky="ew")
        
        filter_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        filter_frame.grid(row=20, column=0, padx=20, pady=10, sticky="ew")
        
        ctk.CTkButton(filter_frame, text="Apply", command=self.apply_filter,
                     width=100, fg_color=GREEN, text_color=NAVY).pack(side="left", padx=(0, 5))
        ctk.CTkButton(filter_frame, text="Reset", command=self.clear_filter,
                     width=60, fg_color=LIGHT_RED).pack(side="right")
    
    def setup_dashboard(self):
        """Vulnerability dashboard with risk overview"""
        self.tab_dashboard.grid_columnconfigure(0, weight=1)
        self.tab_dashboard.grid_rowconfigure(2, weight=1)
        
        # Stats Panel
        stats_frame = ctk.CTkFrame(self.tab_dashboard, fg_color=WHITE, corner_radius=10)
        stats_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=10)
        
        self.stat_total = ctk.CTkLabel(stats_frame, text="Buildings: 0",
                                      font=("Arial", 14, "bold"), text_color=NAVY)
        self.stat_total.pack(side="left", padx=20, pady=10)
        
        self.stat_critical = ctk.CTkLabel(stats_frame, text="Critical: 0",
                                         font=("Arial", 14, "bold"), text_color=LIGHT_RED)
        self.stat_critical.pack(side="left", padx=20, pady=10)
        
        self.stat_high = ctk.CTkLabel(stats_frame, text="High Risk: 0",
                                     font=("Arial", 14, "bold"), text_color=ORANGE)
        self.stat_high.pack(side="left", padx=20, pady=10)
        
        self.stat_moderate = ctk.CTkLabel(stats_frame, text="Moderate: 0",
                                         font=("Arial", 14, "bold"), text_color=YELLOW)
        self.stat_moderate.pack(side="left", padx=20, pady=10)
        
        # Chart Controls
        ctrl_frame = ctk.CTkFrame(self.tab_dashboard, fg_color=GREY, corner_radius=10)
        ctrl_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=10)
        
        self.chart_col = ctk.CTkOptionMenu(ctrl_frame, values=["Select Column"],
                                          fg_color=MAROON, text_color=WHITE)
        self.chart_col.pack(side="left", padx=10, pady=10)
        
        self.chart_type = ctk.CTkSegmentedButton(ctrl_frame,
                                               values=["Risk Distribution", "Age vs Risk", "Material Analysis"],
                                               selected_color=NAVY)
        self.chart_type.set("Risk Distribution")
        self.chart_type.pack(side="left", padx=10, pady=10)
        
        self.chart_btn = ctk.CTkButton(ctrl_frame, text="Generate",
                                      command=self.generate_advanced_chart, fg_color=GREEN)
        self.chart_btn.pack(side="left", padx=10, pady=10)
        
        # Chart Display
        self.chart_frame = ctk.CTkFrame(self.tab_dashboard, fg_color=WHITE, corner_radius=15)
        self.chart_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=10)
        
        self.empty_chart_lbl = ctk.CTkLabel(self.chart_frame,
                                           text="Load data and click 'Generate' to visualize vulnerability distribution",
                                           font=("Arial", 14), text_color="gray")
        self.empty_chart_lbl.place(relx=0.5, rely=0.5, anchor="center")
    
    def setup_inspector(self):
        """Detailed building records with vulnerability scores"""
        self.tab_inspector.grid_columnconfigure(0, weight=1)
        self.tab_inspector.grid_rowconfigure(0, weight=1)
        
        self.scroll_frame = ctk.CTkScrollableFrame(self.tab_inspector, fg_color="transparent",
                                                  label_text="Building Vulnerability Records")
        self.scroll_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)
    
    def setup_comparison(self):
        """Compare two buildings in detail"""
        self.tab_comparison.grid_columnconfigure(0, weight=1)
        self.tab_comparison.grid_columnconfigure(1, weight=1)
        self.tab_comparison.grid_rowconfigure(1, weight=1)
        
        ctrl = ctk.CTkFrame(self.tab_comparison, fg_color=GREY)
        ctrl.grid(row=0, column=0, columnspan=2, sticky="ew", padx=10, pady=10)
        
        self.comp_sel1 = ctk.CTkOptionMenu(ctrl, values=["Building A"], fg_color=NAVY)
        self.comp_sel1.pack(side="left", padx=10, pady=10)
        
        ctk.CTkLabel(ctrl, text="vs", font=("Arial", 12, "bold")).pack(side="left", padx=10)
        
        self.comp_sel2 = ctk.CTkOptionMenu(ctrl, values=["Building B"], fg_color=NAVY)
        self.comp_sel2.pack(side="left", padx=10, pady=10)
        
        self.comp_btn = ctk.CTkButton(ctrl, text="Compare Vulnerability",
                                     command=self.run_comparison, fg_color=MAROON)
        self.comp_btn.pack(side="left", padx=10, pady=10)
        
        self.comp_scroll = ctk.CTkScrollableFrame(self.tab_comparison, fg_color="white")
        self.comp_scroll.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=10, pady=10)
    
    def setup_clustering(self):
        """Building typology and cluster analysis"""
        self.tab_clustering.grid_columnconfigure(0, weight=1)
        self.tab_clustering.grid_rowconfigure(0, weight=1)
        
        btn_frame = ctk.CTkFrame(self.tab_clustering, fg_color=GREY)
        btn_frame.pack(fill="x", padx=10, pady=10)
        
        ctk.CTkButton(btn_frame, text="Analyze All Typologies",
                     command=self.analyze_clusters, fg_color=ORANGE).pack(side="left", padx=5, pady=5)
        
        self.cluster_scroll = ctk.CTkScrollableFrame(self.tab_clustering, fg_color="transparent")
        self.cluster_scroll.pack(fill="both", expand=True, padx=10, pady=10)
    
    def setup_map(self):
        """Geospatial vulnerability visualization"""
        self.tab_map.grid_columnconfigure(0, weight=1)
        self.tab_map.grid_rowconfigure(0, weight=1)
        
        map_frame = ctk.CTkFrame(self.tab_map, fg_color=WHITE)
        map_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        ctk.CTkButton(map_frame, text="🌍 Generate Vulnerability Map",
                     command=self.generate_map, fg_color=NAVY, font=("Arial", 16, "bold"),
                     height=50).place(relx=0.5, rely=0.5, anchor="center")
        
        ctk.CTkLabel(map_frame, text="Requires 'Latitude' and 'Longitude' columns. Uses vulnerability scores for pin colors.",
                    text_color="gray").place(relx=0.5, rely=0.6, anchor="center")
    
    def setup_help(self):
        """Comprehensive user guide"""
        self.tab_help.grid_columnconfigure(0, weight=1)
        self.tab_help.grid_rowconfigure(0, weight=1)
        
        help_text = ctk.CTkScrollableFrame(self.tab_help, fg_color="white")
        help_text.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        help_content = """
═══════════════════════════════════════════════════════════════════════════════
SEISMIC BUILDING VULNERABILITY ANALYZER - COMPLETE USER GUIDE
═══════════════════════════════════════════════════════════════════════════════

📖 TABLE OF CONTENTS
1. Getting Started
2. Loading Data
3. Understanding Vulnerability Scores
4. Using Each Feature
5. Interpreting Results
6. Troubleshooting

═══════════════════════════════════════════════════════════════════════════════
1. GETTING STARTED
═══════════════════════════════════════════════════════════════════════════════

The Seismic Building Vulnerability Analyzer is designed to assess building earthquake
risk based on structural characteristics. It works with Excel files containing building
survey data collected from your seismic survey forms.

KEY CONCEPTS:
• Vulnerability Score: 0-100 (higher = greater seismic risk)
• Risk Levels: Minimal (0-15), Low (15-35), Moderate (35-55), High (55-75), Critical (75-100)
• Building Typology: Category based on age, material, and structure

═══════════════════════════════════════════════════════════════════════════════
2. LOADING DATA
═══════════════════════════════════════════════════════════════════════════════

STEP 1: Prepare Your Excel File
   • Should have columns for: Building ID, material type, construction age, stories, etc.
   • DateTime columns are OK (app handles them)
   • Try to include: Building ID, Year Built, Material, Stories, Foundation Type, Damage Level

STEP 2: Load File
   • Click "📂 Load Excel File" button in left sidebar
   • Select your Excel file (.xlsx or .xls)
   • App automatically detects relevant seismic indicator columns
   • Status shows: "filename.xlsx (123 buildings loaded)"

STEP 3: Merge Multiple Files (Optional)
   • Click "🔗 Merge Datasets" to combine multiple Excel files
   • Selects folder, merges all .xlsx files
   • Removes duplicates by Building ID (keeps latest)

═══════════════════════════════════════════════════════════════════════════════
3. UNDERSTANDING VULNERABILITY SCORES
═══════════════════════════════════════════════════════════════════════════════

The analyzer calculates a COMPOSITE VULNERABILITY INDEX based on:

┌─────────────────────────────────────────────────────────────────┐
│ FACTOR              │ WEIGHT │ EXPLANATION                      │
├─────────────────────────────────────────────────────────────────┤
│ Building Age        │ 35%    │ Pre-1980 = higher risk           │
│ Material Type       │ 25%    │ Unreinforced stone/brick = high  │
│ Foundation Type     │ 20%    │ Shallow/poor = higher risk       │
│ Height (Stories)    │ 10%    │ Taller = more amplification      │
│ Current Condition   │ 10%    │ Existing damage = higher risk    │
└─────────────────────────────────────────────────────────────────┘

RISK LEVEL INTERPRETATION:

🟢 MINIMAL (0-15): Modern, well-maintained building. Meets earthquake codes.
   → Recommendation: Standard monitoring, 5-year condition review

🟡 LOW (15-35): Relatively safe but may need inspection.
   → Recommendation: Non-urgent inspection recommended within 2 years

🟠 MODERATE (35-55): Notable seismic vulnerabilities detected.
   → Recommendation: Detailed engineering assessment recommended

🔴 HIGH (55-75): Significant seismic risk. Intervention needed.
   → Recommendation: Retrofit planning should commence

🔴🔴 CRITICAL (75-100): Immediate seismic risk. Emergency mitigation needed.
   → Recommendation: Urgent structural assessment and retrofit program

═══════════════════════════════════════════════════════════════════════════════
4. USING EACH FEATURE
═══════════════════════════════════════════════════════════════════════════════

📊 VULNERABILITY DASHBOARD
   • Shows: Total buildings, count by risk level
   • Charts: Risk distribution, age vs risk correlation, material analysis
   • Use when: Want quick overview of portfolio vulnerability

   STEPS:
   1. Load data (Data will automatically calculate)
   2. Choose visualization type from dropdown
   3. Select column to analyze (optional, for custom charts)
   4. Click "Generate" to visualize

🗂️ ALL BUILDINGS (Data Inspector)
   • Shows: Detailed card for each building
   • Displays: Building ID, vulnerability score, risk level, recommendations
   • Click building ID to expand details

   FEATURES:
   • Cards color-coded by risk level
   • Shows all captured building parameters
   • Photos/links embedded if available

⚖️ COMPARE TWO BUILDINGS
   • Compares vulnerability factors between buildings
   • Shows: Side-by-side component scores, key differences
   
   STEPS:
   1. Select building from left dropdown
   2. Select building from right dropdown
   3. Click "Compare Vulnerability"
   4. Yellow-highlighted rows = different values

   USE FOR: Finding why similar buildings have different risk scores

🎯 TYPOLOGY ANALYSIS
   • Clusters buildings into types: Early Unreinforced, Transition Era, Modern Code, High-Rise
   • Shows: Count, average age, stories, and aggregate statistics
   
   STEPS:
   1. Click "Analyze All Typologies"
   2. Review cluster summaries
   3. Identify which typology needs most attention

   KEY INSIGHTS:
   • "Early Unreinforced" typically CRITICAL risk
   • "Modern Code" typically LOW-MINIMAL risk
   • "High-Rise" amplifies earthquake shaking

🗺️ MAP VIEW
   • Shows: Geographic distribution of vulnerability
   • Color coding: Red = critical, Yellow = high, Green = low
   • Requires: Latitude & Longitude columns in data
   
   STEPS:
   1. Ensure data has Latitude/Longitude columns
   2. Click "Generate Vulnerability Map"
   3. Saves as "seismic_vulnerability_map.html"
   4. Opens in browser for exploration

   USE FOR: Identifying geographic vulnerability clusters

📘 PASSPORTS + RISK
   • Generates Word document with building profiles + vulnerability scores
   • Includes: Photos, assessment data, risk interpretation, recommendations
   
   STEPS:
   1. Select buildings (filtered if using Query Builder)
   2. Click "📘 Passports + Risk"
   3. Choose output location
   4. Open generated .docx files

📄 RISK REPORT PDF
   • Exports vulnerability statistics and charts to PDF
   • Includes: Risk distribution charts, summary statistics
   
   STEPS:
   1. Click "📄 Risk Report PDF"
   2. Choose location and filename
   3. Review in any PDF reader

📊 EXPORT TO CSV
   • Exports complete dataset with vulnerability scores
   • Useful for: Processing in spreadsheet software, creating custom analyses
   
   STEPS:
   1. Click "📊 Export to CSV"
   2. Open in Excel/Google Sheets

🌍 TO QGIS (GeoJSON)
   • Exports data in GeoJSON format for GIS software
   • Can visualize spatial patterns in QGIS
   • Requires: Latitude/Longitude columns

═══════════════════════════════════════════════════════════════════════════════
5. QUICK FILTER (Left Sidebar)
═══════════════════════════════════════════════════════════════════════════════

Skip to buildings with specific characteristics:

EXAMPLE FILTERS:
   • Column: "Risk Level" | Operator: "=" | Value: "CRITICAL" 
     → Shows only critical risk buildings

   • Column: "Building Age" | Operator: ">" | Value: "40"
     → Shows buildings older than 40 years

   • Column: "Material" | Operator: "contains" | Value: "stone"
     → Shows all stone buildings

   • Column: "Damage Level" | Operator: "contains" | Value: "crack"
     → Shows buildings with recorded cracks

AFTER FILTERING:
   • Dashboard updates automatically
   • Export, comparison, and map respect filtered data
   • Click "Reset" to remove filters

═══════════════════════════════════════════════════════════════════════════════
6. INTERPRETING RESULTS & RECOMMENDATIONS
═══════════════════════════════════════════════════════════════════════════════

The app generates specific recommendations based on vulnerability patterns:

COMMON RECOMMENDATIONS YOU'LL SEE:

1. "URGENT: Pre-code building - Seismic assessment recommended"
   → Building built before modern earthquake code
   → ACTION: Schedule detailed structural assessment (2-4 weeks)

2. "Unreinforced masonry - High priority for retrofit"
   → Stone/brick without seatings or connections
   → ACTION: Develop retrofit plan (3-6 months), budget engineer

3. "Address existing damage before retrofitting"
   → Building shows cracks or deterioration
   → ACTION: Repair structural damage first (1-2 months), then retrofit

4. "IMMEDIATE: Priority for seismic retrofit programme"
   → Very high vulnerability, cannot wait
   → ACTION: Emergency mobilization, proceed with retrofit immediately

5. "Schedule detailed engineering evaluation"
   → Moderate risk, needs professional assessment
   → ACTION: Hire structural engineer for site visit (1 month)

6. "Monitor condition, conduct 5-year review"
   → Low or minimal risk
   → ACTION: Schedule inspection every 5 years, maintain building

═══════════════════════════════════════════════════════════════════════════════
7. TROUBLESHOOTING
═══════════════════════════════════════════════════════════════════════════════

PROBLEM: "Column not found" error
→ SOLUTION: Right-click on column header in your Excel file, rename to match
   expected field names (Building ID, Material, Age, Stories, Foundation, etc.)

PROBLEM: Map doesn't show
→ SOLUTION: Ensure Excel file has "Latitude"/"Longitude" columns with numeric values

PROBLEM: Risk scores seem wrong
→ SOLUTION: Check that numeric columns (age, stories) are formatted as numbers, not text

PROBLEM: Can't load file
→ SOLUTION: Ensure Excel file is closed (not open in another program)

PROBLEM: Some buildings missing from comparison
→ SOLUTION: Make sure Building ID column has unique values for each building

═══════════════════════════════════════════════════════════════════════════════
8. DATA FORMAT REQUIREMENTS
═══════════════════════════════════════════════════════════════════════════════

For best results, ensure your Excel file has these columns (or rename to match):

CRITICAL (Analysis won't work without):
  □ Building ID / Bid / Code

HIGHLY RECOMMENDED:
  □ Year Built / Age / Construction Year  → (for age-based risk)
  □ Material Type / Construction Material  → (for material-based risk)
  □ Stories / Levels / Height / Floors    → (for height-based risk)
  □ Foundation Type / Foundation / Base   → (for foundation assessment)
  □ Damage Level / Condition / Damage     → (for condition assessment)

OPTIONAL (For enhanced analysis):
  □ Latitude / Longitude  → (for map visualization)
  □ Building Type         → (for typology analysis)
  □ Photos / Images       → (for visual documentation)
  □ Comments / Notes      → (for qualitative risk signals)

COLUMN FORMAT:
  • Year Built: 1995, 2010 (numeric, not "1995-01-01")
  • Stories: 2, 5, 10 (numeric whole numbers)
  • Latitude: 28.5555 (decimal degrees, -90 to 90)
  • Longitude: 84.2222 (decimal degrees, -180 to 180)
  • Material: "Unreinforced Stone", "Reinforced Concrete", "Steel Frame"
  • Damage: "None", "Minor Cracks", "Moderate Damage", "Severe"

═══════════════════════════════════════════════════════════════════════════════
9. ADVANCED TIPS & TRICKS
═══════════════════════════════════════════════════════════════════════════════

TIP 1: Merge datasets
   → Combine surveys from multiple days/areas
   → App automatically deduplicates based on Building ID

TIP 2: Use comments field for qualitative assessment
   → Add notes like "severe cracking throughout"
   → App analyzes comments to boost risk where appropriate

TIP 3: Export data for further analysis
   → Save as CSV, open in Python/Excel for custom analysis
   → GeoJSON for spatial analysis in QGIS

TIP 4: Filter before exporting
   → Filter to just "CRITICAL" buildings, then export to PDF
   → Creates targeted retrofit priority list

TIP 5: Use typology analysis for intervention planning
   → All "Early Unreinforced" buildings? Plan retrofit program
   → All "Modern Code"? Focus on maintenance instead

═══════════════════════════════════════════════════════════════════════════════
10. CONTACT & FEEDBACK
═══════════════════════════════════════════════════════════════════════════════

Questions about scores? Check your building parameters against the criteria above.
Unexpected results? Verify your data column formats match requirements.
Feature requests? This app integrates with your seismic survey forms for real-time analysis.

═══════════════════════════════════════════════════════════════════════════════
"""
        
        label = ctk.CTkLabel(help_text, text=help_content, justify="left", 
                           font=("Courier", 14), text_color="black")
        label.pack(padx=10, pady=10)
    
    def load_file(self):
        """Load Excel file with smart column detection - handles both standard and transposed formats"""
        file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx;*.xls")])
        if not file_path:
            return
        
        try:
            df_raw = pd.read_excel(file_path)
            
            # Check if data is transposed format (FIELD NAME in first column, building IDs in headers)
            if df_raw.columns[0].lower() == 'field name' or df_raw.iloc[0, 0] in ['FIELD NAME', 'SURVEYOR NAME', 'BUILDING ID']:
                # Convert transposed to standard format
                self.df = self._convert_transposed_format(df_raw)
            else:
                # Standard format
                self.df = df_raw
            
            self.filtered_df = self.df.copy()
            self.cols = [c for c in self.df.columns if c != "Created_at"]
            
            # Auto-detect seismic fields
            self._auto_detect_columns()
            
            # Parse into typed records
            self._parse_building_records()
            
            # Setup UI
            self.status_label.configure(text=os.path.basename(file_path)[:30] + "...")
            self.query_col.configure(values=self.cols)
            if self.cols:
                self.query_col.set(self.cols[0])
            
            # Setup comparison
            id_col = self.field_mapping.get('building_id') or self._find_column('ID')
            if id_col:
                ids = self.df[id_col].astype(str).unique()[:100]
                if hasattr(self, 'comp_sel1'):
                    self.comp_sel1.configure(values=list(ids))
                    self.comp_sel2.configure(values=list(ids))
                    if len(ids) > 0:
                        self.comp_sel1.set(ids[0])
                    if len(ids) > 1:
                        self.comp_sel2.set(ids[1])
            
            # Auto-calculate vulnerability
            self.calculate_vulnerability()
            if hasattr(self, 'render_cards'):
                self.render_cards(self.filtered_df.head(self.display_limit))
            
            messagebox.showinfo("Success", f"Loaded {len(self.df)} buildings successfully!")
            
        except Exception as e:
            logger.error(f"Load error: {e}", exc_info=True)
            messagebox.showerror("Error", f"Load Failed: {str(e)}")
    
    def _convert_transposed_format(self, df: pd.DataFrame) -> pd.DataFrame:
        """Convert transposed Excel format to standard format"""
        # Set first column as index (field names)
        df = df.set_index('FIELD NAME')
        # Transpose so buildings become rows
        df = df.T
        df = df.reset_index()
        df.rename(columns={'index': 'BUILDING_ID'}, inplace=True)
        return df
    
    def _auto_detect_columns(self):
        """Smart detection of seismic indicator columns"""
        for seismic_field, keywords in SEISMIC_FIELD_KEYWORDS.items():
            for col in self.cols:
                col_lower = col.lower()
                if any(kw in col_lower for kw in keywords):
                    self.field_mapping[seismic_field] = col
                    logger.info(f"Detected {seismic_field} -> {col}")
                    break
    
    def _find_column(self, keyword: str) -> Optional[str]:
        """Find column containing keyword"""
        for col in self.cols:
            if keyword.lower() in col.lower():
                return col
        return None
    
    def _parse_building_records(self):
        """Convert rows to typed BuildingRecord objects"""
        self.building_records = []
        
        for idx, row in self.df.iterrows():
            # Helper to safely extract value from row
            def safe_extract(field_key, fallback_keywords=None):
                """Safely get column value from row"""
                col_name = None
                
                # Try field mapping first
                if field_key in self.field_mapping and self.field_mapping[field_key]:
                    col_name = self.field_mapping[field_key]
                
                # Try fallback keywords if no mapping
                if not col_name and fallback_keywords:
                    for keyword in fallback_keywords:
                        col_name = self._find_column(keyword)
                        if col_name:
                            break
                
                # Extract value if column found
                if col_name and col_name in row.index:
                    val = row[col_name]
                    return val if pd.notna(val) else None
                
                return None
            
            # Extract field values using mapping
            building_id = safe_extract('building_id', ['ID']) or f"B{idx}"
            material_type = safe_extract('material_type', ['material'])
            building_age = DataValidator.safe_float(
                safe_extract('building_age', ['year', 'age', 'built']) or
                self._calculate_age_from_year(safe_extract('building_age', ['year', 'age', 'built']))
            )
            stories = DataValidator.safe_int(safe_extract('stories', ['stories', 'floors', 'levels']))
            foundation_type = safe_extract('foundation_type', ['foundation'])
            damage_level = safe_extract('damage_level', ['damage', 'condition'])
            latitude = DataValidator.safe_float(safe_extract('location', ['latitude', 'lat']))
            longitude = DataValidator.safe_float(safe_extract('location', ['longitude', 'lon']))
            
            record = BuildingRecord(
                building_id=str(building_id),
                survey_date=str(safe_extract('location', ['survey', 'date']) or ""),
                material_type=str(material_type or ""),
                building_age=building_age,
                stories=stories,
                foundation_type=str(foundation_type or ""),
                damage_level=str(damage_level or ""),
                latitude=latitude,
                longitude=longitude,
                raw_data=row.to_dict(),
            )
            
            # Combine all comments fields
            comments = []
            for col in self.cols:
                if 'comment' in col.lower():
                    val = row[col] if col in row.index else None
                    if val and pd.notna(val):
                        comments.append(str(val))
            record.comments = " ".join(comments)
            
            self.building_records.append(record)
    
    def _calculate_age_from_year(self, year_val) -> Optional[float]:
        """Calculate age from year built"""
        try:
            if pd.isna(year_val) or year_val == '':
                return None
            year = int(float(year_val))
            current_year = datetime.now().year
            age = current_year - year
            return age if 0 <= age <= 500 else None
        except:
            return None
    
    def calculate_vulnerability(self):
        """Calculate vulnerability for all loaded buildings"""
        if not self.building_records:
            messagebox.showwarning("Error", "No buildings loaded yet")
            return
        
        try:
            self.vulnerability_cache = {}
            for record in self.building_records:
                score, details = self.risk_engine.calculate(record)
                self.vulnerability_cache[record.building_id] = {
                    'score': score,
                    'level': details.get('risk_level', 'UNKNOWN'),
                    'recommendations': details.get('recommendations', []),
                    'priority': details.get('priority', 0),
                    'components': details.get('components', {}),
                }
            
            # Add to dataframe
            if self.df is not None:
                id_col = self.field_mapping.get('building_id') or self._find_column('ID') or self.cols[0]
                self.df['Vulnerability Score'] = self.df[id_col].apply(
                    lambda x: self.vulnerability_cache.get(str(x), {}).get('score', 0)
                )
                self.df['Risk Level'] = self.df[id_col].apply(
                    lambda x: self.vulnerability_cache.get(str(x), {}).get('level', 'UNKNOWN')
                )
            
            self._update_dashboard_stats()
            self.render_cards(self.filtered_df.head(self.display_limit))
            messagebox.showinfo("Complete", f"Calculated vulnerability for {len(self.building_records)} buildings")
            
        except Exception as e:
            logger.error(f"Calculation error: {e}", exc_info=True)
            messagebox.showerror("Error", f"Calculation failed: {str(e)}")
    
    def _update_dashboard_stats(self):
        """Update vulnerability statistics display"""
        if not self.vulnerability_cache:
            return
        
        levels = [v.get('level') for v in self.vulnerability_cache.values()]
        total = len(levels)
        critical = sum(1 for l in levels if l == 'CRITICAL')
        high = sum(1 for l in levels if l == 'HIGH')
        moderate = sum(1 for l in levels if l == 'MODERATE')
        
        self.stat_total.configure(text=f"Buildings: {total}")
        self.stat_critical.configure(text=f"Critical: {critical}")
        self.stat_high.configure(text=f"High Risk: {high}")
        self.stat_moderate.configure(text=f"Moderate: {moderate}")
    
    def analyze_clusters(self):
        """Analyze building typologies and clusters"""
        if not self.building_records:
            messagebox.showwarning("Error", "No buildings loaded")
            return
        
        try:
            # Classify buildings
            typologies = {}
            for record in self.building_records:
                typo = BuildingTypologyAnalyzer.classify(record)
                if typo not in typologies:
                    typologies[typo] = []
                typologies[typo].append(record)
            
            # Display results
            for widget in self.cluster_scroll.winfo_children():
                widget.destroy()
            
            for typo_name, records in typologies.items():
                frame = ctk.CTkFrame(self.cluster_scroll, fg_color=GREY, corner_radius=10)
                frame.pack(fill="x", padx=10, pady=10)
                
                # Header
                header = ctk.CTkLabel(frame, text=f"{typo_name} ({len(records)} buildings)",
                                     font=("Arial", 12, "bold"), text_color=NAVY)
                header.pack(anchor="w", padx=15, pady=(10, 5))
                
                # Stats
                avg_vuln = np.mean([self.vulnerability_cache.get(r.building_id, {}).get('score', 0) for r in records])
                avg_age = np.mean([r.building_age or 0 for r in records])
                
                stats = ctk.CTkLabel(frame,
                                    text=f"Avg Vulnerability: {avg_vuln:.1f} | Avg Age: {avg_age:.0f} years",
                                    font=("Arial", 10), text_color="gray")
                stats.pack(anchor="w", padx=15, pady=5)
                
                # Building list
                for record in records[:10]:  # Show first 10
                    vuln = self.vulnerability_cache.get(record.building_id, {})
                    bg_color = "#FFE5E5" if vuln.get('level') == 'CRITICAL' else GREY
                    
                    item = ctk.CTkLabel(frame, text=f"  • {record.building_id} - {vuln.get('level')} ({vuln.get('score', 0):.0f})",
                                       font=("Arial", 9), text_color="black")
                    item.pack(anchor="w", padx=25, pady=2)
                
                if len(records) > 10:
                    more = ctk.CTkLabel(frame, text=f"  ... and {len(records) - 10} more",
                                       font=("Arial", 9), text_color="gray")
                    more.pack(anchor="w", padx=25, pady=2)
            
            # Global clustering (KMeans) and anomaly detection (IsolationForest)
            try:
                from sklearn.cluster import KMeans
                from sklearn.ensemble import IsolationForest

                # Prepare feature matrix: [vulnerability, age, stories]
                X = []
                ids = []
                for r in self.building_records:
                    vid = self.vulnerability_cache.get(r.building_id, {}).get('score', 0) or 0
                    age = r.building_age or 0
                    stories = r.stories or 0
                    X.append([vid, age, stories])
                    ids.append(r.building_id)

                import numpy as _np
                X = _np.array(X)
                n_clusters = min(4, max(2, int(_np.sqrt(len(X)))))
                kmeans = KMeans(n_clusters=n_clusters, random_state=42).fit(X)
                clusters = kmeans.labels_

                iso = IsolationForest(contamination=0.05, random_state=42).fit(X)
                anomalies = iso.predict(X)  # -1 = anomaly

                # Summarize clusters
                cluster_frame = ctk.CTkFrame(self.cluster_scroll, fg_color=GREY, corner_radius=8)
                cluster_frame.pack(fill="x", padx=10, pady=8)
                ctk.CTkLabel(cluster_frame, text=f"KMeans Clustering (n={n_clusters})", font=("Arial", 12, "bold"), text_color=NAVY).pack(anchor="w", padx=10, pady=8)
                for ci in range(n_clusters):
                    members = [ids[i] for i, c in enumerate(clusters) if c == ci]
                    avg_score = _np.mean([X[i,0] for i, c in enumerate(clusters) if c == ci])
                    ctk.CTkLabel(cluster_frame, text=f"Cluster {ci}: {len(members)} buildings | Avg Score: {avg_score:.1f}", text_color="black").pack(anchor="w", padx=15)

                # Anomalies
                anom_ids = [ids[i] for i, a in enumerate(anomalies) if a == -1]
                if anom_ids:
                    anom_frame = ctk.CTkFrame(self.cluster_scroll, fg_color="#FFF5F0", corner_radius=8)
                    anom_frame.pack(fill="x", padx=10, pady=8)
                    ctk.CTkLabel(anom_frame, text=f"Anomalies detected: {len(anom_ids)}", font=("Arial", 12, "bold"), text_color=LIGHT_RED).pack(anchor="w", padx=10, pady=6)
                    for aid in anom_ids[:20]:
                        ctk.CTkLabel(anom_frame, text=f"  • {aid}", text_color="black").pack(anchor="w", padx=15)

            except Exception as e:
                logger.warning(f"Clustering skipped (sklearn missing or error): {e}")
                
        except Exception as e:
            logger.error(f"Cluster analysis error: {e}", exc_info=True)
            messagebox.showerror("Error", f"Analysis failed: {str(e)}")
    
    def apply_filter(self):
        """Apply query builder filter"""
        if self.df is None:
            return
        
        col = self.query_col.get()
        op = self.query_op.get()
        val = self.query_val.get()
        
        if not val:
            return
        
        try:
            temp_df = self.df.copy()
            
            if op in [">", "<", ">=", "<="]:
                try:
                    temp_df[col] = pd.to_numeric(temp_df[col])
                    val = float(val)
                except:
                    messagebox.showerror("Type Error", "Column must be numeric")
                    return
            
            if op == "=":
                self.filtered_df = temp_df[temp_df[col].astype(str) == str(val)]
            elif op == "!=":
                self.filtered_df = temp_df[temp_df[col].astype(str) != str(val)]
            elif op == ">":
                self.filtered_df = temp_df[temp_df[col] > val]
            elif op == "<":
                self.filtered_df = temp_df[temp_df[col] < val]
            elif op == ">=":
                self.filtered_df = temp_df[temp_df[col] >= val]
            elif op == "<=":
                self.filtered_df = temp_df[temp_df[col] <= val]
            elif op == "contains":
                self.filtered_df = temp_df[temp_df[col].astype(str).str.contains(val, case=False, na=False)]
            
            self.render_cards(self.filtered_df.head(self.display_limit))
            self.tab_view.set("🗂️ All Buildings")
            
        except Exception as e:
            messagebox.showerror("Query Error", str(e))
    
    def clear_filter(self):
        """Reset filters"""
        if self.df is None:
            return
        self.filtered_df = self.df.copy()
        self.query_val.delete(0, 'end')
        self.render_cards(self.filtered_df.head(self.display_limit))
    
    def render_cards(self, dataframe: pd.DataFrame):
        """Render building cards with vulnerability display"""
        for widget in self.scroll_frame.winfo_children():
            widget.destroy()
        
        if dataframe.empty:
            ctk.CTkLabel(self.scroll_frame, text="No records match query.").pack(pady=20)
            return
        
        id_col = self.field_mapping.get('building_id') or self._find_column('ID') or self.cols[0]
        
        for index, row in dataframe.iterrows():
            bid = str(row[id_col])
            vuln = self.vulnerability_cache.get(bid, {})
            score = vuln.get('score', 0)
            level = vuln.get('level', 'UNKNOWN')
            
            # Color by risk level
            if level == 'CRITICAL':
                bg_color = "#FFE5E5"
                border_color = LIGHT_RED
            elif level == 'HIGH':
                bg_color = "#FFF5E5"
                border_color = ORANGE
            elif level == 'MODERATE':
                bg_color = "#FFFFF0"
                border_color = YELLOW
            else:
                bg_color = WHITE
                border_color = GREEN
            
            card = ctk.CTkFrame(self.scroll_frame, fg_color=bg_color, corner_radius=10, border_width=2, border_color=border_color)
            card.pack(fill="x", padx=10, pady=8)
            
            # Header with score
            header = ctk.CTkFrame(card, fg_color=border_color, height=35, corner_radius=10)
            header.pack(fill="x")
            
            ctk.CTkLabel(header, text=f"{bid}",
                        text_color=WHITE, font=("Arial", 12, "bold")).pack(side="left", padx=15, pady=5)
            
            ctk.CTkLabel(header, text=f"Risk: {level} ({score:.0f}/100)",
                        text_color=WHITE, font=("Arial", 11, "bold")).pack(side="right", padx=15, pady=5)
            
            # Content
            content = ctk.CTkFrame(card, fg_color="transparent")
            content.pack(fill="x", padx=15, pady=10)
            
            # Building parameters
            for key in ['material_type', 'building_age', 'stories', 'foundation_type', 'damage_level']:
                if key in self.field_mapping and self.field_mapping[key]:
                    col = self.field_mapping[key]
                    val = row.get(col)
                    if val and pd.notna(val):
                        label = ctk.CTkLabel(content, text=f"{key.replace('_', ' ').title()}: {val}",
                                            font=("Arial", 10), text_color="black")
                        label.pack(anchor="w", pady=2)
            
            # Recommendations
            if vuln.get('recommendations'):
                rec_text = "  •  ".join(vuln['recommendations'][:2])
                rec_label = ctk.CTkLabel(content, text=f"Recommendations: {rec_text}...",
                                        font=("Arial", 9), text_color="#333333")
                rec_label.pack(anchor="w", pady=(10, 0))
    
    def run_comparison(self):
        """Compare two buildings' vulnerability"""
        id1 = self.comp_sel1.get()
        id2 = self.comp_sel2.get()
        
        if id1 == id2:
            messagebox.showwarning("Same ID", "Select different buildings")
            return
        
        if id1 not in self.vulnerability_cache or id2 not in self.vulnerability_cache:
            messagebox.showwarning("Error", "Buildings not found in analysis")
            return
        
        vuln1 = self.vulnerability_cache[id1]
        vuln2 = self.vulnerability_cache[id2]
        
        for widget in self.comp_scroll.winfo_children():
            widget.destroy()
        
        # Header
        header = ctk.CTkFrame(self.comp_scroll, fg_color="transparent")
        header.pack(fill="x", pady=10, padx=10)
        
        ctk.CTkLabel(header, text="Building ID", font=("Arial", 12, "bold"), width=150).pack(side="left", padx=5)
        ctk.CTkLabel(header, text=id1, font=("Arial", 12, "bold"), text_color=NAVY, width=200).pack(side="left", padx=5)
        ctk.CTkLabel(header, text=id2, font=("Arial", 12, "bold"), text_color=MAROON, width=200).pack(side="left", padx=5)
        
        # Scores
        score_row = ctk.CTkFrame(self.comp_scroll, fg_color=GREY)
        score_row.pack(fill="x", pady=5, padx=10)
        
        ctk.CTkLabel(score_row, text="Vulnerability", width=150, anchor="w").pack(side="left", padx=5)
        ctk.CTkLabel(score_row, text=f"{vuln1.get('score', 0):.1f}", width=200).pack(side="left", padx=5)
        ctk.CTkLabel(score_row, text=f"{vuln2.get('score', 0):.1f}", width=200).pack(side="left", padx=5)
        
        # Risk level
        level_row = ctk.CTkFrame(self.comp_scroll, fg_color=GREY)
        level_row.pack(fill="x", pady=5, padx=10)
        
        ctk.CTkLabel(level_row, text="Risk Level", width=150, anchor="w").pack(side="left", padx=5)
        ctk.CTkLabel(level_row, text=vuln1.get('level', 'UNKNOWN'), width=200, text_color=LIGHT_RED).pack(side="left", padx=5)
        ctk.CTkLabel(level_row, text=vuln2.get('level', 'UNKNOWN'), width=200, text_color=LIGHT_RED).pack(side="left", padx=5)
        
        # Component breakdown
        ctk.CTkLabel(self.comp_scroll, text="Risk Components Breakdown",
                    font=("Arial", 12, "bold"), text_color=NAVY).pack(anchor="w", padx=10, pady=(15, 5))
        
        components1 = vuln1.get('components', {})
        components2 = vuln2.get('components', {})
        
        for comp_name in components1.keys():
            comp_row = ctk.CTkFrame(self.comp_scroll, fg_color=GREY)
            comp_row.pack(fill="x", pady=3, padx=10)
            
            score1 = components1[comp_name].get('score', 0)
            score2 = components2[comp_name].get('score', 0)
            
            ctk.CTkLabel(comp_row, text=comp_name.replace('Calculator', ''), width=150, anchor="w").pack(side="left", padx=5)
            ctk.CTkLabel(comp_row, text=f"{score1:.0f}", width=200).pack(side="left", padx=5)
            ctk.CTkLabel(comp_row, text=f"{score2:.0f}", width=200).pack(side="left", padx=5)
    
    def generate_advanced_chart(self):
        """Generate advanced vulnerability charts"""
        if self.df is None or 'Vulnerability Score' not in self.df.columns:
            messagebox.showwarning("Error", "Calculate vulnerability first")
            return
        
        try:
            import matplotlib.pyplot as plt
            from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg
            
            for widget in self.chart_frame.winfo_children():
                widget.destroy()
            
            chart_type = self.chart_type.get()
            
            if chart_type == "Risk Distribution":
                fig, ax = plt.subplots(figsize=(8, 5), dpi=100)
                self.filtered_df['Risk Level'].value_counts().plot(kind='bar', color=[LIGHT_RED, ORANGE, YELLOW, GREEN], ax=ax)
                ax.set_title("Risk Level Distribution", fontsize=12, fontweight='bold')
                ax.set_xlabel("Risk Level")
                ax.set_ylabel("Count")
            
            elif chart_type == "Age vs Risk":
                if self.field_mapping.get('building_age'):
                    fig, ax = plt.subplots(figsize=(8, 5), dpi=100)
                    age_col = self.field_mapping.get('building_age')
                    self.filtered_df.plot(x=age_col, y='Vulnerability Score', kind='scatter', ax=ax, color=NAVY, s=100)
                    ax.set_title("Building Age vs Vulnerability", fontsize=12, fontweight='bold')
                    ax.set_xlabel("Building Age (years)")
                    ax.set_ylabel("Vulnerability Score")
                else:
                    messagebox.showwarning("Missing Data", "Age column not detected")
                    return
            
            elif chart_type == "Material Analysis":
                if self.field_mapping.get('material_type'):
                    fig, ax = plt.subplots(figsize=(8, 5), dpi=100)
                    mat_col = self.field_mapping.get('material_type')
                    self.filtered_df.groupby(mat_col)['Vulnerability Score'].mean().sort_values(ascending=False).plot(kind='barh', color=NAVY, ax=ax)
                    ax.set_title("Average Vulnerability by Material", fontsize=12, fontweight='bold')
                    ax.set_xlabel("Avg Vulnerability Score")
                else:
                    messagebox.showwarning("Missing Data", "Material column not detected")
                    return
            
            plt.tight_layout()
            
            canvas = FigureCanvasTkAgg(fig, master=self.chart_frame)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="both", expand=True)
            
        except Exception as e:
            logger.error(f"Chart error: {e}", exc_info=True)
            messagebox.showerror("Error", f"Chart generation failed: {str(e)}")
    
    def export_pdf(self):
        """Export vulnerability report to PDF"""
        if self.df is None:
            return
        
        try:
            from matplotlib.backends.backend_pdf import PdfPages
            import matplotlib.pyplot as plt
            
            save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF", "*.pdf")])
            if not save_path:
                return
            
            with PdfPages(save_path) as pdf:
                # Page 1: Summary statistics
                fig, ((ax1, ax2), (ax3, ax4)) = plt.subplots(2, 2, figsize=(11, 8.5))
                
                # Risk distribution
                self.filtered_df['Risk Level'].value_counts().plot(kind='pie', ax=ax1, autopct='%1.1f%%')
                ax1.set_title("Risk Level Distribution")
                
                # Risk histogram
                self.filtered_df['Vulnerability Score'].hist(bins=20, ax=ax2, color=NAVY)
                ax2.set_title("Vulnerability Score Distribution")
                ax2.set_xlabel("Score")
                
                # Stats
                stats_text = f"""VULNERABILITY ANALYSIS REPORT
Buildings Analyzed: {len(self.filtered_df)}

RISK BREAKDOWN:
• Critical (75-100): {sum(self.filtered_df['Risk Level'] == 'CRITICAL')}
• High (55-75): {sum(self.filtered_df['Risk Level'] == 'HIGH')}
• Moderate (35-55): {sum(self.filtered_df['Risk Level'] == 'MODERATE')}
• Low (15-35): {sum(self.filtered_df['Risk Level'] == 'LOW')}
• Minimal (0-15): {sum(self.filtered_df['Risk Level'] == 'MINIMAL')}

AVERAGE METRICS:
• Mean Vulnerability: {self.filtered_df['Vulnerability Score'].mean():.1f}
• Median Vulnerability: {self.filtered_df['Vulnerability Score'].median():.1f}
"""
                ax3.text(0.1, 0.5, stats_text, fontsize=11, verticalalignment='center', family='monospace')
                ax3.axis('off')
                
                # Top vulnerable
                top_vuln = self.filtered_df.nlargest(5, 'Vulnerability Score')
                top_text = "TOP 5 MOST VULNERABLE BUILDINGS:\n\n"
                id_col = self.field_mapping.get('building_id') or self._find_column('ID') or self.cols[0]
                for idx, row in top_vuln.iterrows():
                    top_text += f"• {row[id_col]}: {row['Vulnerability Score']:.0f}/100 ({row['Risk Level']})\n"
                
                ax4.text(0.1, 0.5, top_text, fontsize=11, verticalalignment='center', family='monospace')
                ax4.axis('off')
                
                plt.tight_layout()
                pdf.savefig(fig)
                plt.close()
            
            messagebox.showinfo("Success", "PDF Report Generated!")
            
        except Exception as e:
            logger.error(f"PDF export error: {e}", exc_info=True)
            messagebox.showerror("Error", f"PDF generation failed: {str(e)}")
    
    def export_csv(self):
        """Export data with vulnerability scores"""
        if self.df is None:
            return
        
        try:
            save_path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV", "*.csv")])
            if save_path:
                self.filtered_df.to_csv(save_path, index=False)
                messagebox.showinfo("Success", f"Exported {len(self.filtered_df)} records to CSV")
        except Exception as e:
            messagebox.showerror("Error", str(e))
    
    def export_geojson(self):
        """Export to GeoJSON format"""
        if self.df is None:
            return
        
        try:
            lat_col = self.field_mapping.get('location') or self._find_column('Lat')
            lon_col = self.field_mapping.get('location') or self._find_column('Lon')
            
            if not lat_col or not lon_col:
                messagebox.showwarning("Missing Data", "No latitude/longitude columns found")
                return
            
            features = []
            for _, row in self.filtered_df.iterrows():
                try:
                    props = {k: v for k, v in row.to_dict().items() if isinstance(v, (str, int, float, bool))}
                    
                    feature = {
                        "type": "Feature",
                        "geometry": {
                            "type": "Point",
                            "coordinates": [row[lon_col], row[lat_col]]
                        },
                        "properties": props
                    }
                    features.append(feature)
                except:
                    continue
            
            save_path = filedialog.asksaveasfilename(defaultextension=".geojson")
            if save_path:
                with open(save_path, 'w') as f:
                    json.dump({"type": "FeatureCollection", "features": features}, f)
                messagebox.showinfo("Success", f"Exported {len(features)} points to GeoJSON")
        
        except Exception as e:
            messagebox.showerror("Error", str(e))
    
    def generate_interactive_map(self):
        """Generate interactive map with building locations"""
        try:
            if not self.building_records or not any(r.latitude and r.longitude for r in self.building_records):
                messagebox.showerwarning("Alert", "No GPS coordinates available for mapping")
                return
            
            # Create map centered on average location
            lats = [r.latitude for r in self.building_records if r.latitude]
            lons = [r.longitude for r in self.building_records if r.longitude]
            
            if not lats or not lons:
                messagebox.showerror("Error", "No valid GPS coordinates")
                return
            
            center_lat = np.mean(lats)
            center_lon = np.mean(lons)
            
            m = folium.Map(location=[center_lat, center_lon], zoom_start=11)
            
            # Color code by risk level
            risk_colors = {
                'CRITICAL': 'red',
                'HIGH': 'orange',
                'MODERATE': 'yellow',
                'LOW': 'blue',
                'MINIMAL': 'green'
            }
            
            for record in self.building_records:
                if record.latitude and record.longitude:
                    # Get vulnerability score if calculated
                    vuln_score = self.vulnerability_cache.get(record.building_id, {}).get('score', 0)
                    risk_level = self._interpret_score(vuln_score)
                    color = risk_colors.get(risk_level, 'gray')
                    
                    popup_text = f"""
                    <b>{record.building_id}</b><br>
                    Surveyor: {record.surveyor_name}<br>
                    Risk Score: {vuln_score:.1f}/100<br>
                    Risk Level: {risk_level}<br>
                    Material: {record.material_type or 'N/A'}<br>
                    """
                    
                    folium.CircleMarker(
                        location=[record.latitude, record.longitude],
                        radius=8,
                        popup=folium.Popup(popup_text, max_width=300),
                        color=color,
                        fill=True,
                        fillOpacity=0.7
                    ).add_to(m)
            
            # Save and open
            temp_file = tempfile.NamedTemporaryFile(mode='w', suffix='.html', delete=False)
            m.save(temp_file.name)
            webbrowser.open('file://' + temp_file.name)
            messagebox.showinfo("Success", "Interactive map opened in browser")
        
        except Exception as e:
            logger.error(f"Map generation error: {e}")
            messagebox.showerror("Error", f"Map generation failed: {str(e)}")
    
    def generate_risk_charts(self):
        """Generate interactive charts for data analysis"""
        try:
            if not self.building_records:
                messagebox.showerwarning("Alert", "No data to chart")
                return
            
            # Create a new window for charts
            chart_window = ctk.CTkToplevel(self)
            chart_window.title("Building Analysis Charts")
            chart_window.geometry("1400x900")
            
            # Create figure with subplots
            fig = Figure(figsize=(14, 9), dpi=100)
            
            # Extract data
            records_df = pd.DataFrame([{
                'building_id': r.building_id,
                'material': r.material_type or 'Unknown',
                'age': r.building_age or 0,
                'stories': r.stories or 0,
                'district': r.district or 'Unknown',
                'risk_score': self.vulnerability_cache.get(r.building_id, {}).get('score', 0),
                'risk_level': self._interpret_score(self.vulnerability_cache.get(r.building_id, {}).get('score', 0))
            } for r in self.building_records])
            
            # Chart 1: Risk Level Distribution (Pie)
            ax1 = fig.add_subplot(2, 3, 1)
            risk_counts = records_df['risk_level'].value_counts()
            colors_pie = ['#FF4136', '#FF851B', '#FFDC00', '#0074D9', '#2ECC40']
            ax1.pie(risk_counts.values, labels=risk_counts.index, autopct='%1.1f%%', colors=colors_pie[:len(risk_counts)])
            ax1.set_title("Buildings by Risk Level", fontweight='bold')
            
            # Chart 2: Material Type Distribution
            ax2 = fig.add_subplot(2, 3, 2)
            material_counts = records_df['material'].value_counts()
            ax2.barh(material_counts.index, material_counts.values, color='#39CCCC')
            ax2.set_title("Material Types", fontweight='bold')
            ax2.set_xlabel("Count")
            
            # Chart 3: Average Risk by Material
            ax3 = fig.add_subplot(2, 3, 3)
            mat_risk = records_df.groupby('material')['risk_score'].mean().sort_values(ascending=False)
            ax3.bar(range(len(mat_risk)), mat_risk.values, color='#85144B')
            ax3.set_xticks(range(len(mat_risk)))
            ax3.set_xticklabels(mat_risk.index, rotation=45, ha='right')
            ax3.set_title("Average Risk by Material", fontweight='bold')
            ax3.set_ylabel("Risk Score")
            
            # Chart 4: Buildings by District
            ax4 = fig.add_subplot(2, 3, 4)
            district_counts = records_df['district'].value_counts()
            ax4.bar(range(len(district_counts)), district_counts.values, color='#2ECC40')
            ax4.set_xticks(range(len(district_counts)))
            ax4.set_xticklabels(district_counts.index, rotation=45, ha='right')
            ax4.set_title("Buildings by District", fontweight='bold')
            ax4.set_ylabel("Count")
            
            # Chart 5: Building Age Distribution
            ax5 = fig.add_subplot(2, 3, 5)
            ax5.hist(records_df['age'], bins=20, color='#001F3F', edgecolor='white')
            ax5.set_title("Age Distribution", fontweight='bold')
            ax5.set_xlabel("Years")
            ax5.set_ylabel("Frequency")
            
            # Chart 6: Risk vs Age Scatter
            ax6 = fig.add_subplot(2, 3, 6)
            scatter = ax6.scatter(records_df['age'], records_df['risk_score'], c=records_df['risk_score'], 
                                 cmap='RdYlGn_r', s=100, alpha=0.6)
            ax6.set_title("Risk Score vs Building Age", fontweight='bold')
            ax6.set_xlabel("Age (Years)")
            ax6.set_ylabel("Risk Score")
            fig.colorbar(scatter, ax=ax6)
            
            fig.tight_layout()
            
            # Embed in tkinter
            canvas = FigureCanvasTkAgg(fig, master=chart_window)
            canvas.draw()
            canvas.get_tk_widget().pack(fill="both", expand=True)
            
        except Exception as e:
            logger.error(f"Chart generation error: {e}")
            messagebox.showerror("Error", f"Chart generation failed: {str(e)}")
    
    def _interpret_score(self, score: float) -> str:
        """Interpret vulnerability score"""
        if score >= 75: return "CRITICAL"
        if score >= 55: return "HIGH"
        if score >= 35: return "MODERATE"
        if score >= 15: return "LOW"
        return "MINIMAL"
    
    def generate_passport_thread(self):
        """Generate passports in background thread"""
        threading.Thread(target=self.generate_passports, daemon=True).start()
    
    def generate_passports(self):
        """Generate building passports with vulnerability assessments"""
        if self.df is None:
            messagebox.showwarning("Error", "No data loaded")
            return
        
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            
            save_path = filedialog.asksaveasfilename(defaultextension=".docx",
                                                    filetypes=[("Word Document", "*.docx")])
            if not save_path:
                return
            
            doc = Document()
            doc.add_heading('Seismic Building Vulnerability Passports', 0)
            doc.add_paragraph(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
            doc.add_paragraph(f'Buildings analyzed: {len(self.filtered_df)}')
            doc.add_page_break()
            
            id_col = self.field_mapping.get('building_id') or self._find_column('ID') or self.cols[0]
            
            for index, row in self.filtered_df.iterrows():
                bid = str(row[id_col])
                vuln = self.vulnerability_cache.get(bid, {})
                
                doc.add_heading(f'Building: {bid}', level=1)
                
                # Vulnerability section
                doc.add_heading('Seismic Vulnerability Assessment', level=2)
                doc.add_paragraph(f'Vulnerability Score: {vuln.get("score", 0):.1f} / 100')
                doc.add_paragraph(f'Risk Level: {vuln.get("level", "UNKNOWN")}')
                doc.add_paragraph(f'Priority Ranking: {vuln.get("priority", 0)}/10')
                
                # Component breakdown
                if vuln.get('components'):
                    doc.add_heading('Risk Breakdown by Component', level=3)
                    for comp_name, comp_detail in vuln['components'].items():
                        score = comp_detail.get('score', 0)
                        interp = comp_detail.get('interpretation', '')
                        doc.add_paragraph(f'{comp_name}: {score:.0f} points - {interp}')
                
                # Recommendations
                if vuln.get('recommendations'):
                    doc.add_heading('Recommendations', level=3)
                    for rec in vuln['recommendations']:
                        doc.add_paragraph(rec, style='List Bullet')
                
                # Building details table
                doc.add_heading('Building Parameters', level=2)
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Parameter'
                hdr_cells[1].text = 'Value'
                
                for col in self.cols:
                    if col not in ['Vulnerability Score', 'Risk Level']:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(col)
                        row_cells[1].text = str(row[col])
                
                doc.add_page_break()
            
            doc.save(save_path)
            messagebox.showinfo("Success", "Passports generated successfully!")
        
        except Exception as e:
            logger.error(f"Passport generation error: {e}", exc_info=True)
            messagebox.showerror("Error", f"Generation failed: {str(e)}")
    
    def generate_map(self):
        """Generate interactive vulnerability map"""
        if self.df is None:
            return
        
        try:
            import folium
            
            lat_col = self._find_column('Lat')
            lon_col = self._find_column('Lon')
            
            if not lat_col or not lon_col:
                messagebox.showwarning("Missing Data", "Latitude/Longitude columns required")
                return
            
            center_lat = self.filtered_df[lat_col].mean()
            center_lon = self.filtered_df[lon_col].mean()
            
            m = folium.Map(location=[center_lat, center_lon], zoom_start=12)
            
            id_col = self.field_mapping.get('building_id') or self._find_column('ID') or self.cols[0]
            
            for _, row in self.filtered_df.iterrows():
                try:
                    bid = str(row[id_col])
                    vuln = self.vulnerability_cache.get(bid, {})
                    score = vuln.get('score', 0)
                    level = vuln.get('level', 'UNKNOWN')
                    
                    # Color by risk
                    if level == 'CRITICAL':
                        color = 'red'
                    elif level == 'HIGH':
                        color = 'orange'
                    elif level == 'MODERATE':
                        color = 'yellow'
                    else:
                        color = 'green'
                    
                    popup_text = f"<b>{bid}</b><br>Risk: {level}<br>Score: {score:.0f}/100"
                    folium.Marker(
                        location=[row[lat_col], row[lon_col]],
                        popup=popup_text,
                        icon=folium.Icon(color=color, icon='exclamation' if level in ['CRITICAL', 'HIGH'] else 'info-sign')
                    ).add_to(m)
                
                except:
                    continue
            
            map_path = os.path.abspath("seismic_vulnerability_map.html")
            m.save(map_path)
            
            import webbrowser
            webbrowser.open('file://' + map_path)
            messagebox.showinfo("Map Generated", f"Saved to {map_path}")
        
        except ImportError:
            messagebox.showerror("Dependency Error", "Please install folium: pip install folium")
        except Exception as e:
            logger.error(f"Map generation error: {e}", exc_info=True)
            messagebox.showerror("Error", str(e))
    
    def merge_datasets(self):
        """Merge multiple Excel files"""
        folder_path = filedialog.askdirectory(title="Select Folder with Excel Files")
        if not folder_path:
            return
        
        try:
            import glob
            
            all_files = glob.glob(os.path.join(folder_path, "*.xlsx"))
            if not all_files:
                messagebox.showwarning("Empty", "No .xlsx files found")
                return
            
            df_list = [pd.read_excel(f) for f in all_files]
            merged_df = pd.concat(df_list, ignore_index=True)
            
            id_col = next((c for c in merged_df.columns if "ID" in c), None)
            if id_col:
                merged_df.drop_duplicates(subset=id_col, keep='last', inplace=True)
            
            save_path = filedialog.asksaveasfilename(defaultextension=".xlsx")
            if save_path:
                merged_df.to_excel(save_path, index=False)
                messagebox.showinfo("Success", "Merged successfully")
        
        except Exception as e:
            messagebox.showerror("Error", str(e))

if __name__ == "__main__":
    try:
        app = SeismicAnalyzerApp()
        app.mainloop()
    except Exception as e:
        logger.error(f"App crash: {e}", exc_info=True)
        ctypes.windll.user32.MessageBoxW(0, f"Critical Error:\n{traceback.format_exc()}", "Crash", 0x10)
