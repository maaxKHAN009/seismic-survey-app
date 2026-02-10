import customtkinter as ctk
from tkinter import filedialog, messagebox
import os
import sys
import traceback
import ctypes
import threading
import requests  # For robust image downloading
import shutil

# --- Configuration ---
ctk.set_appearance_mode("Light")
ctk.set_default_color_theme("dark-blue")

# --- Styles ---
NAVY = "#001F3F"
AQUA = "#39CCCC"
MAROON = "#85144B"
GREEN = "#2ECC40"
YELLOW = "#FFDC00"
WHITE = "#FFFFFF"
GREY = "#F5F5F5"

class SeismicAnalyzerApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        # Window Setup
        self.title("Seismic Research Data Analyzer - Ultimate Edition")
        self.geometry("1400x900")
        
        # State
        self.df = None
        self.filtered_df = None
        self.display_limit = 50
        self.cols = []
        self.link_map = {} 

        # --- Layout Grid ---
        self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(0, weight=1)

        # 1. Sidebar (Controls)
        self.create_sidebar()

        # 2. Main Display Area (Tabs)
        self.tab_view = ctk.CTkTabview(self, fg_color="transparent")
        self.tab_view.grid(row=0, column=1, sticky="nsew", padx=20, pady=10)
        
        self.tab_dashboard = self.tab_view.add("📊 Analytics")
        self.tab_inspector = self.tab_view.add("🗂️ Data Inspector")
        self.tab_compare = self.tab_view.add("⚖️ Compare")
        self.tab_map = self.tab_view.add("🗺️ Map View")
        
        # Setup Tabs
        self.setup_dashboard()
        self.setup_inspector()
        self.setup_comparison()
        self.setup_map()

    def create_sidebar(self):
        self.sidebar = ctk.CTkFrame(self, width=300, corner_radius=0, fg_color=NAVY)
        self.sidebar.grid(row=0, column=0, sticky="nsew")
        self.sidebar.grid_rowconfigure(16, weight=1)

        self.logo = ctk.CTkLabel(self.sidebar, text="SEISMIC\nINTELLIGENCE", font=("Arial", 24, "bold"), text_color=WHITE)
        self.logo.grid(row=0, column=0, padx=20, pady=(30, 20))

        # --- File Section ---
        self.lbl_file = ctk.CTkLabel(self.sidebar, text="DATA SOURCE", font=("Arial", 11, "bold"), text_color=AQUA, anchor="w")
        self.lbl_file.grid(row=1, column=0, padx=20, pady=(10, 0), sticky="w")

        self.load_btn = ctk.CTkButton(self.sidebar, text="📂 Load Excel File", command=self.load_file, 
                                      fg_color=AQUA, text_color=NAVY, font=("Arial", 12, "bold"), hover_color="#2eb8b8")
        self.load_btn.grid(row=2, column=0, padx=20, pady=5, sticky="ew")

        self.merge_btn = ctk.CTkButton(self.sidebar, text="🔗 Merge Datasets", command=self.merge_datasets,
                                       fg_color="#334E68", text_color=WHITE, font=("Arial", 12, "bold"))
        self.merge_btn.grid(row=3, column=0, padx=20, pady=5, sticky="ew")

        self.status_label = ctk.CTkLabel(self.sidebar, text="No File Loaded", text_color="gray", font=("Arial", 10))
        self.status_label.grid(row=4, column=0, padx=20, pady=(0, 20), sticky="w")

        # --- Advanced Tools ---
        self.lbl_tools = ctk.CTkLabel(self.sidebar, text="ADVANCED TOOLS", font=("Arial", 11, "bold"), text_color=AQUA, anchor="w")
        self.lbl_tools.grid(row=5, column=0, padx=20, pady=(10, 5), sticky="w")

        self.calc_btn = ctk.CTkButton(self.sidebar, text="🧮 Calculate Risk", command=self.calculate_risk,
                                      fg_color=MAROON, text_color=WHITE, font=("Arial", 12, "bold"))
        self.calc_btn.grid(row=6, column=0, padx=20, pady=5, sticky="ew")

        # Export Grid
        self.export_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        self.export_frame.grid(row=7, column=0, padx=20, pady=5, sticky="ew")
        
        self.btn_geojson = ctk.CTkButton(self.export_frame, text="🌍 To QGIS", width=120, command=self.export_geojson, fg_color="#334E68")
        self.btn_geojson.pack(side="left", padx=(0, 5))
        
        self.btn_imgs = ctk.CTkButton(self.export_frame, text="📷 Dwnld Pics", width=120, command=self.download_images_thread, fg_color="#334E68")
        self.btn_imgs.pack(side="right")

        self.passport_btn = ctk.CTkButton(self.sidebar, text="📘 Generate Passports", command=self.generate_passport_thread,
                                     fg_color="#FF851B", text_color=NAVY, font=("Arial", 12, "bold"))
        self.passport_btn.grid(row=8, column=0, padx=20, pady=5, sticky="ew")

        self.pdf_btn = ctk.CTkButton(self.sidebar, text="📄 Export PDF Report", command=self.export_pdf,
                                     fg_color=WHITE, text_color=NAVY, font=("Arial", 12, "bold"))
        self.pdf_btn.grid(row=9, column=0, padx=20, pady=5, sticky="ew")

        # --- Query Builder ---
        self.lbl_query = ctk.CTkLabel(self.sidebar, text="QUERY BUILDER", font=("Arial", 11, "bold"), text_color=AQUA, anchor="w")
        self.lbl_query.grid(row=10, column=0, padx=20, pady=(20, 5), sticky="w")

        self.query_col = ctk.CTkOptionMenu(self.sidebar, values=["Field"], fg_color="#334E68", button_color="#243B53", text_color=WHITE)
        self.query_col.grid(row=11, column=0, padx=20, pady=5, sticky="ew")

        self.query_op = ctk.CTkOptionMenu(self.sidebar, values=["=", "!=", ">", "<", ">=", "<=", "contains"], fg_color="#334E68", button_color="#243B53", text_color=WHITE)
        self.query_op.set("=")
        self.query_op.grid(row=12, column=0, padx=20, pady=5, sticky="ew")

        self.query_val = ctk.CTkEntry(self.sidebar, placeholder_text="Value...", fg_color="#102A43", border_color="#334E68", text_color=WHITE)
        self.query_val.grid(row=13, column=0, padx=20, pady=5, sticky="ew")

        self.filter_frame = ctk.CTkFrame(self.sidebar, fg_color="transparent")
        self.filter_frame.grid(row=14, column=0, padx=20, pady=10, sticky="ew")
        
        self.apply_btn = ctk.CTkButton(self.filter_frame, text="Apply", command=self.apply_filter, width=100, fg_color=GREEN, text_color=NAVY)
        self.apply_btn.pack(side="left", padx=(0, 5))
        
        self.clear_btn = ctk.CTkButton(self.filter_frame, text="Reset", command=self.clear_filter, width=60, fg_color="#FF4136", text_color=WHITE)
        self.clear_btn.pack(side="right")

    def setup_dashboard(self):
        self.tab_dashboard.grid_columnconfigure(0, weight=1)
        self.tab_dashboard.grid_rowconfigure(2, weight=1)

        self.viz_ctrl_frame = ctk.CTkFrame(self.tab_dashboard, fg_color=GREY, corner_radius=10)
        self.viz_ctrl_frame.grid(row=0, column=0, sticky="ew", padx=10, pady=10)

        self.chart_col = ctk.CTkOptionMenu(self.viz_ctrl_frame, values=["Select Column"], fg_color=MAROON, text_color=WHITE)
        self.chart_col.pack(side="left", padx=10, pady=10)

        self.chart_type = ctk.CTkSegmentedButton(self.viz_ctrl_frame, values=["Bar", "Pie", "Heatmap", "Correlation"], selected_color=NAVY)
        self.chart_type.set("Bar")
        self.chart_type.pack(side="left", padx=10, pady=10)

        self.chart_gen_btn = ctk.CTkButton(self.viz_ctrl_frame, text="Generate Graph", command=self.generate_chart, fg_color=GREEN, text_color=NAVY)
        self.chart_gen_btn.pack(side="left", padx=10, pady=10)

        self.stats_frame = ctk.CTkFrame(self.tab_dashboard, height=60, fg_color=WHITE, corner_radius=10)
        self.stats_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=5)
        
        self.stat_total = ctk.CTkLabel(self.stats_frame, text="Records: 0", font=("Arial", 16, "bold"), text_color=NAVY)
        self.stat_total.pack(side="left", padx=20, pady=10)
        
        self.stat_info = ctk.CTkLabel(self.stats_frame, text="Filter: None", font=("Arial", 12), text_color="gray")
        self.stat_info.pack(side="right", padx=20, pady=10)

        self.chart_frame = ctk.CTkFrame(self.tab_dashboard, fg_color=WHITE, corner_radius=15)
        self.chart_frame.grid(row=2, column=0, sticky="nsew", padx=10, pady=10)
        self.empty_chart_lbl = ctk.CTkLabel(self.chart_frame, text="Select parameters to visualize data", font=("Arial", 16), text_color="gray")
        self.empty_chart_lbl.place(relx=0.5, rely=0.5, anchor="center")

    def setup_inspector(self):
        self.tab_inspector.grid_columnconfigure(0, weight=1)
        self.tab_inspector.grid_rowconfigure(0, weight=1)
        self.scroll_frame = ctk.CTkScrollableFrame(self.tab_inspector, fg_color="transparent", label_text="Building List")
        self.scroll_frame.grid(row=0, column=0, sticky="nsew", padx=5, pady=5)

    def setup_comparison(self):
        self.tab_compare.grid_columnconfigure(0, weight=1)
        self.tab_compare.grid_columnconfigure(1, weight=1)
        self.tab_compare.grid_rowconfigure(1, weight=1)

        self.comp_ctrl = ctk.CTkFrame(self.tab_compare, fg_color=GREY)
        self.comp_ctrl.grid(row=0, column=0, columnspan=2, sticky="ew", padx=10, pady=10)
        
        self.comp_sel1 = ctk.CTkOptionMenu(self.comp_ctrl, values=["ID A"], fg_color=NAVY)
        self.comp_sel1.pack(side="left", padx=10, pady=10)
        
        self.comp_btn = ctk.CTkButton(self.comp_ctrl, text="Compare", command=self.run_comparison, fg_color=MAROON)
        self.comp_btn.pack(side="left", padx=10, pady=10)
        
        self.comp_sel2 = ctk.CTkOptionMenu(self.comp_ctrl, values=["ID B"], fg_color=NAVY)
        self.comp_sel2.pack(side="left", padx=10, pady=10)

        self.comp_scroll = ctk.CTkScrollableFrame(self.tab_compare, fg_color="white")
        self.comp_scroll.grid(row=1, column=0, columnspan=2, sticky="nsew", padx=10, pady=10)

    def setup_map(self):
        self.tab_map.grid_columnconfigure(0, weight=1)
        self.tab_map.grid_rowconfigure(0, weight=1)
        
        self.map_frame = ctk.CTkFrame(self.tab_map, fg_color=WHITE)
        self.map_frame.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)
        
        self.map_btn = ctk.CTkButton(self.map_frame, text="🌍 Generate Geospatial Map", command=self.generate_map, 
                                     font=("Arial", 16, "bold"), height=50, fg_color=NAVY)
        self.map_btn.place(relx=0.5, rely=0.5, anchor="center")
        
        self.map_info = ctk.CTkLabel(self.map_frame, text="Requires 'Lat' and 'Lon' columns in data", text_color="gray")
        self.map_info.place(relx=0.5, rely=0.6, anchor="center")

    # --- LAZY LOADING LOGIC ---
    def load_file(self):
        file_path = filedialog.askopenfilename(filetypes=[("Excel Files", "*.xlsx;*.xls")])
        if not file_path: return

        try:
            import pandas as pd
            import openpyxl

            # 1. Load Data
            self.df = pd.read_excel(file_path)
            self.link_map = {} 
            
            # 2. Extract Hyperlinks
            try:
                wb = openpyxl.load_workbook(file_path)
                ws = wb.active
                for row_idx, row in enumerate(ws.iter_rows(min_row=2), start=0):
                    if row_idx >= len(self.df): break 
                    for col_idx, cell in enumerate(row):
                        if cell.hyperlink and cell.hyperlink.target:
                            col_name = self.df.columns[col_idx]
                            self.link_map[(row_idx, col_name)] = cell.hyperlink.target
            except Exception as e:
                print(f"Link warning: {e}")

            # 3. Setup
            self.filtered_df = self.df.copy()
            # Relaxed filter: Only exclude specific system columns
            exclude = ["Created_at"]
            self.cols = [c for c in self.df.columns if c not in exclude]
            
            self.status_label.configure(text=os.path.basename(file_path)[:30] + "...")
            self.update_stats()
            
            self.query_col.configure(values=self.cols)
            if self.cols: self.query_col.set(self.cols[0])
            self.chart_col.configure(values=self.cols)
            if self.cols: self.chart_col.set(self.cols[0])
            
            # Comparison Setup
            id_col = next((c for c in self.df.columns if "ID" in str(c)), self.df.columns[0])
            ids = self.df[id_col].astype(str).tolist()
            self.comp_sel1.configure(values=ids)
            self.comp_sel2.configure(values=ids)
            if ids: 
                self.comp_sel1.set(ids[0])
                self.comp_sel2.set(ids[min(1, len(ids)-1)])

            self.chart_gen_btn.configure(state="normal")
            self.render_cards(self.filtered_df.head(self.display_limit))
            messagebox.showinfo("Success", "Data Loaded Successfully")

        except Exception as e:
            messagebox.showerror("Error", f"Load Failed: {str(e)}")

    def merge_datasets(self):
        folder_path = filedialog.askdirectory(title="Select Folder with Excel Files")
        if not folder_path: return
        
        import pandas as pd
        import glob
        
        try:
            all_files = glob.glob(os.path.join(folder_path, "*.xlsx"))
            if not all_files:
                return messagebox.showwarning("Empty", "No .xlsx files found in folder.")
                
            df_list = []
            for f in all_files:
                df_list.append(pd.read_excel(f))
                
            merged_df = pd.concat(df_list, ignore_index=True)
            
            id_col = next((c for c in merged_df.columns if "ID" in str(c)), None)
            if id_col:
                merged_df.drop_duplicates(subset=id_col, keep='last', inplace=True)
            
            save_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")])
            if save_path:
                merged_df.to_excel(save_path, index=False)
                messagebox.showinfo("Success", "Merged successfully.")
                
        except Exception as e:
            messagebox.showerror("Merge Error", str(e))

    def apply_filter(self):
        if self.df is None: return
        import pandas as pd
        
        col = self.query_col.get()
        op = self.query_op.get()
        val = self.query_val.get()
        
        if not val: return

        try:
            temp_df = self.df.copy()
            # Numeric conversion for inequalities
            if op in [">", "<", ">=", "<="]:
                try:
                    temp_df[col] = pd.to_numeric(temp_df[col])
                    val = float(val)
                except:
                    messagebox.showerror("Type Error", "Column must be numeric for > < operators.")
                    return

            if op == "=": self.filtered_df = temp_df[temp_df[col].astype(str) == str(val)]
            elif op == "!=": self.filtered_df = temp_df[temp_df[col].astype(str) != str(val)]
            elif op == ">": self.filtered_df = temp_df[temp_df[col] > val]
            elif op == "<": self.filtered_df = temp_df[temp_df[col] < val]
            elif op == ">=": self.filtered_df = temp_df[temp_df[col] >= val]
            elif op == "<=": self.filtered_df = temp_df[temp_df[col] <= val]
            elif op == "contains": self.filtered_df = temp_df[temp_df[col].astype(str).str.contains(val, case=False, na=False)]

            self.update_stats(f"{col} {op} {val}")
            self.render_cards(self.filtered_df.head(self.display_limit))
            self.tab_view.set("🗂️ Data Inspector")

        except Exception as e:
            messagebox.showerror("Query Error", str(e))

    def clear_filter(self):
        if self.df is None: return
        self.filtered_df = self.df.copy()
        self.query_val.delete(0, 'end')
        self.update_stats()
        self.render_cards(self.filtered_df.head(self.display_limit))

    def update_stats(self, active_filter="None"):
        count = len(self.filtered_df)
        self.stat_total.configure(text=f"Visible Records: {count}")
        self.stat_info.configure(text=f"Filter: {active_filter}")

    def calculate_risk(self):
        if self.df is None: return
        import pandas as pd
        risk_keywords = ['crack', 'unreinforced', 'poor', 'damage', 'mud', 'stone']
        try:
            self.filtered_df['Risk Score'] = 0
            for col in self.cols:
                if self.filtered_df[col].dtype == 'object':
                    for keyword in risk_keywords:
                        matches = self.filtered_df[col].astype(str).str.contains(keyword, case=False, na=False)
                        self.filtered_df.loc[matches, 'Risk Score'] += 10
            
            self.render_cards(self.filtered_df.head(self.display_limit))
            self.tab_view.set("🗂️ Data Inspector")
            messagebox.showinfo("Analysis Complete", "Risk Scores calculated.")
        except Exception as e:
            messagebox.showerror("Calc Error", str(e))

    def export_pdf(self):
        if self.df is None: return
        from matplotlib.backends.backend_pdf import PdfPages
        import matplotlib.pyplot as plt
        try:
            save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])
            if not save_path: return
            with PdfPages(save_path) as pdf:
                col = self.chart_col.get()
                if col != "Select Column":
                    fig, ax = plt.subplots(figsize=(8, 6))
                    counts = self.filtered_df[col].value_counts().nlargest(10)
                    counts.plot(kind='bar', color=NAVY, ax=ax)
                    ax.set_title(f"Distribution of {col}", fontsize=14)
                    plt.tight_layout()
                    pdf.savefig(fig)
                    plt.close()
            messagebox.showinfo("Export", "PDF Generated!")
        except Exception as e:
            messagebox.showerror("Export Error", str(e))

    def generate_passport_thread(self):
        threading.Thread(target=self.generate_passports).start()

    def generate_passports(self):
        if self.df is None: return
        from docx import Document
        from docx.shared import Inches
        import requests 
        import shutil
        
        save_path = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if not save_path: return
        
        try:
            doc = Document()
            doc.add_heading('Seismic Vulnerability Building Passports', 0)
            
            for index, row in self.filtered_df.iterrows():
                id_col = next((c for c in self.df.columns if "ID" in str(c)), self.df.columns[0])
                b_id = str(row[id_col])
                
                doc.add_heading(f"Building: {b_id}", level=1)
                
                table = doc.add_table(rows=1, cols=2)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Parameter'
                hdr_cells[1].text = 'Value'
                
                for col in self.cols:
                    if "Photo" in col or "IMG" in col: continue
                    row_cells = table.add_row().cells
                    row_cells[0].text = str(col)
                    row_cells[1].text = str(row[col])
                
                doc.add_paragraph("\nPhotographic Evidence:", style='Heading 2')
                
                for col in self.cols:
                    url = self.link_map.get((index, col))
                    val = str(row[col])
                    if not url and val.startswith('http'): url = val
                    
                    if url and "http" in str(url):
                        try:
                            # FIX: Use requests with User-Agent to avoid blocks
                            img_path = f"temp_{b_id}_{col}.jpg"
                            response = requests.get(url, stream=True, headers={'User-Agent': 'Mozilla/5.0'})
                            if response.status_code == 200:
                                with open(img_path, 'wb') as f:
                                    response.raw.decode_content = True
                                    shutil.copyfileobj(response.raw, f)
                                
                                doc.add_picture(img_path, width=Inches(3))
                                doc.add_paragraph(f"Source: {col}")
                                os.remove(img_path)
                            else:
                                doc.add_paragraph(f"[Image download failed: Status {response.status_code}]")
                        except Exception as img_err:
                            print(f"Image error: {img_err}")
                            doc.add_paragraph(f"[Image load error: {col}]")
                
                doc.add_page_break()
                
            doc.save(save_path)
            messagebox.showinfo("Success", "Passports Generated Successfully!")
            
        except Exception as e:
            messagebox.showerror("Passport Error", str(e))

    def export_geojson(self):
        if self.df is None: return
        import json
        save_path = filedialog.asksaveasfilename(defaultextension=".geojson", filetypes=[("GeoJSON", "*.geojson")])
        if not save_path: return
        
        # Check for Lat/Lon
        lat_col = next((c for c in self.df.columns if "Lat" in c), None)
        lon_col = next((c for c in self.df.columns if "Lon" in c), None)
        if not lat_col or not lon_col: return messagebox.showwarning("Error", "Missing 'Lat'/'Lon' columns.")

        features = []
        for _, row in self.filtered_df.iterrows():
            props = row.to_dict()
            clean_props = {k: v for k, v in props.items() if isinstance(v, (str, int, float, bool))}
            
            feature = {
                "type": "Feature",
                "geometry": {
                    "type": "Point",
                    "coordinates": [row[lon_col], row[lat_col]]
                },
                "properties": clean_props
            }
            features.append(feature)
        
        geojson = {"type": "FeatureCollection", "features": features}
        try:
            with open(save_path, 'w') as f: json.dump(geojson, f)
            messagebox.showinfo("Success", f"Exported {len(features)} points.")
        except Exception as e: messagebox.showerror("Error", str(e))

    def download_images_thread(self):
        threading.Thread(target=self.download_images).start()

    def download_images(self):
        if self.df is None: return
        import requests
        
        folder = filedialog.askdirectory(title="Select Folder")
        if not folder: return
        
        count = 0
        try:
            for index, row in self.filtered_df.iterrows():
                id_col = next((c for c in self.df.columns if "ID" in str(c)), None)
                build_id = str(row[id_col]) if id_col else "Unknown"
                
                for col in self.cols:
                    url = self.link_map.get((index, col))
                    if not url and str(row[col]).startswith('http'): url = row[col]
                    
                    if url and "http" in str(url):
                        b_folder = os.path.join(folder, build_id)
                        os.makedirs(b_folder, exist_ok=True)
                        fname = url.split('/')[-1]
                        try:
                            # FIX: Use requests here too
                            response = requests.get(url, stream=True, headers={'User-Agent': 'Mozilla/5.0'})
                            if response.status_code == 200:
                                with open(os.path.join(b_folder, fname), 'wb') as f:
                                    response.raw.decode_content = True
                                    shutil.copyfileobj(response.raw, f)
                                count += 1
                        except: pass
            
            messagebox.showinfo("Download Complete", f"Downloaded {count} images.")
        except Exception as e:
            messagebox.showerror("Download Error", str(e))

    def generate_map(self):
        if self.df is None: return
        import webbrowser
        
        lat_col = next((c for c in self.df.columns if "Lat" in c), None)
        lon_col = next((c for c in self.df.columns if "Lon" in c), None)
        
        if not lat_col or not lon_col:
            messagebox.showwarning("Missing Data", "Could not find 'Lat' or 'Lon' columns.")
            return

        try:
            import folium
            from folium.plugins import MarkerCluster
            
            center_lat = self.filtered_df[lat_col].mean()
            center_lon = self.filtered_df[lon_col].mean()
            m = folium.Map(location=[center_lat, center_lon], zoom_start=12)
            marker_cluster = MarkerCluster().add_to(m)

            for _, row in self.filtered_df.iterrows():
                try:
                    popup_txt = f"ID: {row.get('Building ID', 'Unknown')}<br>Risk: {row.get('Risk Score', 'N/A')}"
                    folium.Marker(
                        location=[row[lat_col], row[lon_col]],
                        popup=popup_txt,
                        icon=folium.Icon(color="red" if row.get('Risk Score', 0) > 20 else "blue")
                    ).add_to(marker_cluster)
                except: continue

            map_path = os.path.abspath("seismic_map.html")
            m.save(map_path)
            webbrowser.open('file://' + map_path)
            
        except ImportError:
            messagebox.showerror("Dependency Error", "Please install folium")
        except Exception as e:
            messagebox.showerror("Map Error", str(e))

    def open_link(self, url):
        import webbrowser
        webbrowser.open(url)

    def render_cards(self, dataframe):
        import pandas as pd
        for widget in self.scroll_frame.winfo_children(): widget.destroy()

        if dataframe.empty:
            ctk.CTkLabel(self.scroll_frame, text="No records match your query.").pack(pady=20)
            return

        for index, row in dataframe.iterrows():
            card = ctk.CTkFrame(self.scroll_frame, fg_color=WHITE, corner_radius=15, border_width=2, border_color=GREY)
            card.pack(fill="x", padx=10, pady=8)

            id_col = next((c for c in self.df.columns if "ID" in str(c)), dataframe.columns[0])
            header = ctk.CTkFrame(card, fg_color=NAVY, height=35, corner_radius=10)
            header.pack(fill="x")
            ctk.CTkLabel(header, text=f"{str(row[id_col])}", text_color=WHITE, font=("Arial", 12, "bold")).pack(side="left", padx=15, pady=5)
            
            content = ctk.CTkFrame(card, fg_color="transparent")
            content.pack(fill="x", padx=15, pady=10)
            
            col_idx = 0
            row_idx = 0
            
            for col in self.cols:
                raw_val = row[col]
                val = str(raw_val)
                if pd.isna(raw_val) or val == "nan": continue

                url = self.link_map.get((index, col)) 
                # Fallback: check original DF index if filtering messed it up
                if not url:
                     # Attempt to find by ID matching or just check value text
                     if str(val).startswith('http'): url = val

                is_link = url is not None
                
                ctk.CTkLabel(content, text=col.upper(), font=("Arial", 9, "bold"), text_color="#888888").grid(row=row_idx, column=col_idx, sticky="w", padx=(0,5))
                
                if is_link:
                    btn_text = val if len(val) < 20 else "📷 View Photo"
                    link_btn = ctk.CTkButton(content, text=f"📷 {btn_text}", width=120, height=24, 
                                             fg_color=GREEN, text_color=NAVY, font=("Arial", 10, "bold"),
                                             command=lambda u=url: self.open_link(u))
                    link_btn.grid(row=row_idx+1, column=col_idx, sticky="w", padx=(0,20), pady=(0, 10))
                else:
                    ctk.CTkLabel(content, text=val[:40] + ("..." if len(val)>40 else ""), 
                                 font=("Arial", 11), text_color="black").grid(row=row_idx+1, column=col_idx, sticky="w", padx=(0,20), pady=(0, 10))
                
                col_idx += 1
                if col_idx > 3: 
                    col_idx = 0
                    row_idx += 2

    def run_comparison(self):
        id1 = self.comp_sel1.get()
        id2 = self.comp_sel2.get()
        if id1 == id2:
            messagebox.showwarning("Same ID", "Select different buildings.")
            return
        id_col = next((c for c in self.df.columns if "ID" in str(c)), self.df.columns[0])
        row1 = self.df[self.df[id_col].astype(str) == id1].iloc[0]
        row2 = self.df[self.df[id_col].astype(str) == id2].iloc[0]
        for widget in self.comp_scroll.winfo_children(): widget.destroy()
        h_frame = ctk.CTkFrame(self.comp_scroll, fg_color="transparent")
        h_frame.pack(fill="x", pady=5)
        ctk.CTkLabel(h_frame, text="Feature", font=("Arial", 12, "bold"), width=150, anchor="w").grid(row=0, column=0)
        ctk.CTkLabel(h_frame, text=id1, font=("Arial", 12, "bold"), text_color=NAVY, width=200, anchor="w").grid(row=0, column=1)
        ctk.CTkLabel(h_frame, text=id2, font=("Arial", 12, "bold"), text_color=MAROON, width=200, anchor="w").grid(row=0, column=2)
        for col in self.cols:
            val1 = str(row1[col])
            val2 = str(row2[col])
            is_diff = val1 != val2
            bg = "#FFEEEE" if is_diff else WHITE
            row_frame = ctk.CTkFrame(self.comp_scroll, fg_color=bg)
            row_frame.pack(fill="x", pady=2)
            ctk.CTkLabel(row_frame, text=col, width=150, anchor="w", text_color="gray").grid(row=0, column=0, padx=5)
            ctk.CTkLabel(row_frame, text=val1[:30], width=200, anchor="w").grid(row=0, column=1, padx=5)
            ctk.CTkLabel(row_frame, text=val2[:30], width=200, anchor="w").grid(row=0, column=2, padx=5)

    def generate_chart(self):
        import matplotlib.pyplot as plt
        from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg, NavigationToolbar2Tk
        import mplcursors # REQUIRED: pip install mplcursors

        col = self.chart_col.get()
        kind = self.chart_type.get()
        
        for widget in self.chart_frame.winfo_children(): widget.destroy()

        # FIX: Ensure we use only valid data
        data_counts = self.filtered_df[col].value_counts().nlargest(15)

        # Standard Matplotlib Figure
        fig, ax = plt.subplots(figsize=(8, 5), dpi=100)
        
        if kind == "Heatmap":
            dist_col = next((c for c in self.df.columns if "District" in c), None)
            if dist_col and dist_col != col:
                ct = pd.crosstab(self.filtered_df[dist_col], self.filtered_df[col])
                im = ax.imshow(ct, cmap='YlOrRd', aspect='auto')
                ax.set_xticks(range(len(ct.columns)))
                ax.set_yticks(range(len(ct.index)))
                ax.set_xticklabels(ct.columns, rotation=45)
                ax.set_yticklabels(ct.index)
                plt.colorbar(im)
            else:
                self.chart_type.set("Bar")
                return self.generate_chart()
        elif kind == "Correlation":
             numeric_df = self.filtered_df.select_dtypes(include=['float64', 'int64'])
             if numeric_df.empty: return
             corr = numeric_df.corr()
             im = ax.imshow(corr, cmap='coolwarm')
             plt.colorbar(im)
        elif kind == "Bar":
            bars = data_counts.plot(kind='bar', color=NAVY, ax=ax, edgecolor='black')
            # FIX: Add Hover Cursors
            mplcursors.cursor(bars, hover=True).connect("add", lambda sel: sel.annotation.set_text(f"{sel.target[1]:.0f}"))
            plt.xticks(rotation=45, ha='right')
        elif kind == "Pie":
            wedges, _ = ax.pie(data_counts.values, startangle=90, colors=plt.cm.Pastel1.colors)
            ax.legend(wedges, data_counts.index, title=col, loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))
            ax.set_ylabel("")
        elif kind == "Line":
             lines = self.filtered_df[col].value_counts().sort_index().plot(kind='line', color=MAROON, marker='o', ax=ax)
             mplcursors.cursor(lines, hover=True)

        ax.set_title(f"Analysis: {col}", fontsize=12, fontweight='bold', color=NAVY)
        plt.tight_layout()

        # Embed in Tkinter WITH TOOLBAR
        canvas = FigureCanvasTkAgg(fig, master=self.chart_frame)
        canvas.draw()
        
        toolbar = NavigationToolbar2Tk(canvas, self.chart_frame)
        toolbar.update()
        
        canvas.get_tk_widget().pack(fill="both", expand=True)

if __name__ == "__main__":
    try:
        app = SeismicAnalyzerApp()
        app.mainloop()
    except Exception as e:
        ctypes.windll.user32.MessageBoxW(0, f"Critical Error:\n{traceback.format_exc()}", "App Crash", 0x10)